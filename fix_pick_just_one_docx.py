#!/usr/bin/env python3
"""
修復 "If You Had to Pick Just One" DOCX 圖片插入問題
處理 HH:MM:SS 格式與圖片檔名的對應
"""

import os
import sys
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
import glob

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


class PickJustOneImageFixer:
    def __init__(self):
        self.slide_folders = []
        self.all_images = {}  # {time_seconds: (folder_index, img_path)}
        
    def add_slide_folder(self, folder_path):
        """添加投影片資料夾"""
        if os.path.exists(folder_path):
            self.slide_folders.append(folder_path)
            logger.info(f"✓ 添加資料夾: {os.path.basename(folder_path)}")
            return True
        else:
            logger.error(f"✗ 資料夾不存在: {folder_path}")
            return False
    
    def parse_time_hhmmss(self, time_str):
        """解析 HH:MM:SS 或 MM:SS 格式"""
        time_str = time_str.strip('（）()[]')
        
        parts = time_str.split(':')
        try:
            if len(parts) == 3:  # HH:MM:SS
                hours = int(parts[0])
                minutes = int(parts[1])
                seconds = int(parts[2])
                return hours * 3600 + minutes * 60 + seconds
            elif len(parts) == 2:  # MM:SS
                minutes = int(parts[0])
                seconds = int(parts[1])
                return minutes * 60 + seconds
        except:
            pass
        return None
    
    def parse_filename_time(self, filename):
        """從檔名解析時間 (tXmYs 格式)"""
        match = re.search(r't(\d+)m([\d.]+)s', filename)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        return None
    
    def load_all_images(self):
        """載入所有資料夾的圖片"""
        total_images = 0
        
        for folder_idx, folder in enumerate(self.slide_folders):
            folder_name = os.path.basename(folder)
            image_count = 0
            
            # 掃描圖片
            for pattern in ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.PNG']:
                for img_path in glob.glob(os.path.join(folder, pattern)):
                    filename = os.path.basename(img_path)
                    time_sec = self.parse_filename_time(filename)
                    
                    if time_sec is not None:
                        self.all_images[time_sec] = (folder_idx, img_path)
                        image_count += 1
            
            logger.info(f"  {folder_name}: {image_count} 張圖片")
            total_images += image_count
        
        logger.info(f"\n總計: {total_images} 張圖片")
        return total_images > 0
    
    def find_image_for_timestamp(self, time_str):
        """根據時間戳找到對應的圖片"""
        target_seconds = self.parse_time_hhmmss(time_str)
        if target_seconds is None:
            return None
        
        # 找最接近的圖片（容差30秒）
        best_match = None
        min_diff = float('inf')
        
        for img_time, (folder_idx, img_path) in self.all_images.items():
            diff = abs(img_time - target_seconds)
            if diff < min_diff and diff <= 30:
                min_diff = diff
                best_match = img_path
        
        return best_match
    
    def fix_docx(self, docx_path, output_path=None):
        """修復 DOCX 文件"""
        logger.info(f"\n📄 處理 DOCX: {docx_path}")
        
        # 載入所有圖片
        if not self.load_all_images():
            logger.error("沒有找到任何圖片！")
            return False
        
        # 處理 DOCX
        doc = Document(docx_path)
        new_doc = Document()
        
        timestamps_found = 0
        images_inserted = 0
        missing_images = []
        
        for para in doc.paragraphs:
            text = para.text
            
            # 檢查是否包含圖片標記
            if '🖼️' in text and '（' in text:
                # 提取時間戳（支援中文和英文括號）
                matches = re.findall(r'[（\(](\d{2}:\d{2}:\d{2})[）\)]', text)
                
                if matches:
                    timestamps_found += len(matches)
                    
                    # 先添加文字
                    new_para = new_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        try:
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                        except:
                            pass
                    
                    # 插入對應的圖片
                    for ts in matches:
                        logger.info(f"\n⏰ 處理時間戳: {ts}")
                        
                        img_path = self.find_image_for_timestamp(ts)
                        if img_path:
                            try:
                                new_doc.add_paragraph()  # 空行
                                new_doc.add_picture(img_path, width=Inches(5.5))
                                new_doc.add_paragraph()  # 空行
                                
                                images_inserted += 1
                                logger.info(f"  ✓ 插入: {os.path.basename(img_path)}")
                            except Exception as e:
                                logger.error(f"  ✗ 插入失敗: {e}")
                        else:
                            logger.warning(f"  ⚠️  找不到對應圖片")
                            missing_images.append(ts)
                else:
                    # 沒有時間戳，直接複製
                    self._copy_paragraph(para, new_doc)
            else:
                # 普通段落
                self._copy_paragraph(para, new_doc)
        
        # 儲存結果
        if output_path is None:
            output_path = Path(docx_path).with_suffix('.fixed_images.docx')
        
        new_doc.save(str(output_path))
        
        # 顯示結果
        logger.info(f"\n📊 處理結果:")
        logger.info(f"  找到時間戳: {timestamps_found} 個")
        logger.info(f"  成功插入: {images_inserted} 張")
        logger.info(f"  缺失圖片: {len(missing_images)} 個")
        
        if missing_images:
            logger.info(f"\n缺失的時間戳:")
            for ts in missing_images:
                seconds = self.parse_time_hhmmss(ts)
                if seconds:
                    logger.info(f"  {ts} ({seconds}秒)")
        
        logger.info(f"\n✓ 輸出檔案: {output_path}")
        
        return True
    
    def _copy_paragraph(self, source_para, target_doc):
        """複製段落"""
        new_para = target_doc.add_paragraph()
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            try:
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
            except:
                pass
        try:
            new_para.alignment = source_para.alignment
        except:
            pass


def main():
    if len(sys.argv) < 2:
        print("用法: python fix_pick_just_one_docx.py <docx檔案> [圖片資料夾1] [圖片資料夾2] ...")
        print("\n範例:")
        print('python fix_pick_just_one_docx.py output.docx "/path/to/1. GLP-1 RA" "/path/to/2. GLP-1GIP RA"')
        print("\n或使用預設路徑:")
        print('python fix_pick_just_one_docx.py output.docx')
        sys.exit(1)
    
    docx_file = sys.argv[1]
    fixer = PickJustOneImageFixer()
    
    if len(sys.argv) > 2:
        # 使用指定的資料夾
        for folder in sys.argv[2:]:
            fixer.add_slide_folder(folder)
    else:
        # 使用預設路徑
        base_path = "/Volumes/WD_BLACK/國際年會/ADA2025/If You Had to Pick Just One"
        default_folders = [
            "1. GLP-1 RA",
            "2. GLP-1GIP RA",
            "3. SGLT2i",
            "4. Rebuttal GLP-1 RA",
            "5. Rebuttal SGLT2i",
            "6. Rebuttal GLPGIP"
        ]
        
        for folder_name in default_folders:
            folder_path = os.path.join(base_path, folder_name)
            fixer.add_slide_folder(folder_path)
    
    # 執行修復
    fixer.fix_docx(docx_file)


if __name__ == "__main__":
    main()