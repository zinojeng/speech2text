#!/usr/bin/env python3
"""
åˆä½µæ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡å…§å®¹è™•ç†ç¨‹å¼
Merge Transcript and Slides Script for ADA 2025 Conference

æ­¤ç¨‹å¼ç”¨æ–¼æ™ºèƒ½åˆä½µæ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡å…§å®¹ï¼š
1. ä»¥æ¼”è¬›è€…å…§å®¹ç‚ºä¸»è»¸ï¼Œä¿ç•™å®Œæ•´æ¼”è¬›å…§å®¹
2. å°‡æŠ•å½±ç‰‡å…§å®¹ä½œç‚ºè£œå……èªªæ˜
3. æ”¯æ´æŠ•å½±ç‰‡åœ–ç‰‡åˆ†æï¼ˆä½¿ç”¨ OpenAI Vision APIï¼‰
4. ä½¿ç”¨ Gemini-2.5-pro é€²è¡Œå…§å®¹æ•´åˆèˆ‡æ½¤ç¨¿
5. ç”Ÿæˆçµæ§‹åŒ–çš„ Markdown å’Œ Word æ–‡ä»¶

ä½¿ç”¨æ–¹æ³•:
    python merge_transcript_slides.py <transcript_file> <slides_file> [slides_images_folder]
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import Optional, Tuple, List, Dict
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from datetime import datetime
import json
import glob
from dotenv import load_dotenv

# è¼‰å…¥ç’°å¢ƒè®Šæ•¸
load_dotenv()

# å°å…¥åœ–ç‰‡åˆ†æåŠŸèƒ½
try:
    from image_analyzer import analyze_image
    IMAGE_ANALYSIS_AVAILABLE = True
except ImportError:
    IMAGE_ANALYSIS_AVAILABLE = False
    logging.warning("image_analyzer æ¨¡çµ„æœªæ‰¾åˆ°ï¼Œåœ–ç‰‡åˆ†æåŠŸèƒ½å°‡ä¸å¯ç”¨")

# è¨­å®šæ—¥èªŒ
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('merge_transcript_slides.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# API é…ç½®
GOOGLE_API_KEY = 'AIzaSyBUNvJo_D2KZV3UVVgQxvFlZC1aFfXIw9k'
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

# ç³»çµ±æç¤ºè©
SYSTEM_PROMPT = """ä½ æ˜¯ä¸€ä½å°ˆæ¥­çš„é†«å­¸æœƒè­°å…§å®¹ç·¨è¼¯ï¼Œå°ˆç²¾æ–¼æ•´åˆæ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡å…§å®¹ã€‚é€™æ˜¯American Diabetes Association 2025å¹´æœƒçš„å…§å®¹ã€‚

ä½ çš„ä»»å‹™æ˜¯æ™ºèƒ½åˆä½µæ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡å…§å®¹ï¼Œå‰µå»ºä¸€ä»½å®Œæ•´ã€æµæš¢çš„æœƒè­°ç­†è¨˜ã€‚

**æ•´åˆåŸå‰‡ï¼š**

1. **ä»¥æ¼”è¬›è€…å…§å®¹ç‚ºä¸»è»¸**ï¼š
   - ä¿ç•™æ¼”è¬›è€…çš„å®Œæ•´è«–è¿°ï¼ŒåŒ…å«å£èªè¡¨é”çš„ç”Ÿå‹•æ€§
   - å°æ¼”è¬›å…§å®¹é€²è¡Œé©åº¦æ½¤ç¨¿ï¼Œä½¿å…¶æ›´åŠ æµæš¢å°ˆæ¥­
   - ä¿æŒæ¼”è¬›è€…çš„è§€é»å’Œå¼·èª¿é‡é»

2. **æŠ•å½±ç‰‡å…§å®¹çš„é‹ç”¨**ï¼š
   - ä½œç‚ºè£œå……èªªæ˜ï¼Œè±å¯Œæ¼”è¬›å…§å®¹
   - å¡«è£œæ¼”è¬›ä¸­æœªæåŠä½†é‡è¦çš„è³‡è¨Š
   - æä¾›å…·é«”æ•¸æ“šã€åœ–è¡¨èªªæ˜æˆ–åƒè€ƒæ–‡ç»
   - ä¸è¦é‡è¤‡å·²åœ¨æ¼”è¬›ä¸­è©³ç´°èªªæ˜çš„å…§å®¹

3. **å…§å®¹çµ„ç¹”æ–¹å¼**ï¼š
   - æŒ‰ç…§æ¼”è¬›çš„é‚è¼¯é †åºçµ„ç¹”å…§å®¹
   - åœ¨é©ç•¶ä½ç½®æ’å…¥æŠ•å½±ç‰‡çš„è£œå……è³‡è¨Š
   - ä½¿ç”¨ __åº•ç·š__ æ¨™è¨˜å»¶ä¼¸è§£è®€æˆ–é‡è¦è£œå……èªªæ˜
   - ä¿æŒæ®µè½ä¹‹é–“çš„æµæš¢éæ¸¡

4. **æ ¼å¼è¦æ±‚**ï¼š
   - ä½¿ç”¨ # ## ### çµ„ç¹”ç« ç¯€
   - **ç²—é«”**æ¨™è¨˜é‡è¦æ¦‚å¿µã€è—¥ç‰©åç¨±ã€é—œéµæ•¸æ“š
   - __åº•ç·š__ç”¨æ–¼æ¨™è¨˜ï¼š
     * ä¾†è‡ªæŠ•å½±ç‰‡çš„é‡è¦è£œå……
     * å»¶ä¼¸è§£è®€å’Œæ·±å…¥èªªæ˜
     * é—œéµç ”ç©¶ç™¼ç¾æˆ–çµè«–
   - é©åº¦ä½¿ç”¨é …ç›®åˆ—è¡¨å‘ˆç¾ä¸¦åˆ—è³‡è¨Š

5. **å¯«ä½œé¢¨æ ¼**ï¼š
   - ä¿æŒå­¸è¡“å°ˆæ¥­æ€§ï¼Œä½†ä¸å¤±æ¼”è¬›çš„ç”Ÿå‹•æ€§
   - ç¢ºä¿å…§å®¹å®Œæ•´ã€è©³ç´°ã€æ˜“è®€
   - ä½¿ç”¨ç¹é«”ä¸­æ–‡

**ç¯„ä¾‹æ•´åˆæ–¹å¼ï¼š**
æ¼”è¬›è€…èªªï¼šã€Œé€™å€‹ç ”ç©¶é¡¯ç¤ºäº†é¡¯è‘—çš„æ”¹å–„æ•ˆæœ...ã€
æŠ•å½±ç‰‡è£œå……ï¼šå…·é«”æ•¸æ“šç‚ºæ”¹å–„ç‡é”åˆ°78.5% (p<0.001)
æ•´åˆå¾Œï¼šã€Œé€™å€‹ç ”ç©¶é¡¯ç¤ºäº†é¡¯è‘—çš„æ”¹å–„æ•ˆæœï¼Œ__æ ¹æ“šæŠ•å½±ç‰‡æ•¸æ“šï¼Œå…·é«”æ”¹å–„ç‡é”åˆ°78.5% (p<0.001)ï¼Œé€™å€‹çµæœåœ¨çµ±è¨ˆå­¸ä¸Šå…·æœ‰é«˜åº¦é¡¯è‘—æ€§__ã€‚ã€

è«‹ç¢ºä¿è¼¸å‡ºæ˜¯ä¸€ä»½å®Œæ•´ã€å°ˆæ¥­ã€è³‡è¨Šè±å¯Œçš„æœƒè­°ç­†è¨˜ã€‚

**åœ–ç‰‡æ•´åˆæ–¹å¼ï¼š**
ç•¶æä¾›æŠ•å½±ç‰‡åœ–ç‰‡æ™‚ï¼Œè«‹ï¼š
1. æ ¹æ“šæ™‚é–“æˆ³è¨˜åœ¨é©ç•¶ä½ç½®æ’å…¥åœ–ç‰‡åˆ†æ
2. ä½¿ç”¨ä»¥ä¸‹æ ¼å¼ï¼š
   > ğŸ–¼ï¸ **æŠ•å½±ç‰‡åœ–è¡¨èªªæ˜**ï¼ˆ[æ™‚é–“]ï¼‰ï¼š
   > [åœ–ç‰‡åˆ†æå…§å®¹]
   > __[èˆ‡æ¼”è¬›å…§å®¹çš„é—œè¯æˆ–å»¶ä¼¸è§£è®€]__
3. ç¢ºä¿åœ–ç‰‡åˆ†æèˆ‡æ¼”è¬›å…§å®¹ç›¸äº’å‘¼æ‡‰
4. ä¸è¦é‡è¤‡å·²åœ¨æ¼”è¬›ä¸­è©³ç´°èªªæ˜çš„åœ–è¡¨å…§å®¹"""


class TranscriptSlidesProcessor:
    """è™•ç†æ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡åˆä½µçš„è™•ç†å™¨"""
    
    def __init__(self):
        """åˆå§‹åŒ–è™•ç†å™¨"""
        self.setup_api()
        self.image_analyses = {}  # å„²å­˜åœ–ç‰‡åˆ†æçµæœ
    
    def setup_api(self):
        """è¨­å®š Google Gemini API"""
        try:
            genai.configure(api_key=GOOGLE_API_KEY)
            logger.info("Google Gemini API è¨­å®šå®Œæˆ")
        except Exception as e:
            logger.error(f"API è¨­å®šå¤±æ•—: {e}")
            raise
    
    def read_file(self, file_path: str, file_type: str) -> str:
        """
        è®€å–æª”æ¡ˆå…§å®¹
        
        Args:
            file_path: æª”æ¡ˆè·¯å¾‘
            file_type: æª”æ¡ˆé¡å‹æè¿°ï¼ˆç”¨æ–¼æ—¥èªŒï¼‰
            
        Returns:
            æª”æ¡ˆå…§å®¹
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"{file_type}æª”æ¡ˆä¸å­˜åœ¨: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            if not content:
                raise ValueError(f"{file_type}æª”æ¡ˆå…§å®¹ç‚ºç©º")
            
            logger.info(f"æˆåŠŸè®€å–{file_type}æª”æ¡ˆ: {file_path}")
            logger.info(f"æª”æ¡ˆå¤§å°: {len(content)} å­—å…ƒ")
            return content
            
        except Exception as e:
            logger.error(f"è®€å–{file_type}æª”æ¡ˆå¤±æ•—: {e}")
            raise
    
    def parse_slide_time(self, filename: str) -> Optional[float]:
        """
        å¾æª”åè§£ææ™‚é–“æˆ³è¨˜
        ä¾‹å¦‚: slide_009_t1m4.7s.jpg -> 64.7 ç§’
        """
        import re
        match = re.search(r't(\d+)m([\d.]+)s', filename)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        return None
    
    def load_slide_images(self, images_folder: str) -> Dict[float, str]:
        """
        è¼‰å…¥æŠ•å½±ç‰‡åœ–ç‰‡ä¸¦æŒ‰æ™‚é–“æ’åº
        
        Returns:
            Dict[float, str]: æ™‚é–“æˆ³è¨˜åˆ°åœ–ç‰‡è·¯å¾‘çš„æ˜ å°„
        """
        images = {}
        if not os.path.exists(images_folder):
            logger.warning(f"åœ–ç‰‡è³‡æ–™å¤¾ä¸å­˜åœ¨: {images_folder}")
            return images
        
        # æ”¯æ´çš„åœ–ç‰‡æ ¼å¼
        patterns = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']
        
        for pattern in patterns:
            for img_path in glob.glob(os.path.join(images_folder, pattern)):
                time_sec = self.parse_slide_time(os.path.basename(img_path))
                if time_sec is not None:
                    images[time_sec] = img_path
        
        logger.info(f"è¼‰å…¥äº† {len(images)} å¼µæŠ•å½±ç‰‡åœ–ç‰‡")
        return images
    
    def merge_with_gemini(self, transcript: str, slides: str, images_folder: Optional[str] = None) -> str:
        """
        ä½¿ç”¨ Gemini-2.5-pro é€²è¡Œå…§å®¹åˆä½µèˆ‡æ•´åˆ
        
        Args:
            transcript: æ¼”è¬›ç¨¿å…§å®¹
            slides: æŠ•å½±ç‰‡å…§å®¹
            
        Returns:
            æ•´åˆå¾Œçš„ Markdown å…§å®¹
        """
        try:
            logger.info("é–‹å§‹ä½¿ç”¨ Gemini-2.5-pro é€²è¡Œå…§å®¹æ•´åˆ")
            
            # å»ºç«‹æ¨¡å‹
            model = genai.GenerativeModel('gemini-2.5-pro')
            
            # è¼‰å…¥åœ–ç‰‡è³‡è¨Š
            image_info = ""
            if images_folder:
                slide_images = self.load_slide_images(images_folder)
                if slide_images:
                    image_info = f"\n\n=== æŠ•å½±ç‰‡åœ–ç‰‡è³‡è¨Š ===\nå…±æœ‰ {len(slide_images)} å¼µæŠ•å½±ç‰‡åœ–ç‰‡ï¼Œæ™‚é–“ç¯„åœå¾ {min(slide_images.keys()):.1f} ç§’åˆ° {max(slide_images.keys()):.1f} ç§’ã€‚è«‹åœ¨æ•´åˆå…§å®¹æ™‚ï¼Œåœ¨é©ç•¶çš„æ®µè½ä½ç½®æ¨™è¨˜ [IMAGE: {'{'}time{'}'}] ä¾†æŒ‡ç¤ºæ‡‰è©²æ’å…¥å“ªå€‹æ™‚é–“é»çš„åœ–ç‰‡ã€‚"
            
            # æ§‹å»ºæç¤ºè©
            user_prompt = f"""è«‹æ ¹æ“šä»¥ä¸‹æ¼”è¬›ç¨¿å’ŒæŠ•å½±ç‰‡å…§å®¹ï¼Œå‰µå»ºä¸€ä»½æ•´åˆçš„æœƒè­°ç­†è¨˜ï¼š

=== æ¼”è¬›ç¨¿å…§å®¹ ===
{transcript}

=== æŠ•å½±ç‰‡å…§å®¹ ===
{slides}{image_info}

è«‹æŒ‰ç…§æŒ‡ç¤ºæ•´åˆé€™å…©ä»½å…§å®¹ï¼Œç”Ÿæˆå®Œæ•´çš„ Markdown æ ¼å¼æœƒè­°ç­†è¨˜ã€‚è¨˜ä½æŠ•å½±ç‰‡å…§å®¹æ˜¯ç”¨ä¾†è£œå……å’ŒåŠ å¼·æ¼”è¬›è€…çš„è«–è¿°ï¼Œä¸è¦é‡è¤‡ç›¸åŒå…§å®¹ã€‚"""
            
            # ç”Ÿæˆæ•´åˆå…§å®¹
            response = model.generate_content([
                {"role": "user", "parts": [{"text": SYSTEM_PROMPT}]},
                {"role": "user", "parts": [{"text": user_prompt}]}
            ])
            
            merged_content = response.text
            logger.info(f"å…§å®¹æ•´åˆå®Œæˆï¼Œé•·åº¦: {len(merged_content)} å­—å…ƒ")
            return merged_content
            
        except Exception as e:
            logger.error(f"å…§å®¹æ•´åˆå¤±æ•—: {e}")
            raise
    
    def save_markdown(self, content: str, output_path: str, slide_images: Optional[Dict[float, str]] = None) -> str:
        """
        ä¿å­˜ Markdown æª”æ¡ˆï¼Œä¸¦è™•ç†åœ–ç‰‡æ¨™è¨˜
        
        Args:
            content: Markdown å…§å®¹
            output_path: è¼¸å‡ºè·¯å¾‘
            slide_images: åœ–ç‰‡æ™‚é–“æˆ³è¨˜åˆ°è·¯å¾‘çš„æ˜ å°„
            
        Returns:
            ä¿å­˜çš„æª”æ¡ˆè·¯å¾‘
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # è™•ç†åœ–ç‰‡æ¨™è¨˜ï¼Œæ›¿æ›ç‚ºå¯¦éš›çš„ Markdown åœ–ç‰‡èªæ³•
            if slide_images:
                lines = content.split('\n')
                processed_lines = []
                
                for line in lines:
                    if '[IMAGE:' in line:
                        match = re.search(r'\[IMAGE:\s*([\d.]+)\]', line)
                        if match:
                            target_time = float(match.group(1))
                            # æ‰¾åˆ°æœ€æ¥è¿‘çš„åœ–ç‰‡
                            closest_time = min(slide_images.keys(), key=lambda x: abs(x - target_time))
                            if abs(closest_time - target_time) < 30:  # 30ç§’å®¹å·®
                                img_path = slide_images[closest_time]
                                # è½‰æ›ç‚ºç›¸å°è·¯å¾‘
                                img_relative = os.path.relpath(img_path, output_path.parent)
                                # æ›¿æ›ç‚º Markdown åœ–ç‰‡èªæ³•
                                line = f"![æŠ•å½±ç‰‡ {closest_time:.1f}s]({img_relative})"
                                logger.info(f"æ›¿æ›åœ–ç‰‡æ¨™è¨˜: {target_time}s -> {os.path.basename(img_path)}")
                    processed_lines.append(line)
                
                content = '\n'.join(processed_lines)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown æª”æ¡ˆå·²ä¿å­˜: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"ä¿å­˜ Markdown å¤±æ•—: {e}")
            raise
    
    def markdown_to_docx(self, markdown_text: str, output_path: str, slide_images: Optional[Dict[float, str]] = None) -> bool:
        """
        å°‡ Markdown æ–‡å­—è½‰æ›ç‚ºä¿ç•™æ ¼å¼çš„ Word æ–‡ä»¶
        
        Args:
            markdown_text: Markdown æ ¼å¼æ–‡å­—
            output_path: è¼¸å‡ºæª”æ¡ˆè·¯å¾‘
            
        Returns:
            è½‰æ›æ˜¯å¦æˆåŠŸ
        """
        try:
            logger.info(f"é–‹å§‹è½‰æ›ç‚º Word æ–‡ä»¶: {output_path}")
            
            doc = Document()
            
            # æ·»åŠ æ¨™é¡Œ
            title = doc.add_heading('ADA 2025 æœƒè­°ç­†è¨˜ - æ¼”è¬›èˆ‡æŠ•å½±ç‰‡æ•´åˆç‰ˆ', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # æ·»åŠ æ—¥æœŸ
            date_para = doc.add_paragraph()
            date_para.add_run(f"æ•´åˆæ—¥æœŸ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_paragraph()  # ç©ºè¡Œ
            
            # åˆ†è¡Œè™•ç† Markdown å…§å®¹
            lines = markdown_text.split('\n')
            
            for line in lines:
                # è™•ç†åœ–ç‰‡æ’å…¥æ¨™è¨˜
                if slide_images and '[IMAGE:' in line:
                    match = re.search(r'\[IMAGE:\s*([\d.]+)\]', line)
                    if match:
                        target_time = float(match.group(1))
                        # æ‰¾åˆ°æœ€æ¥è¿‘çš„åœ–ç‰‡
                        closest_time = min(slide_images.keys(), key=lambda x: abs(x - target_time))
                        if abs(closest_time - target_time) < 30:  # 30ç§’å®¹å·®
                            img_path = slide_images[closest_time]
                            if os.path.exists(img_path):
                                try:
                                    doc.add_paragraph()  # ç©ºè¡Œ
                                    doc.add_picture(img_path, width=Inches(5.5))
                                    doc.add_paragraph()  # ç©ºè¡Œ
                                    logger.info(f"æ’å…¥åœ–ç‰‡: {os.path.basename(img_path)} (æ™‚é–“: {closest_time}ç§’)")
                                except Exception as e:
                                    logger.warning(f"æ’å…¥åœ–ç‰‡å¤±æ•—: {e}")
                        continue  # è·³éé€™ä¸€è¡Œï¼Œä¸é¡¯ç¤ºæ¨™è¨˜
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # ç©ºè¡Œ
                    continue
                
                # è™•ç†æ¨™é¡Œ
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    
                    if level == 1:
                        doc.add_heading(title_text, level=1)
                    elif level == 2:
                        doc.add_heading(title_text, level=2)
                    elif level == 3:
                        doc.add_heading(title_text, level=3)
                    else:
                        doc.add_heading(title_text, level=4)
                    
                    continue
                
                # è™•ç†åˆ—è¡¨é …ç›®
                if line.startswith(('- ', '* ', '+ ')):
                    list_text = line[2:].strip()
                    paragraph = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(paragraph, list_text)
                elif re.match(r'^\d+\.\s', line):
                    list_text = re.sub(r'^\d+\.\s', '', line).strip()
                    paragraph = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(paragraph, list_text)
                else:
                    # è™•ç†æ™®é€šæ®µè½
                    paragraph = doc.add_paragraph()
                    self._add_formatted_text(paragraph, line)
            
            # å„²å­˜æ–‡ä»¶
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logger.info(f"Word æ–‡ä»¶å·²å„²å­˜: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"è½‰æ›ç‚º Word æ–‡ä»¶å¤±æ•—: {e}")
            return False
    
    def _add_formatted_text(self, paragraph, text):
        """
        è™•ç†æ–‡å­—æ ¼å¼ä¸¦æ·»åŠ åˆ°æ®µè½
        
        Args:
            paragraph: Word æ®µè½å°è±¡
            text: è¦è™•ç†çš„æ–‡å­—
        """
        # ä½¿ç”¨æ›´ç²¾ç¢ºçš„æ­£å‰‡è¡¨é”å¼ä¾†è™•ç†æ ¼å¼
        # å…ˆè™•ç†é›™åº•ç·šï¼ˆå¿…é ˆåœ¨å–®åº•ç·šä¹‹å‰ï¼‰
        parts = re.split(r'(__[^_]+__|_[^_]+_|\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # ç²—é«”
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__') and len(part) > 4:
                # åº•ç·šï¼ˆé›™åº•ç·šï¼‰
                run = paragraph.add_run(part[2:-2])
                run.underline = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                # åº•ç·šï¼ˆå–®åº•ç·šï¼‰
                run = paragraph.add_run(part[1:-1])
                run.underline = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                # æ–œé«”
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # æ™®é€šæ–‡å­—
                paragraph.add_run(part)
    
    def process_files(self, transcript_file: str, slides_file: str, output_base: str = None, images_folder: str = None) -> dict:
        """
        è™•ç†æ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡æª”æ¡ˆ
        
        Args:
            transcript_file: æ¼”è¬›ç¨¿æª”æ¡ˆè·¯å¾‘
            slides_file: æŠ•å½±ç‰‡æª”æ¡ˆè·¯å¾‘
            output_base: è¼¸å‡ºæª”æ¡ˆåŸºç¤åç¨±ï¼ˆå¯é¸ï¼‰
            images_folder: æŠ•å½±ç‰‡åœ–ç‰‡è³‡æ–™å¤¾è·¯å¾‘ï¼ˆå¯é¸ï¼‰
            
        Returns:
            è™•ç†çµæœ
        """
        result = {
            'success': False,
            'transcript_file': transcript_file,
            'slides_file': slides_file,
            'markdown_file': None,
            'docx_file': None,
            'error': None
        }
        
        try:
            # è®€å–æª”æ¡ˆ
            transcript = self.read_file(transcript_file, "æ¼”è¬›ç¨¿")
            slides = self.read_file(slides_file, "æŠ•å½±ç‰‡")
            
            # è¼‰å…¥åœ–ç‰‡ï¼ˆå¦‚æœæœ‰æä¾›ï¼‰
            slide_images = None
            if images_folder:
                slide_images = self.load_slide_images(images_folder)
            
            # ä½¿ç”¨ Gemini é€²è¡Œå…§å®¹æ•´åˆ
            merged_content = self.merge_with_gemini(transcript, slides, images_folder)
            
            # æº–å‚™è¼¸å‡ºè·¯å¾‘
            if not output_base:
                # ä½¿ç”¨æ¼”è¬›ç¨¿æª”æ¡ˆåä½œç‚ºåŸºç¤
                transcript_path = Path(transcript_file)
                output_base = transcript_path.stem
            
            output_dir = Path(transcript_file).parent
            
            # ä¿å­˜ Markdownï¼ˆåŒ…å«åœ–ç‰‡è™•ç†ï¼‰
            markdown_path = output_dir / f"{output_base}_merged.md"
            self.save_markdown(merged_content, str(markdown_path), slide_images)
            result['markdown_file'] = str(markdown_path)
            
            # è½‰æ›ç‚º Word
            docx_path = output_dir / f"{output_base}_merged.docx"
            if self.markdown_to_docx(merged_content, str(docx_path), slide_images):
                result['docx_file'] = str(docx_path)
                result['success'] = True
            else:
                result['error'] = "Word è½‰æ›å¤±æ•—"
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"è™•ç†æª”æ¡ˆå¤±æ•—: {e}")
        
        return result


def main():
    """ä¸»å‡½æ•¸"""
    import argparse
    
    parser = argparse.ArgumentParser(description='åˆä½µæ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡å…§å®¹')
    parser.add_argument('transcript_file', help='æ¼”è¬›ç¨¿æª”æ¡ˆ')
    parser.add_argument('slides_file', help='æŠ•å½±ç‰‡å…§å®¹æª”æ¡ˆ')
    parser.add_argument('output_base', nargs='?', help='è¼¸å‡ºæª”æ¡ˆåŸºç¤åç¨±')
    parser.add_argument('--images', help='æŠ•å½±ç‰‡åœ–ç‰‡è³‡æ–™å¤¾è·¯å¾‘')
    
    args = parser.parse_args()
    
    transcript_file = args.transcript_file
    slides_file = args.slides_file
    output_base = args.output_base
    images_folder = args.images
    
    print(f"\n=== æ¼”è¬›ç¨¿èˆ‡æŠ•å½±ç‰‡æ•´åˆç¨‹å¼ ===")
    print(f"æ¼”è¬›ç¨¿æª”æ¡ˆ: {transcript_file}")
    print(f"æŠ•å½±ç‰‡æª”æ¡ˆ: {slides_file}")
    if images_folder:
        print(f"åœ–ç‰‡è³‡æ–™å¤¾: {images_folder}")
    print(f"ä½¿ç”¨æ¨¡å‹: Gemini-2.5-pro")
    print("è™•ç†ä¸­...\n")
    
    try:
        processor = TranscriptSlidesProcessor()
        result = processor.process_files(transcript_file, slides_file, output_base, images_folder)
        
        if result['success']:
            print("\nâœ… è™•ç†æˆåŠŸï¼")
            print(f"Markdown æª”æ¡ˆ: {result['markdown_file']}")
            print(f"Word æª”æ¡ˆ: {result['docx_file']}")
        else:
            print(f"\nâŒ è™•ç†å¤±æ•—: {result['error']}")
            sys.exit(1)
            
    except Exception as e:
        print(f"\nâŒ ç¨‹å¼åŸ·è¡Œå¤±æ•—: {e}")
        logger.error(f"ç¨‹å¼åŸ·è¡Œå¤±æ•—: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()