"""
批次音訊處理系統 - 報告生成器
Report Generator for Batch Audio Processing System

此模組負責生成處理結果的詳細報告，包括：
- JSON 格式詳細報告
- Markdown 格式摘要報告
- Word 格式報告（包含統計圖表）

Requirements: 6.3, 6.4
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import asdict

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx 未安裝，Word 報告功能將不可用")

from batch_audio_models import FileInfo, TranscriptionResult, SummaryResult
from batch_processing_orchestrator import FileProcessingResult, BatchProcessingResult

# 設定日誌
logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    報告生成器類別
    
    負責生成各種格式的處理報告
    Requirements: 6.3, 6.4
    """
    
    def __init__(self, output_dir: str = "reports"):
        """
        初始化報告生成器
        
        Args:
            output_dir: 報告輸出目錄
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        logger.info(f"報告生成器初始化完成，輸出目錄: {self.output_dir}")
    
    def generate_processing_report(self, batch_result: BatchProcessingResult, 
                                 report_name: Optional[str] = None) -> Dict[str, str]:
        """
        生成完整的處理報告
        
        Args:
            batch_result: 批次處理結果
            report_name: 報告名稱（可選）
            
        Returns:
            包含各種格式報告路徑的字典
        """
        if not report_name:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_name = f"batch_processing_report_{timestamp}"
        
        logger.info(f"開始生成處理報告: {report_name}")
        
        report_paths = {}
        
        try:
            # 生成 JSON 詳細報告
            json_path = self.generate_json_report(batch_result, report_name)
            if json_path:
                report_paths['json'] = str(json_path)
            
            # 生成 Markdown 摘要報告
            markdown_path = self.generate_markdown_report(batch_result, report_name)
            if markdown_path:
                report_paths['markdown'] = str(markdown_path)
            
            # 生成 Word 格式報告
            if DOCX_AVAILABLE:
                docx_path = self.generate_docx_report(batch_result, report_name)
                if docx_path:
                    report_paths['docx'] = str(docx_path)
            else:
                logger.warning("跳過 Word 報告生成：python-docx 未安裝")
            
            logger.info(f"處理報告生成完成: {len(report_paths)} 個檔案")
            
        except Exception as e:
            logger.error(f"生成處理報告時發生錯誤: {e}")
            raise
        
        return report_paths
    
    def generate_json_report(self, batch_result: BatchProcessingResult, 
                           report_name: str) -> Optional[Path]:
        """
        生成 JSON 格式詳細報告
        
        Args:
            batch_result: 批次處理結果
            report_name: 報告名稱
            
        Returns:
            JSON 報告檔案路徑
        """
        try:
            logger.info("生成 JSON 格式詳細報告")
            
            # 準備報告資料
            report_data = self._prepare_detailed_report_data(batch_result)
            
            # 寫入 JSON 檔案
            json_path = self.output_dir / f"{report_name}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2, default=str)
            
            logger.info(f"JSON 報告已生成: {json_path}")
            return json_path
            
        except Exception as e:
            logger.error(f"生成 JSON 報告時發生錯誤: {e}")
            return None
    
    def generate_markdown_report(self, batch_result: BatchProcessingResult, 
                               report_name: str) -> Optional[Path]:
        """
        生成 Markdown 格式摘要報告
        
        Args:
            batch_result: 批次處理結果
            report_name: 報告名稱
            
        Returns:
            Markdown 報告檔案路徑
        """
        try:
            logger.info("生成 Markdown 格式摘要報告")
            
            # 生成 Markdown 內容
            markdown_content = self._generate_markdown_content(batch_result)
            
            # 寫入 Markdown 檔案
            markdown_path = self.output_dir / f"{report_name}.md"
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            logger.info(f"Markdown 報告已生成: {markdown_path}")
            return markdown_path
            
        except Exception as e:
            logger.error(f"生成 Markdown 報告時發生錯誤: {e}")
            return None
    
    def _prepare_detailed_report_data(self, batch_result: BatchProcessingResult) -> Dict[str, Any]:
        """
        準備詳細報告資料
        
        Args:
            batch_result: 批次處理結果
            
        Returns:
            詳細報告資料字典
        """
        # 計算統計資訊
        stats = self._calculate_processing_statistics(batch_result)
        
        # 準備檔案處理詳情
        file_details = []
        for result in batch_result.results:
            file_detail = {
                'file_info': {
                    'name': result.file_info.audio_name,
                    'path': result.file_info.audio_path,
                    'size_mb': result.file_info.file_size_mb,
                    'format': result.file_info.file_format,
                    'agenda_path': result.file_info.agenda_path,
                    'has_agenda': result.file_info.agenda_path is not None
                },
                'processing': {
                    'status': result.status.value,
                    'success': result.success,
                    'total_time': result.total_time,
                    'start_time': result.start_time.isoformat() if result.start_time else None,
                    'end_time': result.end_time.isoformat() if result.end_time else None,
                    'error': result.error
                },
                'transcription': self._extract_transcription_data(result.transcription) if result.transcription else None,
                'summary': self._extract_summary_data(result.summary) if result.summary else None,
                'documents': result.documents
            }
            file_details.append(file_detail)
        
        # 組合完整報告
        report_data = {
            'report_metadata': {
                'generated_at': datetime.now().isoformat(),
                'generator_version': '1.0.0',
                'report_type': 'batch_processing_detailed'
            },
            'batch_summary': {
                'total_files': batch_result.total_files,
                'successful_files': batch_result.successful_files,
                'failed_files': batch_result.failed_files,
                'success_rate': batch_result.get_success_rate(),
                'total_processing_time': batch_result.total_processing_time,
                'start_time': batch_result.start_time.isoformat() if batch_result.start_time else None,
                'end_time': batch_result.end_time.isoformat() if batch_result.end_time else None
            },
            'statistics': stats,
            'file_details': file_details
        }
        
        return report_data
    
    def _generate_markdown_content(self, batch_result: BatchProcessingResult) -> str:
        """
        生成 Markdown 報告內容
        
        Args:
            batch_result: 批次處理結果
            
        Returns:
            Markdown 格式的報告內容
        """
        # 計算統計資訊
        stats = self._calculate_processing_statistics(batch_result)
        
        # 開始生成 Markdown 內容
        content = []
        
        # 標題和基本資訊
        content.append("# 批次音訊處理報告")
        content.append("")
        content.append(f"**生成時間**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")
        
        # 處理摘要
        content.append("## 處理摘要")
        content.append("")
        content.append(f"- **總檔案數**: {batch_result.total_files}")
        content.append(f"- **成功處理**: {batch_result.successful_files}")
        content.append(f"- **處理失敗**: {batch_result.failed_files}")
        content.append(f"- **成功率**: {batch_result.get_success_rate():.1%}")
        content.append(f"- **總處理時間**: {batch_result.total_processing_time:.1f} 秒")
        
        if batch_result.start_time and batch_result.end_time:
            duration = (batch_result.end_time - batch_result.start_time).total_seconds()
            content.append(f"- **實際耗時**: {duration:.1f} 秒")
        
        content.append("")
        
        # 統計資訊
        content.append("## 詳細統計")
        content.append("")
        
        # 檔案統計
        content.append("### 檔案統計")
        content.append("")
        content.append(f"- **總檔案大小**: {stats['file_stats']['total_size_mb']:.1f} MB")
        content.append(f"- **平均檔案大小**: {stats['file_stats']['average_size_mb']:.1f} MB")
        content.append(f"- **最大檔案**: {stats['file_stats']['max_size_mb']:.1f} MB")
        content.append(f"- **最小檔案**: {stats['file_stats']['min_size_mb']:.1f} MB")
        content.append(f"- **有議程檔案**: {stats['file_stats']['files_with_agenda']}")
        content.append("")
        
        # 轉錄統計
        if stats['transcription_stats']['total_transcriptions'] > 0:
            content.append("### 轉錄統計")
            content.append("")
            content.append(f"- **轉錄成功**: {stats['transcription_stats']['successful_transcriptions']}")
            content.append(f"- **轉錄失敗**: {stats['transcription_stats']['failed_transcriptions']}")
            content.append(f"- **平均轉錄時間**: {stats['transcription_stats']['average_time']:.1f} 秒")
            content.append(f"- **總內容長度**: {stats['transcription_stats']['total_content_length']:,} 字符")
            content.append("")
        
        # 摘要統計
        if stats['summary_stats']['total_summaries'] > 0:
            content.append("### 摘要統計")
            content.append("")
            content.append(f"- **摘要成功**: {stats['summary_stats']['successful_summaries']}")
            content.append(f"- **摘要失敗**: {stats['summary_stats']['failed_summaries']}")
            content.append(f"- **平均摘要時間**: {stats['summary_stats']['average_time']:.1f} 秒")
            content.append(f"- **使用議程**: {stats['summary_stats']['summaries_with_agenda']}")
            content.append(f"- **插入圖片**: {stats['summary_stats']['total_images_inserted']}")
            content.append("")
        
        # 處理結果詳情
        content.append("## 處理結果詳情")
        content.append("")
        
        # 成功處理的檔案
        successful_results = [r for r in batch_result.results if r.success]
        if successful_results:
            content.append("### ✅ 成功處理的檔案")
            content.append("")
            content.append("| 檔案名稱 | 大小(MB) | 處理時間(秒) | 轉錄 | 摘要 | 文件數 |")
            content.append("|---------|---------|-------------|------|------|--------|")
            
            for result in successful_results:
                transcription_status = "✅" if result.transcription and result.transcription.success else "❌"
                summary_status = "✅" if result.summary and result.summary.success else "❌"
                
                content.append(f"| {result.file_info.audio_name} | "
                             f"{result.file_info.file_size_mb:.1f} | "
                             f"{result.total_time:.1f} | "
                             f"{transcription_status} | "
                             f"{summary_status} | "
                             f"{len(result.documents)} |")
            
            content.append("")
        
        # 失敗處理的檔案
        failed_results = [r for r in batch_result.results if not r.success]
        if failed_results:
            content.append("### ❌ 處理失敗的檔案")
            content.append("")
            content.append("| 檔案名稱 | 大小(MB) | 錯誤訊息 |")
            content.append("|---------|---------|----------|")
            
            for result in failed_results:
                error_msg = result.error or "未知錯誤"
                # 截斷過長的錯誤訊息
                if len(error_msg) > 50:
                    error_msg = error_msg[:47] + "..."
                
                content.append(f"| {result.file_info.audio_name} | "
                             f"{result.file_info.file_size_mb:.1f} | "
                             f"{error_msg} |")
            
            content.append("")
        
        # 效能分析
        content.append("## 效能分析")
        content.append("")
        
        if batch_result.total_files > 0:
            avg_time_per_file = batch_result.total_processing_time / batch_result.total_files
            content.append(f"- **平均每檔案處理時間**: {avg_time_per_file:.1f} 秒")
            
            total_size = sum(r.file_info.file_size_mb for r in batch_result.results)
            if total_size > 0:
                processing_speed = total_size / batch_result.total_processing_time * 60  # MB/分鐘
                content.append(f"- **處理速度**: {processing_speed:.1f} MB/分鐘")
        
        content.append("")
        
        # 建議和注意事項
        content.append("## 建議和注意事項")
        content.append("")
        
        if batch_result.failed_files > 0:
            failure_rate = batch_result.failed_files / batch_result.total_files
            if failure_rate > 0.1:  # 失敗率超過 10%
                content.append("⚠️ **高失敗率警告**: 處理失敗率較高，建議檢查:")
                content.append("- API 金鑰是否有效")
                content.append("- 網路連線是否穩定")
                content.append("- 檔案格式是否支援")
                content.append("")
        
        # 大檔案處理建議
        large_files = [r for r in batch_result.results if r.file_info.file_size_mb > 25]
        if large_files:
            content.append("📁 **大檔案處理**: 發現大型檔案，建議:")
            content.append("- 考慮啟用自動分割功能")
            content.append("- 增加處理超時時間")
            content.append("- 監控記憶體使用情況")
            content.append("")
        
        return "\n".join(content)
    
    def generate_docx_report(self, batch_result: BatchProcessingResult, 
                           report_name: str) -> Optional[Path]:
        """
        生成 Word 格式報告
        
        Args:
            batch_result: 批次處理結果
            report_name: 報告名稱
            
        Returns:
            Word 報告檔案路徑
        """
        if not DOCX_AVAILABLE:
            logger.error("無法生成 Word 報告：python-docx 未安裝")
            return None
        
        try:
            logger.info("生成 Word 格式報告")
            
            # 創建 Word 文件
            doc = Document()
            
            # 設定文件樣式
            self._setup_docx_styles(doc)
            
            # 添加標題
            title = doc.add_heading('批次音訊處理報告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加生成時間
            time_para = doc.add_paragraph()
            time_para.add_run(f"生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # 空行
            
            # 添加處理摘要
            self._add_summary_section(doc, batch_result)
            
            # 添加統計圖表
            self._add_statistics_section(doc, batch_result)
            
            # 添加詳細結果
            self._add_detailed_results_section(doc, batch_result)
            
            # 添加建議和注意事項
            self._add_recommendations_section(doc, batch_result)
            
            # 保存文件
            docx_path = self.output_dir / f"{report_name}.docx"
            doc.save(str(docx_path))
            
            logger.info(f"Word 報告已生成: {docx_path}")
            return docx_path
            
        except Exception as e:
            logger.error(f"生成 Word 報告時發生錯誤: {e}")
            return None
    
    def _setup_docx_styles(self, doc: Document) -> None:
        """設定 Word 文件樣式"""
        try:
            # 設定預設字體
            style = doc.styles['Normal']
            font = style.font
            font.name = '微軟正黑體'
            font.size = Pt(11)
        except Exception as e:
            logger.warning(f"設定文件樣式時發生錯誤: {e}")
    
    def _add_summary_section(self, doc: Document, batch_result: BatchProcessingResult) -> None:
        """添加處理摘要區段"""
        # 摘要標題
        doc.add_heading('處理摘要', level=1)
        
        # 創建摘要表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 設定表格標題
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '項目'
        hdr_cells[1].text = '數值'
        
        # 添加摘要資料
        summary_data = [
            ('總檔案數', str(batch_result.total_files)),
            ('成功處理', str(batch_result.successful_files)),
            ('處理失敗', str(batch_result.failed_files)),
            ('成功率', f"{batch_result.get_success_rate():.1%}"),
            ('總處理時間', f"{batch_result.total_processing_time:.1f} 秒")
        ]
        
        if batch_result.start_time and batch_result.end_time:
            duration = (batch_result.end_time - batch_result.start_time).total_seconds()
            summary_data.append(('實際耗時', f"{duration:.1f} 秒"))
        
        for item, value in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = value
        
        doc.add_paragraph()  # 空行
    
    def _add_statistics_section(self, doc: Document, batch_result: BatchProcessingResult) -> None:
        """添加統計資訊區段"""
        doc.add_heading('詳細統計', level=1)
        
        # 計算統計資訊
        stats = self._calculate_processing_statistics(batch_result)
        
        # 檔案統計
        doc.add_heading('檔案統計', level=2)
        file_stats_table = doc.add_table(rows=1, cols=2)
        file_stats_table.style = 'Table Grid'
        
        hdr_cells = file_stats_table.rows[0].cells
        hdr_cells[0].text = '統計項目'
        hdr_cells[1].text = '數值'
        
        file_stats_data = [
            ('總檔案大小', f"{stats['file_stats']['total_size_mb']:.1f} MB"),
            ('平均檔案大小', f"{stats['file_stats']['average_size_mb']:.1f} MB"),
            ('最大檔案', f"{stats['file_stats']['max_size_mb']:.1f} MB"),
            ('最小檔案', f"{stats['file_stats']['min_size_mb']:.1f} MB"),
            ('有議程檔案', str(stats['file_stats']['files_with_agenda']))
        ]
        
        for item, value in file_stats_data:
            row_cells = file_stats_table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = value
        
        doc.add_paragraph()  # 空行
        
        # 轉錄統計
        if stats['transcription_stats']['total_transcriptions'] > 0:
            doc.add_heading('轉錄統計', level=2)
            trans_stats_table = doc.add_table(rows=1, cols=2)
            trans_stats_table.style = 'Table Grid'
            
            hdr_cells = trans_stats_table.rows[0].cells
            hdr_cells[0].text = '統計項目'
            hdr_cells[1].text = '數值'
            
            trans_stats_data = [
                ('轉錄成功', str(stats['transcription_stats']['successful_transcriptions'])),
                ('轉錄失敗', str(stats['transcription_stats']['failed_transcriptions'])),
                ('平均轉錄時間', f"{stats['transcription_stats']['average_time']:.1f} 秒"),
                ('總內容長度', f"{stats['transcription_stats']['total_content_length']:,} 字符")
            ]
            
            for item, value in trans_stats_data:
                row_cells = trans_stats_table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = value
            
            doc.add_paragraph()  # 空行
        
        # 摘要統計
        if stats['summary_stats']['total_summaries'] > 0:
            doc.add_heading('摘要統計', level=2)
            summary_stats_table = doc.add_table(rows=1, cols=2)
            summary_stats_table.style = 'Table Grid'
            
            hdr_cells = summary_stats_table.rows[0].cells
            hdr_cells[0].text = '統計項目'
            hdr_cells[1].text = '數值'
            
            summary_stats_data = [
                ('摘要成功', str(stats['summary_stats']['successful_summaries'])),
                ('摘要失敗', str(stats['summary_stats']['failed_summaries'])),
                ('平均摘要時間', f"{stats['summary_stats']['average_time']:.1f} 秒"),
                ('使用議程', str(stats['summary_stats']['summaries_with_agenda'])),
                ('插入圖片', str(stats['summary_stats']['total_images_inserted']))
            ]
            
            for item, value in summary_stats_data:
                row_cells = summary_stats_table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = value
            
            doc.add_paragraph()  # 空行
    
    def _add_detailed_results_section(self, doc: Document, batch_result: BatchProcessingResult) -> None:
        """添加詳細結果區段"""
        doc.add_heading('處理結果詳情', level=1)
        
        # 成功處理的檔案
        successful_results = [r for r in batch_result.results if r.success]
        if successful_results:
            doc.add_heading('✅ 成功處理的檔案', level=2)
            
            success_table = doc.add_table(rows=1, cols=6)
            success_table.style = 'Table Grid'
            
            # 設定表格標題
            hdr_cells = success_table.rows[0].cells
            headers = ['檔案名稱', '大小(MB)', '處理時間(秒)', '轉錄', '摘要', '文件數']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # 添加成功處理的檔案資料
            for result in successful_results:
                row_cells = success_table.add_row().cells
                row_cells[0].text = result.file_info.audio_name
                row_cells[1].text = f"{result.file_info.file_size_mb:.1f}"
                row_cells[2].text = f"{result.total_time:.1f}"
                row_cells[3].text = "✅" if result.transcription and result.transcription.success else "❌"
                row_cells[4].text = "✅" if result.summary and result.summary.success else "❌"
                row_cells[5].text = str(len(result.documents))
            
            doc.add_paragraph()  # 空行
        
        # 失敗處理的檔案
        failed_results = [r for r in batch_result.results if not r.success]
        if failed_results:
            doc.add_heading('❌ 處理失敗的檔案', level=2)
            
            failed_table = doc.add_table(rows=1, cols=3)
            failed_table.style = 'Table Grid'
            
            # 設定表格標題
            hdr_cells = failed_table.rows[0].cells
            headers = ['檔案名稱', '大小(MB)', '錯誤訊息']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # 添加失敗處理的檔案資料
            for result in failed_results:
                row_cells = failed_table.add_row().cells
                row_cells[0].text = result.file_info.audio_name
                row_cells[1].text = f"{result.file_info.file_size_mb:.1f}"
                
                error_msg = result.error or "未知錯誤"
                # 截斷過長的錯誤訊息
                if len(error_msg) > 100:
                    error_msg = error_msg[:97] + "..."
                row_cells[2].text = error_msg
            
            doc.add_paragraph()  # 空行
    
    def _add_recommendations_section(self, doc: Document, batch_result: BatchProcessingResult) -> None:
        """添加建議和注意事項區段"""
        doc.add_heading('建議和注意事項', level=1)
        
        recommendations = []
        
        # 高失敗率警告
        if batch_result.failed_files > 0:
            failure_rate = batch_result.failed_files / batch_result.total_files
            if failure_rate > 0.1:  # 失敗率超過 10%
                recommendations.append({
                    'title': '⚠️ 高失敗率警告',
                    'content': [
                        '處理失敗率較高，建議檢查:',
                        '• API 金鑰是否有效',
                        '• 網路連線是否穩定',
                        '• 檔案格式是否支援'
                    ]
                })
        
        # 大檔案處理建議
        large_files = [r for r in batch_result.results if r.file_info.file_size_mb > 25]
        if large_files:
            recommendations.append({
                'title': '📁 大檔案處理建議',
                'content': [
                    '發現大型檔案，建議:',
                    '• 考慮啟用自動分割功能',
                    '• 增加處理超時時間',
                    '• 監控記憶體使用情況'
                ]
            })
        
        # 效能優化建議
        if batch_result.total_files > 5:
            avg_time_per_file = batch_result.total_processing_time / batch_result.total_files
            if avg_time_per_file > 60:  # 平均每檔案超過1分鐘
                recommendations.append({
                    'title': '⚡ 效能優化建議',
                    'content': [
                        '處理時間較長，建議:',
                        '• 考慮啟用並行處理',
                        '• 檢查網路連線速度',
                        '• 優化檔案大小和格式'
                    ]
                })
        
        # 添加建議內容
        if recommendations:
            for rec in recommendations:
                doc.add_heading(rec['title'], level=2)
                for content_line in rec['content']:
                    doc.add_paragraph(content_line)
                doc.add_paragraph()  # 空行
        else:
            doc.add_paragraph('✅ 處理過程順利，無特別建議事項。')
        
        # 添加效能分析
        doc.add_heading('效能分析', level=2)
        
        if batch_result.total_files > 0:
            avg_time_per_file = batch_result.total_processing_time / batch_result.total_files
            perf_para = doc.add_paragraph()
            perf_para.add_run(f"平均每檔案處理時間: {avg_time_per_file:.1f} 秒").bold = True
            
            total_size = sum(r.file_info.file_size_mb for r in batch_result.results)
            if total_size > 0 and batch_result.total_processing_time > 0:
                processing_speed = total_size / batch_result.total_processing_time * 60  # MB/分鐘
                speed_para = doc.add_paragraph()
                speed_para.add_run(f"處理速度: {processing_speed:.1f} MB/分鐘").bold = True
    
    def _calculate_processing_statistics(self, batch_result: BatchProcessingResult) -> Dict[str, Any]:
        """
        計算處理統計資訊
        
        Args:
            batch_result: 批次處理結果
            
        Returns:
            統計資訊字典
        """
        stats = {
            'file_stats': {
                'total_files': len(batch_result.results),
                'total_size_mb': 0.0,
                'average_size_mb': 0.0,
                'max_size_mb': 0.0,
                'min_size_mb': float('inf'),
                'files_with_agenda': 0
            },
            'transcription_stats': {
                'total_transcriptions': 0,
                'successful_transcriptions': 0,
                'failed_transcriptions': 0,
                'total_time': 0.0,
                'average_time': 0.0,
                'total_content_length': 0
            },
            'summary_stats': {
                'total_summaries': 0,
                'successful_summaries': 0,
                'failed_summaries': 0,
                'total_time': 0.0,
                'average_time': 0.0,
                'summaries_with_agenda': 0,
                'total_images_inserted': 0
            }
        }
        
        if not batch_result.results:
            return stats
        
        # 計算檔案統計
        file_sizes = []
        for result in batch_result.results:
            file_size = result.file_info.file_size_mb
            file_sizes.append(file_size)
            stats['file_stats']['total_size_mb'] += file_size
            stats['file_stats']['max_size_mb'] = max(stats['file_stats']['max_size_mb'], file_size)
            stats['file_stats']['min_size_mb'] = min(stats['file_stats']['min_size_mb'], file_size)
            
            if result.file_info.agenda_path:
                stats['file_stats']['files_with_agenda'] += 1
        
        if file_sizes:
            stats['file_stats']['average_size_mb'] = sum(file_sizes) / len(file_sizes)
            if stats['file_stats']['min_size_mb'] == float('inf'):
                stats['file_stats']['min_size_mb'] = 0.0
        
        # 計算轉錄統計
        transcription_times = []
        for result in batch_result.results:
            if result.transcription:
                stats['transcription_stats']['total_transcriptions'] += 1
                if result.transcription.success:
                    stats['transcription_stats']['successful_transcriptions'] += 1
                    transcription_times.append(result.transcription.processing_time)
                    stats['transcription_stats']['total_content_length'] += len(result.transcription.content or "")
                else:
                    stats['transcription_stats']['failed_transcriptions'] += 1
                
                stats['transcription_stats']['total_time'] += result.transcription.processing_time
        
        if transcription_times:
            stats['transcription_stats']['average_time'] = sum(transcription_times) / len(transcription_times)
        
        # 計算摘要統計
        summary_times = []
        for result in batch_result.results:
            if result.summary:
                stats['summary_stats']['total_summaries'] += 1
                if result.summary.success:
                    stats['summary_stats']['successful_summaries'] += 1
                    summary_times.append(result.summary.processing_time)
                    if result.summary.agenda_used:
                        stats['summary_stats']['summaries_with_agenda'] += 1
                    stats['summary_stats']['total_images_inserted'] += result.summary.images_inserted
                else:
                    stats['summary_stats']['failed_summaries'] += 1
                
                stats['summary_stats']['total_time'] += result.summary.processing_time
        
        if summary_times:
            stats['summary_stats']['average_time'] = sum(summary_times) / len(summary_times)
        
        return stats
    
    def _extract_transcription_data(self, transcription: TranscriptionResult) -> Dict[str, Any]:
        """
        提取轉錄結果資料
        
        Args:
            transcription: 轉錄結果
            
        Returns:
            轉錄資料字典
        """
        return {
            'success': transcription.success,
            'content_length': len(transcription.content) if transcription.content else 0,
            'processing_time': transcription.processing_time,
            'token_count': transcription.token_count,
            'model_used': transcription.model_used,
            'language_detected': transcription.language_detected,
            'confidence_score': transcription.confidence_score,
            'segments_processed': transcription.segments_processed,
            'error': transcription.error
        }
    
    def _extract_summary_data(self, summary: SummaryResult) -> Dict[str, Any]:
        """
        提取摘要結果資料
        
        Args:
            summary: 摘要結果
            
        Returns:
            摘要資料字典
        """
        return {
            'success': summary.success,
            'content_length': len(summary.content) if summary.content else 0,
            'processing_time': summary.processing_time,
            'token_count': summary.token_count,
            'model_used': summary.model_used,
            'agenda_used': summary.agenda_used,
            'images_inserted': summary.images_inserted,
            'error': summary.error
        }


if __name__ == "__main__":
    """測試報告生成器功能"""
    import sys
    from batch_audio_models import FileInfo, TranscriptionResult, SummaryResult
    from batch_processing_orchestrator import FileProcessingResult, BatchProcessingResult, ProcessingStatus
    
    # 設定日誌
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    
    print("=== 批次音訊處理系統 - 報告生成器測試 ===\n")
    
    try:
        # 創建測試資料
        print("1. 創建測試資料...")
        
        # 創建檔案資訊
        file_info1 = FileInfo(
            audio_path="test1.mp3",
            audio_name="test1",
            file_size_mb=15.5,
            agenda_path="test1_agenda.txt"
        )
        
        file_info2 = FileInfo(
            audio_path="test2.wav",
            audio_name="test2",
            file_size_mb=32.1
        )
        
        # 創建轉錄結果
        transcription1 = TranscriptionResult(
            success=True,
            content="這是測試轉錄內容1...",
            processing_time=45.2,
            model_used="gpt-transcribe",
            segments_processed=1
        )
        
        transcription2 = TranscriptionResult(
            success=True,
            content="這是測試轉錄內容2，內容較長...",
            processing_time=78.5,
            model_used="gpt-transcribe",
            segments_processed=2
        )
        
        # 創建摘要結果
        summary1 = SummaryResult(
            success=True,
            content="這是測試摘要內容1...",
            processing_time=12.3,
            model_used="gemini-3.7-flash",
            agenda_used=True,
            images_inserted=2
        )
        
        summary2 = SummaryResult(
            success=False,
            error="API 呼叫失敗",
            processing_time=5.1
        )
        
        # 創建檔案處理結果
        result1 = FileProcessingResult(file_info=file_info1)
        result1.transcription = transcription1
        result1.summary = summary1
        result1.documents = {"markdown": "test1.md", "docx": "test1.docx"}
        result1.mark_completed(success=True)
        
        result2 = FileProcessingResult(file_info=file_info2)
        result2.transcription = transcription2
        result2.summary = summary2
        result2.mark_completed(success=False, error="摘要處理失敗")
        
        # 創建批次處理結果
        batch_result = BatchProcessingResult(total_files=2)
        batch_result.add_result(result1)
        batch_result.add_result(result2)
        batch_result.mark_completed()
        
        print("✅ 測試資料創建完成")
        
        # 測試報告生成器
        print("\n2. 測試報告生成器...")
        
        generator = ReportGenerator(output_dir="test_reports")
        
        # 生成報告
        report_paths = generator.generate_processing_report(batch_result, "test_report")
        
        print(f"✅ 報告生成完成:")
        for format_type, path in report_paths.items():
            print(f"   {format_type.upper()}: {path}")
        
        # 驗證檔案存在
        print("\n3. 驗證報告檔案...")
        for format_type, path in report_paths.items():
            if Path(path).exists():
                file_size = Path(path).stat().st_size
                print(f"✅ {format_type.upper()} 報告存在: {path} ({file_size} bytes)")
            else:
                print(f"❌ {format_type.upper()} 報告不存在: {path}")
        
        print("\n=== 測試完成 ===")
        
    except Exception as e:
        print(f"❌ 測試失敗: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)