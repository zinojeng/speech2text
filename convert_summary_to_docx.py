"""
將 ADA2025 摘要轉換為 Word 文件
Convert ADA2025 summary to Word document while preserving Markdown formatting
"""

import os
import sys
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Tuple

# 設定日誌
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
    logger.info("python-docx 模組載入成功")
except ImportError as e:
    DOCX_AVAILABLE = False
    logger.warning(f"python-docx 未安裝，Word 轉換功能將不可用: {e}")

class MarkdownToDocxConverter:
    """
    Markdown 到 Word 文件轉換器
    
    支援的 Markdown 格式：
    - 標題 (# ## ###)
    - 粗體 (**text**)
    - 斜體 (*text*)
    - 項目符號 (- item)
    - 編號列表 (1. item)
    - 程式碼 (`code`)
    - 分隔線 (---)
    - 表格
    """
    
    def __init__(self):
        """初始化轉換器"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝，無法使用 Word 轉換功能")
        
        self.doc = Document()
        self._setup_styles()
        logger.info("Markdown 到 Word 轉換器初始化完成")
    
    def _setup_styles(self):
        """設定 Word 文件樣式"""
        try:
            # 設定預設字體
            style = self.doc.styles['Normal']
            font = style.font
            font.name = '微軟正黑體'
            font.size = Pt(11)
            
            # 創建標題樣式
            for i in range(1, 4):
                heading_style_name = f'Heading {i}'
                if heading_style_name in self.doc.styles:
                    heading_style = self.doc.styles[heading_style_name]
                else:
                    heading_style = self.doc.styles.add_style(heading_style_name, WD_STYLE_TYPE.PARAGRAPH)
                
                heading_font = heading_style.font
                heading_font.name = '微軟正黑體'
                heading_font.bold = True
                heading_font.size = Pt(16 - i * 2)  # H1=16pt, H2=14pt, H3=12pt
                
                if i == 1:
                    heading_font.color.rgb = RGBColor(0, 51, 102)  # 深藍色
                elif i == 2:
                    heading_font.color.rgb = RGBColor(0, 102, 204)  # 藍色
                else:
                    heading_font.color.rgb = RGBColor(51, 51, 51)  # 深灰色
            
            # 創建程式碼樣式
            code_style = self.doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(10)
            code_font.color.rgb = RGBColor(199, 37, 78)  # 紅色
            
            # 創建強調樣式
            emphasis_style = self.doc.styles.add_style('Emphasis', WD_STYLE_TYPE.CHARACTER)
            emphasis_font = emphasis_style.font
            emphasis_font.bold = True
            emphasis_font.color.rgb = RGBColor(0, 102, 51)  # 深綠色
            
            logger.info("Word 文件樣式設定完成")
            
        except Exception as e:
            logger.warning(f"設定樣式時發生錯誤: {e}")
    
    def _parse_markdown_line(self, line: str) -> Tuple[str, str, str]:
        """
        解析 Markdown 行
        
        Args:
            line: Markdown 行內容
            
        Returns:
            (type, content, level) - 類型、內容、層級
        """
        line = line.strip()
        
        if not line:
            return 'empty', '', ''
        
        # 標題
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            content = line.lstrip('#').strip()
            return 'heading', content, str(level)
        
        # 分隔線
        if line.startswith('---') or line.startswith('***'):
            return 'separator', '', ''
        
        # 項目符號
        if line.startswith('- ') or line.startswith('* '):
            content = line[2:].strip()
            return 'bullet', content, '1'
        
        # 縮排項目符號
        if line.startswith('  - ') or line.startswith('  * '):
            content = line[4:].strip()
            return 'bullet', content, '2'
        
        # 編號列表
        if re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            return 'number', content, '1'
        
        # 表格標題
        if '|' in line and line.count('|') >= 2:
            return 'table', line, ''
        
        # 一般段落
        return 'paragraph', line, ''
    
    def _format_text(self, text: str, paragraph):
        """
        格式化文字（處理粗體、斜體、程式碼）
        
        Args:
            text: 原始文字
            paragraph: Word 段落物件
        """
        # 處理粗體 **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗體文字
                bold_text = part[2:-2]
                # 再處理程式碼
                code_parts = re.split(r'(`.*?`)', bold_text)
                for code_part in code_parts:
                    if code_part.startswith('`') and code_part.endswith('`'):
                        run = paragraph.add_run(code_part[1:-1])
                        run.style = 'Code'
                        run.bold = True
                    else:
                        run = paragraph.add_run(code_part)
                        run.bold = True
            else:
                # 處理程式碼 `code`
                code_parts = re.split(r'(`.*?`)', part)
                for code_part in code_parts:
                    if code_part.startswith('`') and code_part.endswith('`'):
                        run = paragraph.add_run(code_part[1:-1])
                        run.style = 'Code'
                    else:
                        # 處理斜體 *text*
                        italic_parts = re.split(r'(\*.*?\*)', code_part)
                        for italic_part in italic_parts:
                            if italic_part.startswith('*') and italic_part.endswith('*') and not italic_part.startswith('**'):
                                run = paragraph.add_run(italic_part[1:-1])
                                run.italic = True
                            else:
                                paragraph.add_run(italic_part)
    
    def _add_table(self, table_lines: List[str]):
        """
        添加表格
        
        Args:
            table_lines: 表格行列表
        """
        if len(table_lines) < 2:
            return
        
        # 解析表格標題
        header_line = table_lines[0]
        headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
        
        if not headers:
            return
        
        # 跳過分隔線，處理資料行
        data_lines = table_lines[2:] if len(table_lines) > 2 else []
        
        # 創建表格
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # 設定標題行
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = header
                # 設定標題樣式
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # 添加資料行
        for data_line in data_lines:
            if '|' in data_line:
                cells_data = [cell.strip() for cell in data_line.split('|') if cell.strip()]
                if cells_data:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(cells_data):
                        if i < len(row_cells):
                            row_cells[i].text = cell_data
        
        logger.info(f"已添加表格: {len(headers)} 列 x {len(data_lines)} 行")
    
    def convert_markdown_to_docx(self, markdown_content: str, title: str = "文件"):
        """
        將 Markdown 內容轉換為 Word 文件
        
        Args:
            markdown_content: Markdown 內容
            title: 文件標題
        """
        lines = markdown_content.split('\n')
        table_lines = []
        in_table = False
        
        for line in lines:
            line_type, content, level = self._parse_markdown_line(line)
            
            # 處理表格
            if line_type == 'table':
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(content)
                continue
            else:
                if in_table:
                    # 結束表格處理
                    self._add_table(table_lines)
                    in_table = False
                    table_lines = []
            
            # 處理其他類型
            if line_type == 'empty':
                continue
            
            elif line_type == 'heading':
                if level in ['1', '2', '3']:
                    heading = self.doc.add_heading(content, level=int(level))
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # 超過 3 級的標題當作粗體段落處理
                    para = self.doc.add_paragraph()
                    run = para.add_run(content)
                    run.bold = True
                    run.font.size = Pt(12)
            
            elif line_type == 'separator':
                # 添加分隔線
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run('─' * 50)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            elif line_type == 'bullet':
                para = self.doc.add_paragraph(style='List Bullet')
                if level == '2':
                    # 縮排項目
                    para.paragraph_format.left_indent = Inches(0.5)
                self._format_text(content, para)
            
            elif line_type == 'number':
                para = self.doc.add_paragraph(style='List Number')
                self._format_text(content, para)
            
            elif line_type == 'paragraph':
                para = self.doc.add_paragraph()
                self._format_text(content, para)
        
        # 處理最後的表格
        if in_table and table_lines:
            self._add_table(table_lines)
        
        logger.info("Markdown 轉換完成")
    
    def save_document(self, output_path: str):
        """
        保存 Word 文件
        
        Args:
            output_path: 輸出檔案路徑
        """
        self.doc.save(output_path)
        logger.info(f"Word 文件已保存: {output_path}")

def convert_ada2025_summary_to_docx():
    """將 ADA2025 改進摘要轉換為 Word 文件"""
    
    if not DOCX_AVAILABLE:
        print("❌ python-docx 未安裝，無法轉換 Word 文件")
        print("請執行: pip install python-docx")
        return False
    
    folder_path = "/Volumes/WD_BLACK/國際年會/ADA2025/Standards of Care in Diabetes 2025 Updates"
    summary_file = "CT-SY33-1.5 - Standards of Care in Diabetes 2025 Updates_改進摘要_20250727_095645.md"
    
    print("=== ADA2025 摘要轉換為 Word 文件 ===")
    print(f"資料夾: {folder_path}")
    print(f"摘要檔案: {summary_file}")
    print()
    
    try:
        # 檢查摘要檔案是否存在
        summary_path = Path(folder_path) / summary_file
        
        if not summary_path.exists():
            print(f"❌ 摘要檔案不存在: {summary_path}")
            return False
        
        print("1. 讀取改進摘要內容...")
        
        with open(summary_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        print(f"   ✅ 摘要內容已讀取 ({len(markdown_content)} 字符)")
        
        # 提取摘要的主要內容（跳過檔案資訊部分）
        content_lines = markdown_content.split('\n')
        main_content_start = False
        main_content_lines = []
        
        for line in content_lines:
            if line.strip() == "=" * 50:
                main_content_start = True
                continue
            elif main_content_start:
                main_content_lines.append(line)
        
        if not main_content_lines:
            # 如果沒有找到分隔線，使用全部內容
            main_content_lines = content_lines
        
        main_content = '\n'.join(main_content_lines)
        
        print("2. 初始化 Word 轉換器...")
        
        converter = MarkdownToDocxConverter()
        
        print("3. 轉換 Markdown 為 Word 格式...")
        
        # 添加文件標題和資訊
        title = "ADA 2025 年糖尿病照護標準更新摘要"
        
        # 添加標題頁
        title_para = converter.doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加文件資訊
        info_para = converter.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        info_run.italic = True
        
        converter.doc.add_paragraph()  # 空行
        
        # 轉換主要內容
        converter.convert_markdown_to_docx(main_content, title)
        
        print("4. 保存 Word 文件...")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = "CT-SY33-1.5 - Standards of Care in Diabetes 2025 Updates"
        
        # 保存到原資料夾
        docx_file = Path(folder_path) / f"{base_name}_改進摘要_{timestamp}.docx"
        converter.save_document(str(docx_file))
        
        print(f"   ✅ Word 文件已保存: {docx_file.name}")
        
        # 也保存到測試目錄
        test_docx_file = Path("temp/test_results") / f"{base_name}_改進摘要_{timestamp}.docx"
        test_docx_file.parent.mkdir(parents=True, exist_ok=True)
        converter.save_document(str(test_docx_file))
        
        print(f"   ✅ 測試備份已保存: {test_docx_file}")
        
        # 檢查檔案大小
        if docx_file.exists():
            file_size = docx_file.stat().st_size
            print(f"   📊 檔案大小: {file_size:,} bytes ({file_size/1024:.1f} KB)")
        
        # 創建轉換報告
        print("5. 創建轉換報告...")
        
        conversion_report = Path(folder_path) / f"{base_name}_轉換報告_{timestamp}.md"
        
        with open(conversion_report, 'w', encoding='utf-8') as f:
            f.write(f"# ADA2025 摘要 Word 轉換報告\n\n")
            f.write(f"**轉換時間**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"**原始檔案**: {summary_file}\n")
            f.write(f"**Word 檔案**: {docx_file.name}\n")
            f.write(f"**檔案大小**: {file_size:,} bytes\n\n")
            
            f.write("## 轉換功能\n\n")
            f.write("### ✅ 支援的 Markdown 格式：\n")
            f.write("- 標題層級 (# ## ###)\n")
            f.write("- 粗體文字 (**text**)\n")
            f.write("- 斜體文字 (*text*)\n")
            f.write("- 程式碼 (`code`)\n")
            f.write("- 項目符號列表 (- item)\n")
            f.write("- 編號列表 (1. item)\n")
            f.write("- 表格格式\n")
            f.write("- 分隔線 (---)\n\n")
            
            f.write("### 🎨 Word 格式特色：\n")
            f.write("- 微軟正黑體中文字型\n")
            f.write("- 階層式標題顏色\n")
            f.write("- 程式碼特殊格式 (Consolas 字型)\n")
            f.write("- 粗體重點標示\n")
            f.write("- 專業表格樣式\n")
            f.write("- 適合列印和分享\n\n")
            
            f.write("## 使用建議\n\n")
            f.write("1. **Word 檔案**適合：\n")
            f.write("   - 正式文件分享\n")
            f.write("   - 列印和存檔\n")
            f.write("   - 進一步編輯和註解\n\n")
            f.write("2. **Markdown 檔案**適合：\n")
            f.write("   - 線上閱讀\n")
            f.write("   - 版本控制\n")
            f.write("   - 快速瀏覽\n\n")
        
        print(f"   ✅ 轉換報告已保存: {conversion_report.name}")
        
        print(f"\n=== 轉換完成 ===")
        print(f"✅ ADA2025 摘要已成功轉換為 Word 文件")
        print(f"📁 Word 檔案: {docx_file}")
        print(f"📄 Markdown 原檔: {summary_path}")
        print(f"📊 轉換報告: {conversion_report}")
        print(f"💾 測試備份: {test_docx_file}")
        print(f"🎯 兩種格式都已保留，可根據需要使用")
        
        return True
        
    except Exception as e:
        print(f"❌ 轉換失敗: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = convert_ada2025_summary_to_docx()
    sys.exit(0 if success else 1)