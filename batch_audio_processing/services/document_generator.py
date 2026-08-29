"""
批次音訊處理系統 - 文件生成器
Document Generator for Batch Audio Processing System

此模組實作文件生成功能，包括：
- MarkdownProcessor: Markdown 格式處理和生成
- DocxConverter: Word 文件轉換器
- DocumentGenerator: 文件生成協調器

Requirements: 5.1, 5.2, 5.3, 5.4, 5.5
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime

# 第三方套件
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx 未安裝，Word 文件生成功能將不可用")

try:
    from markdown import markdown
    from bs4 import BeautifulSoup
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False
    logging.warning("markdown 和 beautifulsoup4 未安裝，Markdown 處理功能將受限")

# 本地模組
from batch_audio_models import FileInfo, TranscriptionResult, SummaryResult

# 設定日誌
logger = logging.getLogger(__name__)


@dataclass
class DocumentResult:
    """
    文件生成結果類別
    
    包含生成的文件路徑和相關資訊
    """
    success: bool
    markdown_path: Optional[str] = None
    docx_path: Optional[str] = None
    error: Optional[str] = None
    processing_time: float = 0.0
    file_size_kb: float = 0.0
    
    def get_display_summary(self) -> str:
        """取得用於顯示的結果摘要"""
        if self.success:
            parts = ["✅ 文件生成成功"]
            if self.markdown_path:
                parts.append(f"Markdown: {Path(self.markdown_path).name}")
            if self.docx_path:
                parts.append(f"Word: {Path(self.docx_path).name}")
            if self.file_size_kb > 0:
                parts.append(f"大小: {self.file_size_kb:.1f}KB")
            return " | ".join(parts)
        else:
            return f"❌ 文件生成失敗: {self.error or '未知錯誤'}"


class MarkdownProcessor:
    """
    Markdown 處理器類別
    
    負責 Markdown 格式的生成、處理和格式化
    Requirements: 5.1, 5.2
    """
    
    def __init__(self):
        """初始化 Markdown 處理器"""
        self.image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
        self.link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        self.bold_pattern = re.compile(r'\*\*([^*]+)\*\*')
        self.italic_pattern = re.compile(r'\*([^*]+)\*')
        self.underline_pattern = re.compile(r'__([^_]+)__')
        
        logger.info("MarkdownProcessor 初始化完成")
    
    def generate_markdown(self, content: str, title: Optional[str] = None, 
                         metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        生成標準化的 Markdown 內容
        
        Args:
            content: 原始內容
            title: 文件標題
            metadata: 元資料資訊
            
        Returns:
            格式化的 Markdown 字串
        """
        try:
            markdown_parts = []
            
            # 添加標題
            if title:
                markdown_parts.append(f"# {title}\n")
            
            # 添加元資料
            if metadata:
                markdown_parts.append(self._generate_metadata_section(metadata))
            
            # 處理內容格式
            formatted_content = self._format_content(content)
            markdown_parts.append(formatted_content)
            
            # 添加生成時間戳記
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            markdown_parts.append(f"\n---\n*生成時間: {timestamp}*")
            
            result = "\n".join(markdown_parts)
            logger.debug(f"Markdown 生成完成，長度: {len(result)} 字符")
            
            return result
            
        except Exception as e:
            logger.error(f"生成 Markdown 時發生錯誤: {e}")
            return content  # 返回原始內容作為備用
    
    def _generate_metadata_section(self, metadata: Dict[str, Any]) -> str:
        """生成元資料區段"""
        lines = ["## 文件資訊\n"]
        
        for key, value in metadata.items():
            if value is not None:
                # 格式化不同類型的值
                if isinstance(value, (int, float)):
                    formatted_value = str(value)
                elif isinstance(value, bool):
                    formatted_value = "是" if value else "否"
                elif isinstance(value, list):
                    formatted_value = ", ".join(str(v) for v in value)
                else:
                    formatted_value = str(value)
                
                lines.append(f"- **{key}**: {formatted_value}")
        
        lines.append("")  # 空行分隔
        return "\n".join(lines)
    
    def _format_content(self, content: str) -> str:
        """
        格式化內容，確保 Markdown 語法正確
        
        Args:
            content: 原始內容
            
        Returns:
            格式化後的內容
        """
        if not content:
            return ""
        
        # 分行處理
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            # 處理標題格式
            formatted_line = self._format_headings(line)
            
            # 處理列表格式
            formatted_line = self._format_lists(formatted_line)
            
            # 處理強調格式
            formatted_line = self._format_emphasis(formatted_line)
            
            formatted_lines.append(formatted_line)
        
        return '\n'.join(formatted_lines)
    
    def _format_headings(self, line: str) -> str:
        """格式化標題"""
        line = line.strip()
        
        # 檢查是否已經是正確的 Markdown 標題格式
        if re.match(r'^#{1,6}\s+', line):
            return line
        
        # 檢查是否是標題樣式的文字
        if line and (line.isupper() or line.endswith('：') or line.endswith(':')):
            # 根據內容長度決定標題級別
            if len(line) < 20:
                return f"## {line}"
            else:
                return f"### {line}"
        
        return line
    
    def _format_lists(self, line: str) -> str:
        """格式化列表"""
        line = line.strip()
        
        # 檢查是否已經是正確的列表格式
        if re.match(r'^[-*+]\s+', line) or re.match(r'^\d+\.\s+', line):
            return line
        
        # 檢查是否是列表項目
        if line.startswith('•') or line.startswith('◦') or line.startswith('▪'):
            return f"- {line[1:].strip()}"
        
        # 檢查數字列表
        if re.match(r'^\d+[.)]\s*', line):
            return re.sub(r'^(\d+)[.)](\s*)', r'\1. ', line)
        
        return line
    
    def _format_emphasis(self, line: str) -> str:
        """格式化強調文字"""
        # 處理粗體（保持現有的 ** 格式）
        if '**' not in line:
            # 尋找可能的粗體文字（全大寫或特殊標記）
            line = re.sub(r'【([^】]+)】', r'**\1**', line)
            line = re.sub(r'《([^》]+)》', r'**\1**', line)
        
        # 處理底線（轉換為 Markdown 格式）
        line = re.sub(r'_([^_]+)_', r'*\1*', line)
        
        return line
    
    def process_image_links(self, content: str, base_path: str = "") -> Tuple[str, int]:
        """
        處理圖片連結，確保路徑正確
        
        Args:
            content: 包含圖片連結的內容
            base_path: 圖片的基礎路徑
            
        Returns:
            處理後的內容和處理的圖片數量
        """
        try:
            processed_count = 0
            
            def replace_image(match):
                nonlocal processed_count
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # 處理相對路徑
                if base_path and not os.path.isabs(image_path):
                    full_path = os.path.join(base_path, image_path)
                    
                    # 檢查檔案是否存在
                    if os.path.exists(full_path):
                        processed_count += 1
                        return f"![{alt_text}]({full_path})"
                    else:
                        logger.warning(f"圖片檔案不存在: {full_path}")
                        return f"![{alt_text}]({image_path}) *(檔案不存在)*"
                
                processed_count += 1
                return match.group(0)  # 保持原樣
            
            processed_content = self.image_pattern.sub(replace_image, content)
            
            logger.debug(f"處理了 {processed_count} 個圖片連結")
            return processed_content, processed_count
            
        except Exception as e:
            logger.error(f"處理圖片連結時發生錯誤: {e}")
            return content, 0
    
    def insert_image_at_position(self, content: str, image_path: str, 
                               position: int, alt_text: str = "") -> str:
        """
        在指定位置插入圖片
        
        Args:
            content: 原始內容
            image_path: 圖片路徑
            position: 插入位置（行號）
            alt_text: 圖片替代文字
            
        Returns:
            插入圖片後的內容
        """
        try:
            lines = content.split('\n')
            
            if 0 <= position <= len(lines):
                image_markdown = f"![{alt_text}]({image_path})"
                lines.insert(position, image_markdown)
                lines.insert(position + 1, "")  # 添加空行
                
                logger.debug(f"在第 {position} 行插入圖片: {image_path}")
                return '\n'.join(lines)
            else:
                logger.warning(f"插入位置 {position} 超出範圍")
                return content
                
        except Exception as e:
            logger.error(f"插入圖片時發生錯誤: {e}")
            return content
    
    def validate_markdown_syntax(self, content: str) -> Dict[str, Any]:
        """
        驗證 Markdown 語法
        
        Args:
            content: Markdown 內容
            
        Returns:
            驗證結果字典
        """
        validation_result = {
            'valid': True,
            'warnings': [],
            'errors': [],
            'statistics': {}
        }
        
        try:
            lines = content.split('\n')
            
            # 統計資訊
            validation_result['statistics'] = {
                'total_lines': len(lines),
                'headings': len(re.findall(r'^#{1,6}\s+', content, re.MULTILINE)),
                'images': len(self.image_pattern.findall(content)),
                'links': len(self.link_pattern.findall(content)),
                'bold_text': len(self.bold_pattern.findall(content)),
                'italic_text': len(self.italic_pattern.findall(content))
            }
            
            # 檢查常見問題
            for i, line in enumerate(lines, 1):
                # 檢查標題格式
                if line.strip().startswith('#'):
                    if not re.match(r'^#{1,6}\s+', line.strip()):
                        validation_result['warnings'].append(
                            f"第 {i} 行: 標題格式可能不正確"
                        )
                
                # 檢查圖片連結
                for match in self.image_pattern.finditer(line):
                    image_path = match.group(2)
                    if not os.path.exists(image_path) and not image_path.startswith('http'):
                        validation_result['warnings'].append(
                            f"第 {i} 行: 圖片檔案可能不存在: {image_path}"
                        )
            
            # 檢查是否有未配對的格式標記
            if content.count('**') % 2 != 0:
                validation_result['warnings'].append("粗體標記 (**) 可能未正確配對")
            
            if content.count('*') % 2 != 0:
                validation_result['warnings'].append("斜體標記 (*) 可能未正確配對")
            
            # 如果有警告或錯誤，標記為無效
            if validation_result['warnings'] or validation_result['errors']:
                validation_result['valid'] = len(validation_result['errors']) == 0
            
            logger.debug(f"Markdown 語法驗證完成: {validation_result['statistics']}")
            
        except Exception as e:
            validation_result['valid'] = False
            validation_result['errors'].append(f"驗證過程發生錯誤: {e}")
            logger.error(f"Markdown 語法驗證失敗: {e}")
        
        return validation_result
    
    def save_markdown_file(self, content: str, output_path: str) -> bool:
        """
        儲存 Markdown 檔案
        
        Args:
            content: Markdown 內容
            output_path: 輸出檔案路徑
            
        Returns:
            是否成功儲存
        """
        try:
            # 確保輸出目錄存在
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
            
            # 寫入檔案
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # 驗證檔案是否成功寫入
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Markdown 檔案已儲存: {output_path} ({file_size} bytes)")
                return True
            else:
                logger.error(f"Markdown 檔案儲存失敗: {output_path}")
                return False
                
        except Exception as e:
            logger.error(f"儲存 Markdown 檔案時發生錯誤: {e}")
            return False


if __name__ == "__main__":
    """測試 MarkdownProcessor 功能"""
    # 設定日誌
    logging.basicConfig(level=logging.INFO)
    
    print("=== MarkdownProcessor 測試 ===\n")
    
    # 創建處理器實例
    processor = MarkdownProcessor()
    
    # 測試內容
    test_content = """
    這是一個測試標題
    
    這是一段普通文字，包含一些**粗體文字**和*斜體文字*。
    
    • 這是一個列表項目
    • 這是另一個列表項目
    
    1) 這是編號列表
    2) 第二個項目
    
    【重要提醒】這應該變成粗體
    
    這裡有一張圖片: ![測試圖片](test.jpg)
    """
    
    # 測試 Markdown 生成
    print("1. 測試 Markdown 生成...")
    metadata = {
        "檔案名稱": "test.mp3",
        "處理時間": "2024-01-01 12:00:00",
        "模型": "gpt-transcribe",
        "成功": True
    }
    
    result = processor.generate_markdown(test_content, "測試文件", metadata)
    print("✅ Markdown 生成完成")
    print(f"內容長度: {len(result)} 字符")
    
    # 測試語法驗證
    print("\n2. 測試語法驗證...")
    validation = processor.validate_markdown_syntax(result)
    print(f"✅ 語法驗證完成")
    print(f"有效: {validation['valid']}")
    print(f"統計: {validation['statistics']}")
    if validation['warnings']:
        print(f"警告: {len(validation['warnings'])} 個")
    
    # 測試圖片處理
    print("\n3. 測試圖片處理...")
    processed_content, image_count = processor.process_image_links(result, "./images")
    print(f"✅ 圖片處理完成，處理了 {image_count} 個圖片")
    
    # 測試檔案儲存
    print("\n4. 測試檔案儲存...")
    test_output = "test_output.md"
    success = processor.save_markdown_file(result, test_output)
    if success:
        print(f"✅ 檔案儲存成功: {test_output}")
        # 清理測試檔案
        if os.path.exists(test_output):
            os.remove(test_output)
    else:
        print("❌ 檔案儲存失敗")
    
    print("\n=== 測試完成 ===")

class DocxConverter:
    """
    Word 文件轉換器類別
    
    負責將 Markdown 內容轉換為 Word 文件格式
    Requirements: 5.3, 5.4
    """
    
    def __init__(self):
        """初始化 DocxConverter"""
        if not HAS_DOCX:
            raise ImportError("python-docx 套件未安裝，無法使用 Word 文件轉換功能")
        
        if not HAS_MARKDOWN:
            logger.warning("markdown 套件未安裝，將使用簡化的 Markdown 解析")
        
        # 中文字體設定
        self.chinese_font = "Microsoft JhengHei"  # 微軟正黑體
        self.fallback_fonts = ["SimSun", "Arial Unicode MS", "Arial"]
        
        logger.info("DocxConverter 初始化完成")
    
    def convert_markdown_to_docx(self, markdown_content: str, output_path: str, 
                                title: Optional[str] = None, 
                                metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        將 Markdown 內容轉換為 Word 文件
        
        Args:
            markdown_content: Markdown 格式內容
            output_path: 輸出檔案路徑
            title: 文件標題
            metadata: 元資料資訊
            
        Returns:
            轉換是否成功
        """
        try:
            logger.info(f"開始轉換 Markdown 為 Word 文件: {output_path}")
            
            # 創建新文件
            doc = Document()
            
            # 設定文件屬性
            self._set_document_properties(doc, title, metadata)
            
            # 添加標題
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_chinese_font(title_para)
            
            # 添加元資料區段
            if metadata:
                self._add_metadata_section(doc, metadata)
            
            # 處理 Markdown 內容
            if HAS_MARKDOWN:
                self._process_markdown_with_parser(doc, markdown_content)
            else:
                self._process_markdown_simple(doc, markdown_content)
            
            # 確保輸出目錄存在
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
            
            # 儲存文件
            doc.save(output_path)
            
            # 驗證檔案是否成功儲存
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Word 文件已儲存: {output_path} ({file_size} bytes)")
                return True
            else:
                logger.error(f"Word 文件儲存失敗: {output_path}")
                return False
                
        except Exception as e:
            logger.error(f"轉換 Markdown 為 Word 文件時發生錯誤: {e}")
            return False
    
    def _set_document_properties(self, doc: Document, title: Optional[str], 
                               metadata: Optional[Dict[str, Any]]) -> None:
        """設定文件屬性"""
        try:
            core_props = doc.core_properties
            
            if title:
                core_props.title = title
            
            core_props.author = "批次音訊處理系統"
            core_props.created = datetime.now()
            core_props.modified = datetime.now()
            
            if metadata:
                # 設定主題和關鍵字
                if 'model' in metadata:
                    core_props.subject = f"使用 {metadata['model']} 處理"
                
                if 'keywords' in metadata:
                    core_props.keywords = metadata['keywords']
            
        except Exception as e:
            logger.warning(f"設定文件屬性時發生錯誤: {e}")
    
    def _add_metadata_section(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """添加元資料區段"""
        try:
            # 添加分隔線
            doc.add_paragraph("─" * 50)
            
            # 添加元資料標題
            meta_heading = doc.add_heading("文件資訊", level=2)
            self._set_chinese_font(meta_heading)
            
            # 添加元資料表格
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # 設定表格標題
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '項目'
            hdr_cells[1].text = '內容'
            
            # 設定標題格式
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        self._set_font_properties(run)
            
            # 添加元資料行
            for key, value in metadata.items():
                if value is not None:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = self._format_metadata_value(value)
                    
                    # 設定字體
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self._set_font_properties(run)
            
            # 添加空行
            doc.add_paragraph()
            
        except Exception as e:
            logger.warning(f"添加元資料區段時發生錯誤: {e}")
    
    def _format_metadata_value(self, value: Any) -> str:
        """格式化元資料值"""
        if isinstance(value, bool):
            return "是" if value else "否"
        elif isinstance(value, (int, float)):
            return str(value)
        elif isinstance(value, list):
            return ", ".join(str(v) for v in value)
        elif isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        else:
            return str(value)
    
    def _process_markdown_with_parser(self, doc: Document, markdown_content: str) -> None:
        """使用 markdown 套件處理內容"""
        try:
            # 轉換為 HTML
            html = markdown(markdown_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            # 處理 HTML 元素
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'pre', 'blockquote']):
                self._process_html_element(doc, element)
                
        except Exception as e:
            logger.warning(f"使用 markdown 套件處理時發生錯誤，改用簡化處理: {e}")
            self._process_markdown_simple(doc, markdown_content)
    
    def _process_markdown_simple(self, doc: Document, markdown_content: str) -> None:
        """簡化的 Markdown 處理"""
        try:
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    doc.add_paragraph()  # 空行
                    continue
                
                # 處理標題
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    
                    heading = doc.add_heading(title_text, level=min(level, 6))
                    self._set_chinese_font(heading)
                    continue
                
                # 處理列表項目
                if line.startswith(('- ', '* ', '+ ')):
                    list_text = line[2:].strip()
                    paragraph = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(paragraph, list_text)
                elif re.match(r'^\d+\.\s', line):
                    list_text = re.sub(r'^\d+\.\s', '', line).strip()
                    paragraph = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(paragraph, list_text)
                elif line.startswith('>'):
                    # 引用文字
                    quote_text = line[1:].strip()
                    paragraph = doc.add_paragraph(style='Quote')
                    self._add_formatted_text(paragraph, quote_text)
                else:
                    # 處理普通段落
                    paragraph = doc.add_paragraph()
                    self._add_formatted_text(paragraph, line)
                    
        except Exception as e:
            logger.error(f"簡化 Markdown 處理時發生錯誤: {e}")
    
    def _process_html_element(self, doc: Document, element) -> None:
        """處理 HTML 元素"""
        try:
            if element.name == 'p':
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, element.get_text())
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
                self._set_chinese_font(heading)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    paragraph = doc.add_paragraph(style=style)
                    self._add_formatted_text(paragraph, li.get_text())
            elif element.name == 'pre':
                paragraph = doc.add_paragraph(style='Normal')
                run = paragraph.add_run(element.get_text())
                run.font.name = 'Courier New'  # 等寬字體
                run.font.size = Pt(10)
            elif element.name == 'blockquote':
                paragraph = doc.add_paragraph(style='Quote')
                self._add_formatted_text(paragraph, element.get_text())
                
        except Exception as e:
            logger.warning(f"處理 HTML 元素時發生錯誤: {e}")
    
    def _add_formatted_text(self, paragraph, text: str) -> None:
        """
        處理文字格式並添加到段落
        
        Args:
            paragraph: Word 段落對象
            text: 要處理的文字
        """
        try:
            # 處理圖片連結
            text = self._process_image_links(paragraph, text)
            
            # 使用正則表達式分割格式化文字
            # 處理順序：雙底線 -> 雙星號 -> 單底線 -> 單星號
            parts = re.split(r'(__[^_]+__|_[^_]+_|\*\*[^*]+\*\*|\*[^*]+\*)', text)
            
            for part in parts:
                if not part:
                    continue
                
                run = None
                
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    # 粗體
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('__') and part.endswith('__') and len(part) > 4:
                    # 底線（雙底線）
                    run = paragraph.add_run(part[2:-2])
                    run.underline = True
                elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                    # 底線（單底線）
                    run = paragraph.add_run(part[1:-1])
                    run.underline = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    # 斜體
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    # 普通文字
                    run = paragraph.add_run(part)
                
                # 設定字體
                if run:
                    self._set_font_properties(run)
                    
        except Exception as e:
            logger.warning(f"處理格式化文字時發生錯誤: {e}")
            # 備用方案：直接添加純文字
            run = paragraph.add_run(text)
            self._set_font_properties(run)
    
    def _process_image_links(self, paragraph, text: str) -> str:
        """
        處理圖片連結，嘗試插入圖片到文件中
        
        Args:
            paragraph: Word 段落對象
            text: 包含圖片連結的文字
            
        Returns:
            移除圖片連結後的文字
        """
        try:
            # 尋找 Markdown 圖片語法
            image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
            
            def replace_image(match):
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # 嘗試插入圖片
                if os.path.exists(image_path):
                    try:
                        # 在當前段落前插入圖片
                        img_paragraph = paragraph._element.getparent().insert_paragraph_before(paragraph._element)
                        img_paragraph = paragraph._element.getparent().paragraphs[-1]
                        
                        # 添加圖片
                        run = img_paragraph.add_run()
                        run.add_picture(image_path, width=Inches(5.5))
                        
                        # 添加圖片說明
                        if alt_text:
                            caption_para = paragraph._element.getparent().add_paragraph()
                            caption_run = caption_para.add_run(f"圖片: {alt_text}")
                            caption_run.italic = True
                            self._set_font_properties(caption_run)
                        
                        logger.debug(f"成功插入圖片: {image_path}")
                        return ""  # 移除原始的圖片連結文字
                        
                    except Exception as e:
                        logger.warning(f"插入圖片失敗 {image_path}: {e}")
                        return f"[圖片: {alt_text}] (檔案不存在或無法載入)"
                else:
                    logger.warning(f"圖片檔案不存在: {image_path}")
                    return f"[圖片: {alt_text}] (檔案不存在)"
            
            # 替換所有圖片連結
            processed_text = image_pattern.sub(replace_image, text)
            return processed_text
            
        except Exception as e:
            logger.warning(f"處理圖片連結時發生錯誤: {e}")
            return text
    
    def _set_chinese_font(self, paragraph) -> None:
        """設定段落的中文字體"""
        try:
            for run in paragraph.runs:
                self._set_font_properties(run)
        except Exception as e:
            logger.warning(f"設定中文字體時發生錯誤: {e}")
    
    def _set_font_properties(self, run) -> None:
        """設定文字執行的字體屬性"""
        try:
            # 設定中文字體
            run.font.name = self.chinese_font
            run.font.size = Pt(12)
            
            # 設定東亞字體（確保中文顯示正確）
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.chinese_font)
            
        except Exception as e:
            logger.warning(f"設定字體屬性時發生錯誤: {e}")
    
    def add_image_to_document(self, doc: Document, image_path: str, 
                            width: float = 5.5, caption: str = "") -> bool:
        """
        添加圖片到文件
        
        Args:
            doc: Word 文件對象
            image_path: 圖片路徑
            width: 圖片寬度（英寸）
            caption: 圖片說明
            
        Returns:
            是否成功添加
        """
        try:
            if not os.path.exists(image_path):
                logger.warning(f"圖片檔案不存在: {image_path}")
                return False
            
            # 添加空行
            doc.add_paragraph()
            
            # 添加圖片
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            # 置中對齊
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加圖片說明
            if caption:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"圖片: {caption}")
                caption_run.italic = True
                self._set_font_properties(caption_run)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加空行
            doc.add_paragraph()
            
            logger.debug(f"成功添加圖片: {image_path}")
            return True
            
        except Exception as e:
            logger.error(f"添加圖片時發生錯誤: {e}")
            return False
    
    def create_document_from_results(self, file_info: FileInfo, 
                                   transcription_result: TranscriptionResult,
                                   summary_result: SummaryResult,
                                   output_path: str) -> bool:
        """
        從處理結果創建完整的 Word 文件
        
        Args:
            file_info: 檔案資訊
            transcription_result: 轉錄結果
            summary_result: 摘要結果
            output_path: 輸出路徑
            
        Returns:
            是否成功創建
        """
        try:
            # 準備元資料
            metadata = {
                "檔案名稱": file_info.audio_name,
                "檔案大小": f"{file_info.file_size_mb:.2f} MB",
                "檔案格式": file_info.file_format,
                "處理時間": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "轉錄模型": transcription_result.model_used,
                "摘要模型": summary_result.model_used,
                "轉錄成功": transcription_result.success,
                "摘要成功": summary_result.success,
                "使用議程": summary_result.agenda_used,
                "插入圖片": summary_result.images_inserted
            }
            
            # 準備內容
            content_parts = []
            
            # 添加摘要內容
            if summary_result.success and summary_result.content:
                content_parts.append("## 智能摘要\n")
                content_parts.append(summary_result.content)
                content_parts.append("\n---\n")
            
            # 添加轉錄內容
            if transcription_result.success and transcription_result.content:
                content_parts.append("## 完整轉錄\n")
                content_parts.append(transcription_result.content)
            
            markdown_content = "\n".join(content_parts)
            
            # 轉換為 Word 文件
            title = f"ADA 2025 會議筆記 - {file_info.audio_name}"
            return self.convert_markdown_to_docx(markdown_content, output_path, title, metadata)
            
        except Exception as e:
            logger.error(f"從處理結果創建文件時發生錯誤: {e}")
            return False


if __name__ == "__main__":
    """測試 DocxConverter 功能"""
    # 設定日誌
    logging.basicConfig(level=logging.INFO)
    
    print("=== DocxConverter 測試 ===\n")
    
    if not HAS_DOCX:
        print("❌ python-docx 未安裝，無法進行測試")
        exit(1)
    
    # 創建轉換器實例
    converter = DocxConverter()
    
    # 測試內容
    test_markdown = """
# 測試文件標題

## 第一章節

這是一段包含**粗體文字**和*斜體文字*的普通段落。

### 子標題

這裡有一個列表：

- 第一個項目
- 第二個項目，包含**重要內容**
- 第三個項目

編號列表：

1. 首先做這個
2. 然後做那個
3. 最後完成

> 這是一段引用文字，用來測試引用格式。

## 第二章節

這是另一個段落，包含一些__底線文字__。

![測試圖片](test.jpg)

**結論**：這是測試文件的結論部分。
"""
    
    # 測試元資料
    test_metadata = {
        "檔案名稱": "test_audio.mp3",
        "處理時間": datetime.now(),
        "轉錄模型": "gpt-transcribe",
        "成功": True,
        "檔案大小": "15.5 MB"
    }
    
    # 測試轉換
    print("1. 測試 Markdown 到 Word 轉換...")
    output_path = "test_output.docx"
    success = converter.convert_markdown_to_docx(
        test_markdown, 
        output_path, 
        "測試文件標題", 
        test_metadata
    )
    
    if success:
        print(f"✅ 轉換成功: {output_path}")
        print(f"檔案大小: {os.path.getsize(output_path)} bytes")
        
        # 清理測試檔案
        if os.path.exists(output_path):
            os.remove(output_path)
            print("🧹 測試檔案已清理")
    else:
        print("❌ 轉換失敗")
    
    print("\n=== 測試完成 ===")

class DocumentGenerator:
    """
    文件生成協調器類別
    
    整合 MarkdownProcessor 和 DocxConverter，提供統一的文件生成介面
    Requirements: 5.5
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化文件生成器
        
        Args:
            config: 配置選項
        """
        self.config = config or {}
        
        # 初始化處理器
        self.markdown_processor = MarkdownProcessor()
        
        # 只有在需要時才初始化 DocxConverter
        self.docx_converter = None
        if HAS_DOCX:
            try:
                self.docx_converter = DocxConverter()
            except ImportError as e:
                logger.warning(f"DocxConverter 初始化失敗: {e}")
        
        # 預設設定
        self.default_output_formats = ["markdown", "docx"]
        self.default_image_width = 5.5  # 英寸
        
        logger.info("DocumentGenerator 初始化完成")
    
    def generate_documents(self, file_info: FileInfo, 
                         transcription_result: TranscriptionResult,
                         summary_result: SummaryResult,
                         output_dir: str,
                         formats: Optional[List[str]] = None) -> DocumentResult:
        """
        生成多種格式的文件
        
        Args:
            file_info: 檔案資訊
            transcription_result: 轉錄結果
            summary_result: 摘要結果
            output_dir: 輸出目錄
            formats: 要生成的格式列表 (markdown, docx)
            
        Returns:
            文件生成結果
        """
        start_time = datetime.now()
        result = DocumentResult(success=False)
        
        try:
            logger.info(f"開始生成文件: {file_info.audio_name}")
            
            # 使用指定格式或預設格式
            target_formats = formats or self.default_output_formats
            
            # 確保輸出目錄存在
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成檔案名稱
            base_filename = self._generate_filename(file_info.audio_name)
            
            # 準備內容和元資料
            markdown_content, metadata = self._prepare_content_and_metadata(
                file_info, transcription_result, summary_result
            )
            
            # 生成 Markdown 文件
            if "markdown" in target_formats:
                markdown_path = os.path.join(output_dir, f"{base_filename}.md")
                if self._generate_markdown_document(markdown_content, markdown_path, metadata):
                    result.markdown_path = markdown_path
                    logger.info(f"Markdown 文件已生成: {markdown_path}")
            
            # 生成 Word 文件
            if "docx" in target_formats and self.docx_converter:
                docx_path = os.path.join(output_dir, f"{base_filename}.docx")
                title = f"ADA 2025 會議筆記 - {file_info.audio_name}"
                if self.docx_converter.convert_markdown_to_docx(markdown_content, docx_path, title, metadata):
                    result.docx_path = docx_path
                    logger.info(f"Word 文件已生成: {docx_path}")
            elif "docx" in target_formats and not self.docx_converter:
                logger.warning("Word 文件生成被跳過：DocxConverter 不可用")
            
            # 計算處理時間和檔案大小
            result.processing_time = (datetime.now() - start_time).total_seconds()
            result.file_size_kb = self._calculate_total_file_size(result)
            
            # 檢查是否至少生成了一個文件
            result.success = bool(result.markdown_path or result.docx_path)
            
            if result.success:
                logger.info(f"文件生成完成: {file_info.audio_name} ({result.processing_time:.1f}秒)")
            else:
                result.error = "未能生成任何文件"
                logger.error(f"文件生成失敗: {file_info.audio_name}")
            
        except Exception as e:
            result.success = False
            result.error = str(e)
            result.processing_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"文件生成過程發生錯誤: {e}")
        
        return result
    
    def _generate_filename(self, audio_name: str) -> str:
        """
        生成安全的檔案名稱
        
        Args:
            audio_name: 原始音訊檔案名稱
            
        Returns:
            安全的檔案名稱（不含副檔名）
        """
        try:
            # 移除副檔名
            base_name = os.path.splitext(audio_name)[0]
            
            # 替換不安全的字符
            safe_chars = re.sub(r'[<>:"/\\|?*]', '_', base_name)
            
            # 移除多餘的空格和底線
            safe_chars = re.sub(r'[_\s]+', '_', safe_chars).strip('_')
            
            # 限制長度
            if len(safe_chars) > 100:
                safe_chars = safe_chars[:100]
            
            # 確保不為空
            if not safe_chars:
                safe_chars = "document"
            
            # 添加時間戳記以避免衝突
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            final_name = f"{safe_chars}_{timestamp}"
            
            logger.debug(f"生成檔案名稱: {audio_name} -> {final_name}")
            return final_name
            
        except Exception as e:
            logger.warning(f"生成檔案名稱時發生錯誤: {e}")
            # 備用方案
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"document_{timestamp}"
    
    def _prepare_content_and_metadata(self, file_info: FileInfo,
                                    transcription_result: TranscriptionResult,
                                    summary_result: SummaryResult) -> Tuple[str, Dict[str, Any]]:
        """
        準備文件內容和元資料
        
        Args:
            file_info: 檔案資訊
            transcription_result: 轉錄結果
            summary_result: 摘要結果
            
        Returns:
            (markdown_content, metadata) 元組
        """
        try:
            # 準備元資料
            metadata = {
                "檔案名稱": file_info.audio_name,
                "檔案路徑": file_info.audio_path,
                "檔案大小": f"{file_info.file_size_mb:.2f} MB",
                "檔案格式": file_info.file_format or "未知",
                "處理時間": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "轉錄模型": transcription_result.model_used or "未知",
                "轉錄成功": transcription_result.success,
                "轉錄處理時間": f"{transcription_result.processing_time:.1f}秒",
                "摘要模型": summary_result.model_used or "未知",
                "摘要成功": summary_result.success,
                "摘要處理時間": f"{summary_result.processing_time:.1f}秒",
                "使用議程": summary_result.agenda_used,
                "插入圖片數量": summary_result.images_inserted
            }
            
            # 添加議程資訊
            if file_info.agenda_path:
                metadata["議程檔案"] = os.path.basename(file_info.agenda_path)
            
            # 添加估計時長
            if file_info.estimated_duration:
                metadata["估計時長"] = f"{file_info.estimated_duration:.1f}秒"
            
            # 準備內容
            content_parts = []
            
            # 添加摘要內容（如果存在且成功）
            if summary_result.success and summary_result.content:
                content_parts.append("## 智能摘要")
                content_parts.append("")
                content_parts.append(summary_result.content)
                content_parts.append("")
                content_parts.append("---")
                content_parts.append("")
            
            # 添加轉錄內容（如果存在且成功）
            if transcription_result.success and transcription_result.content:
                content_parts.append("## 完整轉錄")
                content_parts.append("")
                content_parts.append(transcription_result.content)
            
            # 如果都沒有成功，添加錯誤資訊
            if not transcription_result.success and not summary_result.success:
                content_parts.append("## 處理結果")
                content_parts.append("")
                content_parts.append("**轉錄狀態**: 失敗")
                if transcription_result.error:
                    content_parts.append(f"**轉錄錯誤**: {transcription_result.error}")
                content_parts.append("")
                content_parts.append("**摘要狀態**: 失敗")
                if summary_result.error:
                    content_parts.append(f"**摘要錯誤**: {summary_result.error}")
            
            markdown_content = "\n".join(content_parts)
            
            logger.debug(f"內容和元資料準備完成，內容長度: {len(markdown_content)} 字符")
            return markdown_content, metadata
            
        except Exception as e:
            logger.error(f"準備內容和元資料時發生錯誤: {e}")
            # 返回基本內容
            basic_metadata = {
                "檔案名稱": file_info.audio_name,
                "處理時間": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "錯誤": str(e)
            }
            basic_content = f"# 處理結果\n\n處理過程中發生錯誤: {e}"
            return basic_content, basic_metadata
    
    def _generate_markdown_document(self, content: str, output_path: str, 
                                  metadata: Dict[str, Any]) -> bool:
        """
        生成 Markdown 文件
        
        Args:
            content: 文件內容
            output_path: 輸出路徑
            metadata: 元資料
            
        Returns:
            是否成功生成
        """
        try:
            # 使用 MarkdownProcessor 生成格式化內容
            title = metadata.get("檔案名稱", "文件")
            formatted_content = self.markdown_processor.generate_markdown(
                content, title, metadata
            )
            
            # 處理圖片連結
            if os.path.dirname(output_path):
                base_path = os.path.dirname(output_path)
                formatted_content, _ = self.markdown_processor.process_image_links(
                    formatted_content, base_path
                )
            
            # 儲存檔案
            return self.markdown_processor.save_markdown_file(formatted_content, output_path)
            
        except Exception as e:
            logger.error(f"生成 Markdown 文件時發生錯誤: {e}")
            return False
    
    def _calculate_total_file_size(self, result: DocumentResult) -> float:
        """
        計算生成文件的總大小
        
        Args:
            result: 文件生成結果
            
        Returns:
            總檔案大小（KB）
        """
        try:
            total_size = 0.0
            
            if result.markdown_path and os.path.exists(result.markdown_path):
                total_size += os.path.getsize(result.markdown_path)
            
            if result.docx_path and os.path.exists(result.docx_path):
                total_size += os.path.getsize(result.docx_path)
            
            return total_size / 1024  # 轉換為 KB
            
        except Exception as e:
            logger.warning(f"計算檔案大小時發生錯誤: {e}")
            return 0.0
    
    def generate_batch_report(self, results: List[DocumentResult], 
                            output_dir: str) -> Optional[str]:
        """
        生成批次處理報告
        
        Args:
            results: 文件生成結果列表
            output_dir: 輸出目錄
            
        Returns:
            報告檔案路徑，如果失敗則返回 None
        """
        try:
            logger.info("開始生成批次處理報告")
            
            # 統計資訊
            total_files = len(results)
            successful_files = sum(1 for r in results if r.success)
            failed_files = total_files - successful_files
            total_processing_time = sum(r.processing_time for r in results)
            total_file_size = sum(r.file_size_kb for r in results)
            
            # 準備報告內容
            report_content = f"""# 批次文件生成報告

## 處理摘要

- **總檔案數**: {total_files}
- **成功處理**: {successful_files}
- **處理失敗**: {failed_files}
- **成功率**: {(successful_files/total_files*100):.1f}%
- **總處理時間**: {total_processing_time:.1f}秒
- **總檔案大小**: {total_file_size:.1f}KB
- **報告生成時間**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 詳細結果

"""
            
            # 添加每個檔案的詳細資訊
            for i, result in enumerate(results, 1):
                status = "✅ 成功" if result.success else "❌ 失敗"
                report_content += f"### {i}. 檔案處理結果\n\n"
                report_content += f"- **狀態**: {status}\n"
                report_content += f"- **處理時間**: {result.processing_time:.1f}秒\n"
                report_content += f"- **檔案大小**: {result.file_size_kb:.1f}KB\n"
                
                if result.markdown_path:
                    report_content += f"- **Markdown**: {os.path.basename(result.markdown_path)}\n"
                
                if result.docx_path:
                    report_content += f"- **Word**: {os.path.basename(result.docx_path)}\n"
                
                if result.error:
                    report_content += f"- **錯誤**: {result.error}\n"
                
                report_content += "\n"
            
            # 儲存報告
            report_path = os.path.join(output_dir, f"batch_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md")
            
            if self.markdown_processor.save_markdown_file(report_content, report_path):
                logger.info(f"批次處理報告已生成: {report_path}")
                return report_path
            else:
                logger.error("批次處理報告生成失敗")
                return None
                
        except Exception as e:
            logger.error(f"生成批次處理報告時發生錯誤: {e}")
            return None
    
    def cleanup_temp_files(self, temp_dir: str) -> None:
        """
        清理臨時檔案
        
        Args:
            temp_dir: 臨時檔案目錄
        """
        try:
            if os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
                logger.info(f"臨時檔案已清理: {temp_dir}")
        except Exception as e:
            logger.warning(f"清理臨時檔案時發生錯誤: {e}")
    
    def validate_output_formats(self, formats: List[str]) -> List[str]:
        """
        驗證輸出格式
        
        Args:
            formats: 要驗證的格式列表
            
        Returns:
            有效的格式列表
        """
        valid_formats = []
        
        for fmt in formats:
            if fmt == "markdown":
                valid_formats.append(fmt)
            elif fmt == "docx":
                if self.docx_converter:
                    valid_formats.append(fmt)
                else:
                    logger.warning("Word 格式不可用：DocxConverter 未初始化")
            else:
                logger.warning(f"不支援的輸出格式: {fmt}")
        
        return valid_formats


if __name__ == "__main__":
    """測試 DocumentGenerator 功能"""
    # 設定日誌
    logging.basicConfig(level=logging.INFO)
    
    print("=== DocumentGenerator 測試 ===\n")
    
    # 創建生成器實例
    generator = DocumentGenerator()
    
    # 創建測試資料
    from batch_audio_models import FileInfo, TranscriptionResult, SummaryResult
    
    # 測試檔案資訊
    test_file_info = FileInfo(
        audio_path="test_audio.mp3",
        audio_name="test_audio.mp3",
        file_size_mb=15.5,
        file_format="mp3"
    )
    
    # 測試轉錄結果
    test_transcription = TranscriptionResult(
        success=True,
        content="這是測試的轉錄內容，包含一些**重要資訊**和詳細說明。",
        processing_time=45.2,
        model_used="gpt-transcribe",
        token_count=150
    )
    
    # 測試摘要結果
    test_summary = SummaryResult(
        success=True,
        content="## 會議摘要\n\n這是測試的摘要內容，包含主要重點和結論。\n\n### 重要決議\n\n- 決議一\n- 決議二",
        processing_time=12.8,
        model_used="gemini-3.7-flash",
        agenda_used=True,
        images_inserted=2
    )
    
    # 測試文件生成
    print("1. 測試文件生成...")
    output_dir = "test_output"
    result = generator.generate_documents(
        test_file_info,
        test_transcription,
        test_summary,
        output_dir,
        formats=["markdown", "docx"]
    )
    
    print(f"生成結果: {result.get_display_summary()}")
    
    # 測試批次報告
    print("\n2. 測試批次報告生成...")
    batch_results = [result]  # 模擬批次結果
    report_path = generator.generate_batch_report(batch_results, output_dir)
    
    if report_path:
        print(f"✅ 批次報告已生成: {report_path}")
    else:
        print("❌ 批次報告生成失敗")
    
    # 清理測試檔案
    print("\n3. 清理測試檔案...")
    generator.cleanup_temp_files(output_dir)
    
    print("\n=== 測試完成 ===")