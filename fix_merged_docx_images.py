#!/usr/bin/env python3
"""
修復合併後的 DOCX 文件中的圖片插入問題
處理兩種格式：
1. [IMAGE: filename.jpg] - 只有檔名
2. > 🖼️ 投影片圖表說明（HH:MM:SS）- 只有時間戳
"""

import os
import sys
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
import glob

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


class MergedDocxImageFixer:
    def __init__(self):
        self.slide_folders = []
        self.filename_to_path = {}  # 檔名 -> 完整路徑
        self.time_to_path = {}      # 時間(秒) -> 完整路徑
        
    def add_slide_folders(self, folders):
        """添加多個投影片資料夾"""
        for folder in folders:
            if os.path.exists(folder):
                self.slide_folders.append(folder)
                logger.info(f"✓ 找到資料夾: {os.path.basename(folder)}")
            else:
                logger.warning(f"✗ 資料夾不存在: {folder}")
    
    def load_all_images(self):
        """載入所有圖片並建立映射"""
        total_images = 0
        
        for folder in self.slide_folders:
            folder_name = os.path.basename(folder)
            image_count = 0
            
            # 掃描所有圖片（包括子目錄）
            patterns = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']
            for pattern in patterns:
                # 使用 ** 來遞歸搜索子目錄
                for img_path in glob.glob(os.path.join(folder, '**', pattern), recursive=True):
                    filename = os.path.basename(img_path)
                    
                    # 建立檔名映射
                    self.filename_to_path[filename] = img_path
                    
                    # 嘗試解析時間
                    time_match = re.search(r't(\d+)m([\d.]+)s', filename)
                    if time_match:
                        minutes = int(time_match.group(1))
                        seconds = float(time_match.group(2))
                        time_sec = minutes * 60 + seconds
                        self.time_to_path[time_sec] = img_path
                    
                    image_count += 1
                
                # 也直接在資料夾根目錄搜索
                for img_path in glob.glob(os.path.join(folder, pattern)):
                    filename = os.path.basename(img_path)
                    
                    # 避免重複
                    if filename not in self.filename_to_path:
                        self.filename_to_path[filename] = img_path
                        
                        time_match = re.search(r't(\d+)m([\d.]+)s', filename)
                        if time_match:
                            minutes = int(time_match.group(1))
                            seconds = float(time_match.group(2))
                            time_sec = minutes * 60 + seconds
                            self.time_to_path[time_sec] = img_path
                        
                        image_count += 1
            
            logger.info(f"  {folder_name}: {image_count} 張圖片")
            total_images += image_count
        
        logger.info(f"\n總計載入: {total_images} 張圖片")
        return total_images > 0
    
    def parse_time_hhmmss(self, time_str):
        """解析 HH:MM:SS 或 MM:SS 格式"""
        time_str = time_str.strip('（）()[]')
        
        parts = time_str.split(':')
        try:
            if len(parts) == 3:  # HH:MM:SS
                hours = int(parts[0])
                minutes = int(parts[1])
                seconds = int(parts[2])
                return hours * 3600 + minutes * 60 + seconds
            elif len(parts) == 2:  # MM:SS
                minutes = int(parts[0])
                seconds = int(parts[1])
                return minutes * 60 + seconds
        except:
            pass
        return None
    
    def find_image_by_filename(self, filename):
        """根據檔名找圖片"""
        # 直接查找
        if filename in self.filename_to_path:
            return self.filename_to_path[filename]
        
        # 嘗試不同的變體
        # 移除可能的路徑部分
        if '/' in filename:
            filename = filename.split('/')[-1]
            if filename in self.filename_to_path:
                return self.filename_to_path[filename]
        
        # 嘗試相似的檔名（處理可能的編碼問題）
        for stored_name, path in self.filename_to_path.items():
            if filename.lower() == stored_name.lower():
                return path
            # 移除空格比較
            if filename.replace(' ', '') == stored_name.replace(' ', ''):
                return path
        
        return None
    
    def find_image_by_time(self, time_seconds):
        """根據時間找最接近的圖片"""
        if not self.time_to_path:
            return None
            
        # 找最接近的時間（容差30秒）
        best_match = None
        min_diff = float('inf')
        
        for img_time, path in self.time_to_path.items():
            diff = abs(img_time - time_seconds)
            if diff < min_diff and diff <= 30:
                min_diff = diff
                best_match = path
        
        return best_match
    
    def fix_docx(self, docx_path, output_path=None):
        """修復 DOCX 文件中的圖片"""
        logger.info(f"\n📄 處理 DOCX: {docx_path}")
        
        # 載入所有圖片
        if not self.load_all_images():
            logger.error("沒有找到任何圖片！")
            return False
        
        # 載入 DOCX
        doc = Document(docx_path)
        new_doc = Document()
        
        # 統計
        images_found = 0
        images_inserted = 0
        missing_images = []
        
        # 處理每個段落
        for para in doc.paragraphs:
            text = para.text
            images_to_insert = []
            
            # 檢查格式 1: [IMAGE: filename.jpg]
            filename_matches = re.findall(r'\[IMAGE:\s*([^\]]+)\]', text)
            if filename_matches:
                images_found += len(filename_matches)
                for match in filename_matches:
                    match = match.strip()
                    
                    # 判斷是檔名還是時間格式
                    if re.match(r'^\d{2}:\d{2}(:\d{2})?$', match):
                        # 這是時間格式 (MM:SS 或 HH:MM:SS)
                        time_seconds = self.parse_time_hhmmss(match)
                        if time_seconds is not None:
                            img_path = self.find_image_by_time(time_seconds)
                            if img_path:
                                images_to_insert.append(img_path)
                            else:
                                logger.warning(f"  ⚠️  找不到時間 {match} ({time_seconds}秒) 的圖片")
                                missing_images.append(f"時間: {match}")
                    else:
                        # 這是檔名
                        img_path = self.find_image_by_filename(match)
                        if img_path:
                            images_to_insert.append(img_path)
                        else:
                            logger.warning(f"  ⚠️  找不到檔案: {match}")
                            missing_images.append(f"檔名: {match}")
            
            # 檢查格式 2: > 🖼️ 投影片圖表說明（HH:MM:SS）
            if '🖼️' in text and '（' in text:
                # 提取時間戳
                time_matches = re.findall(r'[（\(](\d{2}:\d{2}:\d{2})[）\)]', text)
                if not time_matches:
                    # 嘗試 MM:SS 格式
                    time_matches = re.findall(r'[（\(](\d{1,2}:\d{2})[）\)]', text)
                
                if time_matches:
                    images_found += len(time_matches)
                    for time_str in time_matches:
                        time_seconds = self.parse_time_hhmmss(time_str)
                        if time_seconds is not None:
                            img_path = self.find_image_by_time(time_seconds)
                            if img_path:
                                images_to_insert.append(img_path)
                            else:
                                logger.warning(f"  ⚠️  找不到時間 {time_str} ({time_seconds}秒) 的圖片")
                                missing_images.append(f"時間: {time_str}")
            
            # 檢查格式 3: 單獨的時間戳 (HH:MM:SS) - 只有小括號
            elif re.search(r'\(\d{2}:\d{2}:\d{2}\)', text):
                time_matches = re.findall(r'\((\d{2}:\d{2}:\d{2})\)', text)
                if time_matches:
                    images_found += len(time_matches)
                    for time_str in time_matches:
                        time_seconds = self.parse_time_hhmmss(time_str)
                        if time_seconds is not None:
                            img_path = self.find_image_by_time(time_seconds)
                            if img_path:
                                images_to_insert.append(img_path)
                            else:
                                logger.warning(f"  ⚠️  找不到時間 ({time_str}) = {time_seconds}秒 的圖片")
                                missing_images.append(f"時間: ({time_str})")
            
            # 插入段落文字
            new_para = new_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                try:
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                except:
                    pass
            try:
                new_para.alignment = para.alignment
            except:
                pass
            
            # 插入找到的圖片
            for img_path in images_to_insert:
                try:
                    new_doc.add_paragraph()  # 空行
                    new_doc.add_picture(img_path, width=Inches(5.5))
                    new_doc.add_paragraph()  # 空行
                    images_inserted += 1
                    logger.info(f"  ✓ 插入: {os.path.basename(img_path)}")
                except Exception as e:
                    logger.error(f"  ✗ 插入失敗 {os.path.basename(img_path)}: {e}")
        
        # 儲存結果
        if output_path is None:
            output_path = Path(docx_path).with_suffix('.fixed_images.docx')
        
        new_doc.save(str(output_path))
        
        # 顯示結果
        logger.info(f"\n📊 處理結果:")
        logger.info(f"  圖片標記: {images_found} 個")
        logger.info(f"  成功插入: {images_inserted} 張")
        logger.info(f"  缺失圖片: {len(missing_images)} 個")
        
        if missing_images:
            logger.info(f"\n缺失的圖片:")
            for missing in missing_images[:10]:  # 只顯示前10個
                logger.info(f"  - {missing}")
            if len(missing_images) > 10:
                logger.info(f"  ... 還有 {len(missing_images)-10} 個")
        
        logger.info(f"\n✓ 輸出檔案: {output_path}")
        
        return True


def main():
    if len(sys.argv) < 2:
        print("用法: python fix_merged_docx_images.py <docx檔案>")
        print("\n會自動尋找 'If You Had to Pick Just One' 的所有圖片資料夾")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    # 設定預設的圖片資料夾
    base_path = "/Volumes/WD_BLACK/國際年會/ADA2025/If You Had to Pick Just One"
    folders = [
        os.path.join(base_path, "1. GLP-1 RA"),
        os.path.join(base_path, "2. GLP-1GIP RA"),
        os.path.join(base_path, "3. SGLT2i"),
        os.path.join(base_path, "4. Rebuttal GLP-1 RA"),
        os.path.join(base_path, "5. Rebuttal SGLT2i"),
        os.path.join(base_path, "6. Rebuttal GLPGIP"),
    ]
    
    # 執行修復
    fixer = MergedDocxImageFixer()
    fixer.add_slide_folders(folders)
    fixer.fix_docx(docx_file)


if __name__ == "__main__":
    main()