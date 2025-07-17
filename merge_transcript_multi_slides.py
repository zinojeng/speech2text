#!/usr/bin/env python3
"""
合併演講稿與多個投影片內容處理程式
Merge Transcript with Multiple Slides Script for ADA 2025 Conference

此程式用於智能合併演講稿與多個投影片內容：
1. 以演講者內容為主軸，保留完整演講內容
2. 將多個投影片內容按順序作為補充說明
3. 支援每個投影片對應的圖片資料夾
4. 使用 Gemini-2.5-pro 進行內容整合與潤稿
5. 生成結構化的 Markdown 和 Word 文件

使用方法:
    # 單一投影片
    python merge_transcript_multi_slides.py transcript.txt slides1.md
    
    # 多個投影片（無圖片）
    python merge_transcript_multi_slides.py transcript.txt slides1.md slides2.md --output output_name
    
    # 多個投影片與對應圖片
    python merge_transcript_multi_slides.py transcript.txt slides1.md:images1/ slides2.md:images2/ --output output_name
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import Optional, Tuple, List, Dict
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from datetime import datetime
import json
import glob
from dotenv import load_dotenv

# 載入環境變數
load_dotenv()

# 導入圖片分析功能
try:
    from image_analyzer import analyze_image
    IMAGE_ANALYSIS_AVAILABLE = True
except ImportError:
    IMAGE_ANALYSIS_AVAILABLE = False
    logging.warning("image_analyzer 模組未找到，圖片分析功能將不可用")

# 設定日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('merge_transcript_multi_slides.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# API 配置
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

# 系統提示詞
SYSTEM_PROMPT = """你是一位專業的醫學會議內容編輯，專精於整合演講稿與多個投影片內容。這是American Diabetes Association 2025年會的內容。

你的任務是智能合併演講稿與多個投影片內容，創建一份完整、流暢的會議筆記。

**整合原則：**

1. **以演講者內容為主軸**：
   - 保留演講者的完整論述，包含口語表達的生動性
   - 對演講內容進行適度潤稿，使其更加流暢專業
   - 保持演講者的觀點和強調重點

2. **多個投影片內容的運用**：
   - 按照提供的順序整合多個投影片內容
   - 作為補充說明，豐富演講內容
   - 填補演講中未提及但重要的資訊
   - 不要重複已在演講中詳細說明的內容
   - 當有多個投影片時，確保內容銜接流暢

3. **內容組織方式**：
   - 按照演講的邏輯順序組織內容
   - 在適當位置插入各個投影片的補充資訊
   - 使用 __底線__ 標記延伸解讀或重要補充說明
   - 保持段落之間的流暢過渡

4. **格式要求**：
   - 使用 # ## ### 組織章節
   - **粗體**標記重要概念、藥物名稱、關鍵數據
   - __底線__用於標記：
     * 來自投影片的重要補充
     * 延伸解讀和深入說明
     * 關鍵研究發現或結論
   - 適度使用項目列表呈現並列資訊

5. **寫作風格**：
   - 保持學術專業性，但不失演講的生動性
   - 確保內容完整、詳細、易讀
   - 使用繁體中文

**圖片整合方式：**
當提供投影片圖片時，請：
1. 根據時間戳記在適當位置插入圖片標記
2. 使用以下兩種格式之一：
   - 簡單標記：[IMAGE: MM:SS] 或 [IMAGE: HH:MM:SS]
   - 帶說明標記：
     > 🖼️ **投影片圖表說明**（HH:MM:SS）：
     > [圖片分析內容]
     > __[與演講內容的關聯或延伸解讀]__
3. 重要：只使用時間格式，不要使用檔名（如 slide001.jpg）
4. 確保圖片分析與演講內容相互呼應
5. 不要重複已在演講中詳細說明的圖表內容

請確保輸出是一份完整、專業、資訊豐富的會議筆記。"""


class MultiSlidesProcessor:
    """處理演講稿與多個投影片合併的處理器"""
    
    def __init__(self):
        """初始化處理器"""
        self.setup_api()
        self.all_slide_images = {}  # 儲存所有投影片的圖片 {time_sec: [(slide_index, img_path), ...]}
    
    def setup_api(self):
        """設定 Google Gemini API"""
        try:
            if not GOOGLE_API_KEY:
                raise ValueError("請在 .env 檔案中設定 GOOGLE_API_KEY")
            genai.configure(api_key=GOOGLE_API_KEY)
            logger.info("Google Gemini API 設定完成")
        except Exception as e:
            logger.error(f"API 設定失敗: {e}")
            raise
    
    def parse_slide_input(self, slide_input: str) -> Tuple[str, Optional[str]]:
        """
        解析投影片輸入，支援 "slides.md:images/" 格式
        
        Args:
            slide_input: 投影片輸入字串
            
        Returns:
            (投影片檔案路徑, 圖片資料夾路徑)
        """
        if ':' in slide_input:
            parts = slide_input.split(':', 1)
            return parts[0], parts[1]
        return slide_input, None
    
    def read_file(self, file_path: str, file_type: str) -> str:
        """
        讀取檔案內容
        
        Args:
            file_path: 檔案路徑
            file_type: 檔案類型描述（用於日誌）
            
        Returns:
            檔案內容
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"{file_type}檔案不存在: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            if not content:
                raise ValueError(f"{file_type}檔案內容為空")
            
            logger.info(f"成功讀取{file_type}檔案: {file_path}")
            logger.info(f"檔案大小: {len(content)} 字元")
            return content
            
        except Exception as e:
            logger.error(f"讀取{file_type}檔案失敗: {e}")
            raise
    
    def parse_slide_time(self, filename: str) -> Optional[float]:
        """
        從檔名解析時間戳記
        例如: slide_009_t1m4.7s.jpg -> 64.7 秒
        """
        import re
        match = re.search(r't(\d+)m([\d.]+)s', filename)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        return None
    
    def parse_time_format(self, time_str: str) -> Optional[float]:
        """
        解析各種時間格式
        支援: "3m34.7s", "214.7", "214.7s", "t1m4.7s", "0:38", "1:14", "00:04:28"
        """
        import re
        
        # 移除開頭的 't' 前綴（如果存在）
        if time_str.startswith('t'):
            time_str = time_str[1:]
        
        # 純數字
        try:
            return float(time_str)
        except ValueError:
            pass
        
        # 數字+s
        if time_str.endswith('s'):
            try:
                return float(time_str[:-1])
            except ValueError:
                pass
        
        # HH:MM:SS 格式（如 00:04:28）
        match = re.match(r'(\d{1,2}):(\d{2}):(\d{2})', time_str)
        if match:
            hours = int(match.group(1))
            minutes = int(match.group(2))
            seconds = int(match.group(3))
            return hours * 3600 + minutes * 60 + seconds
        
        # M:SS 格式（如 0:38, 1:14）
        match = re.match(r'(\d+):(\d+\.?\d*)', time_str)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        
        # 分鐘格式
        match = re.match(r'(\d+)m([\d.]+)s', time_str)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        
        return None
    
    def load_slide_images(self, images_folder: str, slide_index: int) -> Dict[float, str]:
        """
        載入投影片圖片並按時間排序
        
        Args:
            images_folder: 圖片資料夾路徑
            slide_index: 投影片索引（用於區分不同投影片的圖片）
            
        Returns:
            Dict[float, str]: 時間戳記到圖片路徑的映射
        """
        images = {}
        if not images_folder or not os.path.exists(images_folder):
            if images_folder:
                logger.warning(f"圖片資料夾不存在: {images_folder}")
            return images
        
        # 支援的圖片格式
        patterns = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']
        
        for pattern in patterns:
            for img_path in glob.glob(os.path.join(images_folder, pattern)):
                time_sec = self.parse_slide_time(os.path.basename(img_path))
                if time_sec is not None:
                    images[time_sec] = img_path
                    # 使用列表來儲存多個投影片在相同時間的圖片
                    if time_sec not in self.all_slide_images:
                        self.all_slide_images[time_sec] = []
                    self.all_slide_images[time_sec].append((slide_index, img_path))
        
        logger.info(f"從投影片 {slide_index+1} 載入了 {len(images)} 張圖片")
        if images:
            logger.debug(f"投影片 {slide_index+1} 圖片時間範圍: {min(images.keys()):.1f}s - {max(images.keys()):.1f}s")
        logger.debug(f"總共儲存了 {len(self.all_slide_images)} 個不同時間點的圖片")
        return images
    
    def merge_with_gemini(self, transcript: str, slides_contents: List[Tuple[str, str, Optional[str]]]) -> str:
        """
        使用 Gemini-2.5-pro 進行內容合併與整合
        
        Args:
            transcript: 演講稿內容
            slides_contents: [(檔名, 內容, 圖片資料夾), ...] 的列表
            
        Returns:
            整合後的 Markdown 內容
        """
        try:
            logger.info("開始使用 Gemini-2.5-pro 進行內容整合")
            
            # 建立模型
            model = genai.GenerativeModel('gemini-2.5-pro')
            
            # 構建投影片內容部分
            slides_text = ""
            image_info = ""
            
            for i, (filename, content, images_folder) in enumerate(slides_contents):
                slides_text += f"\n\n=== 投影片 {i+1}: {filename} ===\n{content}"
                
                if images_folder:
                    slide_images = self.load_slide_images(images_folder, i)
                    if slide_images:
                        image_info += f"\n\n投影片 {i+1} 包含 {len(slide_images)} 張圖片"
            
            # 載入圖片資訊
            if self.all_slide_images:
                total_images = sum(len(img_list) for img_list in self.all_slide_images.values())
                image_info = f"\n\n=== 投影片圖片資訊 ===\n共有 {total_images} 張投影片圖片（{len(self.all_slide_images)} 個不同時間點）。\n重要：請在整合內容時，使用以下格式插入圖片：\n1. 對於有時間戳的圖片：使用 [IMAGE: 時間] 格式，例如 [IMAGE: 3m28s] 或 [IMAGE: 00:03:28]\n2. 不要使用檔名格式如 [IMAGE: slide001.jpg]\n3. 時間格式可以是 MM:SS 或 HH:MM:SS"
                logger.info(f"總共載入 {total_images} 張圖片，分佈在 {len(self.all_slide_images)} 個時間點")
            
            # 構建提示詞
            user_prompt = f"""請根據以下演講稿和多個投影片內容，創建一份整合的會議筆記：

=== 演講稿內容 ===
{transcript}

=== 投影片內容 ==={slides_text}{image_info}

請按照指示整合這些內容，生成完整的 Markdown 格式會議筆記。記住投影片內容是用來補充和加強演講者的論述，不要重複相同內容。"""
            
            # 生成整合內容
            response = model.generate_content([
                {"role": "user", "parts": [{"text": SYSTEM_PROMPT}]},
                {"role": "user", "parts": [{"text": user_prompt}]}
            ])
            
            merged_content = response.text
            logger.info(f"內容整合完成，長度: {len(merged_content)} 字元")
            return merged_content
            
        except Exception as e:
            logger.error(f"內容整合失敗: {e}")
            raise
    
    def save_markdown(self, content: str, output_path: str) -> str:
        """
        保存 Markdown 檔案，並處理圖片標記
        
        Args:
            content: Markdown 內容
            output_path: 輸出路徑
            
        Returns:
            保存的檔案路徑
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 處理圖片標記，替換為實際的 Markdown 圖片語法
            if self.all_slide_images:
                total_images = sum(len(img_list) for img_list in self.all_slide_images.values())
                logger.info(f"處理 Markdown 中的圖片標記，共有 {total_images} 張圖片在 {len(self.all_slide_images)} 個時間點")
                lines = content.split('\n')
                processed_lines = []
                
                for line in lines:
                    # 用於收集該行所有圖片
                    images_to_insert = []
                    
                    # 處理原始格式 [IMAGE: time] (支援 t1m4.7s 格式)
                    if '[IMAGE:' in line:
                        matches = re.findall(r'\[IMAGE:\s*([^\]]+)\]', line)
                        for time_str in matches:
                            target_time = self.parse_time_format(time_str)
                            if target_time is not None:
                                img_line = self._insert_image_markdown('', target_time, output_path.parent)
                                if img_line:
                                    images_to_insert.append(img_line)
                    
                    # 處理 Gemini 生成的格式：> 🖼️ **投影片圖表說明**（[時間戳]）：
                    # 支援多個時間戳，如：（[53m25.2s], [54m0.2s]）或（[53m25.2s] 與 [54m0.2s]）
                    if '🖼️' in line:
                        # 嘗試匹配中文括號格式
                        bracket_match = re.search(r'（\[([^\)]+)\]）', line)
                        if not bracket_match:
                            # 嘗試匹配英文括號格式
                            bracket_match = re.search(r'\(\[([^\)]+)\]\)', line)
                        
                        if bracket_match:
                            # 提取括號內的所有內容
                            bracket_content = bracket_match.group(1)
                            # 找出所有時間戳
                            time_matches = re.findall(r'\[?([^\[\],]+?)(?:\]|,|$)', bracket_content)
                            
                            for time_str in time_matches:
                                # 清理時間字串
                                time_str = time_str.strip().strip(']').strip()
                                if time_str and time_str not in ['與', 'and', '和']:  # 排除連接詞
                                    target_time = self.parse_time_format(time_str)
                                    if target_time is not None:
                                        img_line = self._insert_image_markdown('', target_time, output_path.parent)
                                        if img_line:
                                            images_to_insert.append(img_line)
                    
                    # 插入所有找到的圖片
                    if images_to_insert:
                        for img in images_to_insert:
                            processed_lines.append(img)
                        processed_lines.append('')  # 空行
                    
                    processed_lines.append(line)
                
                content = '\n'.join(processed_lines)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown 檔案已保存: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"保存 Markdown 失敗: {e}")
            raise
    
    def _insert_image_markdown(self, line: str, target_time: float, base_path: Path) -> Optional[str]:
        """
        在 Markdown 中插入圖片
        
        Returns:
            替換後的行，或 None 如果沒有找到合適的圖片
        """
        if not self.all_slide_images:
            return None
            
        # 找到最接近的圖片
        closest_time = min(self.all_slide_images.keys(), 
                         key=lambda x: abs(x - target_time))
        if abs(closest_time - target_time) < 30:  # 30秒容差
            # 從列表中選擇最合適的圖片
            slide_images = self.all_slide_images[closest_time]
            if slide_images:
                # 如果有多個圖片，優先選擇較早的投影片（按slide_index排序）
                slide_images_sorted = sorted(slide_images, key=lambda x: x[0])
                slide_index, img_path = slide_images_sorted[0]
                # 轉換為相對路徑
                img_relative = os.path.relpath(img_path, base_path)
                logger.info(f"插入圖片: {target_time}s -> {os.path.basename(img_path)} (投影片 {slide_index+1})")
                return f"![投影片 {closest_time:.1f}s]({img_relative})"
        return None
    
    def _insert_image_docx(self, doc: Document, target_time: float) -> bool:
        """
        在 Word 文件中插入圖片
        
        Returns:
            是否成功插入
        """
        if not self.all_slide_images:
            return False
            
        # 找到最接近的圖片
        closest_time = min(self.all_slide_images.keys(), 
                         key=lambda x: abs(x - target_time))
        if abs(closest_time - target_time) < 30:  # 30秒容差
            # 從列表中選擇最合適的圖片
            slide_images = self.all_slide_images[closest_time]
            if slide_images:
                # 如果有多個圖片，優先選擇較早的投影片（按slide_index排序）
                slide_images_sorted = sorted(slide_images, key=lambda x: x[0])
                slide_index, img_path = slide_images_sorted[0]
                if os.path.exists(img_path):
                    try:
                        doc.add_paragraph()  # 空行
                        doc.add_picture(img_path, width=Inches(5.5))
                        doc.add_paragraph()  # 空行
                        logger.info(f"插入圖片到 Word: {os.path.basename(img_path)} (時間: {closest_time}秒, 投影片 {slide_index+1})")
                        return True
                    except Exception as e:
                        logger.warning(f"插入圖片失敗: {e}")
        return False
    
    def _insert_image_docx_by_path(self, doc: Document, img_path: str) -> bool:
        """
        在 Word 文件中根據路徑插入圖片
        
        Args:
            doc: Word document object
            img_path: Image file path (can be relative or contain folder)
            
        Returns:
            是否成功插入
        """
        # Try to find the image in various locations
        possible_paths = []
        
        # Direct path
        possible_paths.append(img_path)
        
        # If it's a relative path with folder
        if '/' in img_path:
            # Try from current directory
            possible_paths.append(Path.cwd() / img_path)
            # Try from document directory (if we know it)
            possible_paths.append(Path(img_path))
        
        # Try each possible path
        for path in possible_paths:
            if os.path.exists(path):
                try:
                    doc.add_paragraph()  # 空行
                    doc.add_picture(str(path), width=Inches(5.5))
                    doc.add_paragraph()  # 空行
                    logger.info(f"插入圖片到 Word: {path}")
                    return True
                except Exception as e:
                    logger.warning(f"插入圖片失敗 {path}: {e}")
        
        logger.warning(f"找不到圖片: {img_path}")
        return False
    
    def markdown_to_docx(self, markdown_text: str, output_path: str) -> bool:
        """
        將 Markdown 文字轉換為保留格式的 Word 文件
        
        Args:
            markdown_text: Markdown 格式文字
            output_path: 輸出檔案路徑
            
        Returns:
            轉換是否成功
        """
        try:
            logger.info(f"開始轉換為 Word 文件: {output_path}")
            
            doc = Document()
            
            # 添加標題
            title = doc.add_heading('ADA 2025 會議筆記 - 演講與多投影片整合版', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加日期
            date_para = doc.add_paragraph()
            date_para.add_run(f"整合日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_paragraph()  # 空行
            
            # 分行處理 Markdown 內容
            lines = markdown_text.split('\n')
            
            for line in lines:
                # 用於收集該行所有圖片
                images_inserted = []
                
                # 處理原始格式 [IMAGE: time] (支援 t1m4.7s 格式)
                # 也支援 [IMAGE: folder/filename.jpg] 格式
                if '[IMAGE:' in line:
                    matches = re.findall(r'\[IMAGE:\s*([^\]]+)\]', line)
                    for match in matches:
                        # Check if it's a file path (contains / or .)
                        if '/' in match or '.' in match:
                            # It's a file path
                            if self._insert_image_docx_by_path(doc, match):
                                images_inserted.append(match)
                        else:
                            # It's a timestamp
                            target_time = self.parse_time_format(match)
                            if target_time is not None and self.all_slide_images:
                                if self._insert_image_docx(doc, target_time):
                                    images_inserted.append(target_time)
                    
                    # 如果插入了圖片，跳過這一行文字
                    if images_inserted and '[IMAGE:' in line and line.strip().startswith('[IMAGE:'):
                        continue
                
                # 處理 Gemini 生成的格式：> 🖼️ **投影片圖表說明**（[時間戳]）：
                # 支援多個時間戳，如：（[53m25.2s], [54m0.2s]）或（[53m25.2s] 與 [54m0.2s]）
                if self.all_slide_images and '🖼️' in line:
                    # 嘗試匹配中文括號格式
                    bracket_match = re.search(r'（\[([^\)]+)\]）', line)
                    if not bracket_match:
                        # 嘗試匹配英文括號格式
                        bracket_match = re.search(r'\(\[([^\)]+)\]\)', line)
                    
                    if bracket_match:
                        # 提取括號內的所有內容
                        bracket_content = bracket_match.group(1)
                        # 找出所有時間戳
                        time_matches = re.findall(r'\[?([^\[\],]+?)(?:\]|,|$)', bracket_content)
                        
                        for time_str in time_matches:
                            # 清理時間字串
                            time_str = time_str.strip().strip(']').strip()
                            if time_str and time_str not in ['與', 'and', '和']:  # 排除連接詞
                                target_time = self.parse_time_format(time_str)
                                if target_time is not None:
                                    if self._insert_image_docx(doc, target_time):
                                        images_inserted.append(target_time)
                    
                    # 繼續處理此行文字（但不包含圖片標記）
                
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # 空行
                    continue
                
                # 處理標題
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    
                    if level == 1:
                        doc.add_heading(title_text, level=1)
                    elif level == 2:
                        doc.add_heading(title_text, level=2)
                    elif level == 3:
                        doc.add_heading(title_text, level=3)
                    else:
                        doc.add_heading(title_text, level=4)
                    
                    continue
                
                # 處理列表項目
                if line.startswith(('- ', '* ', '+ ')):
                    list_text = line[2:].strip()
                    paragraph = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(paragraph, list_text)
                elif re.match(r'^\d+\.\s', line):
                    list_text = re.sub(r'^\d+\.\s', '', line).strip()
                    paragraph = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(paragraph, list_text)
                else:
                    # 處理普通段落
                    paragraph = doc.add_paragraph()
                    self._add_formatted_text(paragraph, line)
            
            # 儲存文件
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logger.info(f"Word 文件已儲存: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"轉換為 Word 文件失敗: {e}")
            return False
    
    def _add_formatted_text(self, paragraph, text):
        """
        處理文字格式並添加到段落
        
        Args:
            paragraph: Word 段落對象
            text: 要處理的文字
        """
        # 使用更精確的正則表達式來處理格式
        # 先處理雙底線（必須在單底線之前）
        parts = re.split(r'(__[^_]+__|_[^_]+_|\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # 粗體
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__') and len(part) > 4:
                # 底線（雙底線）
                run = paragraph.add_run(part[2:-2])
                run.underline = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                # 底線（單底線）
                run = paragraph.add_run(part[1:-1])
                run.underline = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                # 斜體
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # 普通文字
                paragraph.add_run(part)
    
    def process_files(self, transcript_file: str, slides_inputs: List[str], output_base: str = None) -> dict:
        """
        處理演講稿與多個投影片檔案
        
        Args:
            transcript_file: 演講稿檔案路徑
            slides_inputs: 投影片輸入列表（可包含圖片資料夾）
            output_base: 輸出檔案基礎名稱（可選）
            
        Returns:
            處理結果
        """
        result = {
            'success': False,
            'transcript_file': transcript_file,
            'slides_files': [],
            'markdown_file': None,
            'docx_file': None,
            'error': None
        }
        
        try:
            # 讀取演講稿
            transcript = self.read_file(transcript_file, "演講稿")
            
            # 讀取所有投影片
            slides_contents = []
            for slide_input in slides_inputs:
                slide_file, images_folder = self.parse_slide_input(slide_input)
                slide_content = self.read_file(slide_file, f"投影片")
                slides_contents.append((os.path.basename(slide_file), slide_content, images_folder))
                result['slides_files'].append(slide_file)
            
            # 使用 Gemini 進行內容整合
            merged_content = self.merge_with_gemini(transcript, slides_contents)
            
            # 準備輸出路徑
            if not output_base:
                # 使用演講稿檔案名作為基礎
                transcript_path = Path(transcript_file)
                output_base = transcript_path.stem
            
            output_dir = Path(transcript_file).parent
            
            # 保存 Markdown（包含圖片處理）
            markdown_path = output_dir / f"{output_base}_multi_merged.md"
            self.save_markdown(merged_content, str(markdown_path))
            result['markdown_file'] = str(markdown_path)
            
            # 轉換為 Word
            docx_path = output_dir / f"{output_base}_multi_merged.docx"
            if self.markdown_to_docx(merged_content, str(docx_path)):
                result['docx_file'] = str(docx_path)
                result['success'] = True
            else:
                result['error'] = "Word 轉換失敗"
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"處理檔案失敗: {e}")
        
        return result


def main():
    """主函數"""
    import argparse
    
    # 自定義幫助訊息
    usage_text = """
使用範例:
  # 單一投影片（無圖片）
  python %(prog)s transcript.txt slides1.md
  
  # 多個投影片（無圖片）
  python %(prog)s transcript.txt slides1.md slides2.md --output merged_notes
  
  # 單一投影片與圖片
  python %(prog)s transcript.txt slides1.md:images1/
  
  # 多個投影片與對應圖片
  python %(prog)s transcript.txt slides1.md:images1/ slides2.md:images2/ --output final_notes
  
  # 混合使用（部分有圖片）
  python %(prog)s transcript.txt slides1.md slides2.md:images2/ slides3.md:images3/
"""
    
    parser = argparse.ArgumentParser(
        description='合併演講稿與多個投影片內容（支援圖片）',
        epilog=usage_text,
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument('transcript_file', help='演講稿檔案路徑')
    parser.add_argument('slides', nargs='+', 
                       help='投影片檔案，格式: slides.md 或 slides.md:images/')
    parser.add_argument('--output', '-o', dest='output_base',
                       help='輸出檔案基礎名稱（預設使用演講稿檔名）')
    
    args = parser.parse_args()
    
    print(f"\n=== 演講稿與多投影片整合程式 ===")
    print(f"演講稿檔案: {args.transcript_file}")
    print(f"投影片數量: {len(args.slides)}")
    
    # 顯示投影片資訊
    for i, slide in enumerate(args.slides, 1):
        if ':' in slide:
            slide_file, img_folder = slide.split(':', 1)
            print(f"  投影片 {i}: {slide_file} (圖片: {img_folder})")
        else:
            print(f"  投影片 {i}: {slide}")
    
    print(f"使用模型: Gemini-2.5-pro")
    print("處理中...\n")
    
    try:
        processor = MultiSlidesProcessor()
        result = processor.process_files(args.transcript_file, args.slides, args.output_base)
        
        if result['success']:
            print("\n✅ 處理成功！")
            print(f"處理了 {len(result['slides_files'])} 個投影片檔案")
            print(f"Markdown 檔案: {result['markdown_file']}")
            print(f"Word 檔案: {result['docx_file']}")
        else:
            print(f"\n❌ 處理失敗: {result['error']}")
            sys.exit(1)
            
    except Exception as e:
        print(f"\n❌ 程式執行失敗: {e}")
        logger.error(f"程式執行失敗: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()