#!/usr/bin/env python3
"""
直接修復 DOCX 文件，插入對應時間的圖片
"""

import os
import sys
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
import glob

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


class DocxImageFixer:
    def __init__(self):
        self.image_map = {}  # 時間(秒) -> 圖片路徑
        
    def parse_time(self, time_str):
        """解析各種時間格式"""
        time_str = time_str.strip('[]')
        
        # MM:SS 或 M:SS 格式
        if ':' in time_str:
            parts = time_str.split(':')
            if len(parts) == 2:
                try:
                    minutes = int(parts[0])
                    seconds = int(parts[1])
                    return minutes * 60 + seconds
                except:
                    pass
            elif len(parts) == 3:  # HH:MM:SS
                try:
                    hours = int(parts[0])
                    minutes = int(parts[1])
                    seconds = int(parts[2])
                    return hours * 3600 + minutes * 60 + seconds
                except:
                    pass
        
        return None
    
    def load_images_from_folder(self, folder_path):
        """載入資料夾中的所有圖片並建立時間映射"""
        if not os.path.exists(folder_path):
            logger.error(f"❌ 圖片資料夾不存在: {folder_path}")
            return False
        
        logger.info(f"📁 掃描圖片資料夾: {folder_path}")
        
        # 尋找所有圖片
        patterns = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']
        image_count = 0
        
        for pattern in patterns:
            for img_path in glob.glob(os.path.join(folder_path, pattern)):
                filename = os.path.basename(img_path)
                
                # 解析檔名中的時間 (格式: slide_XXX_tXmYs.jpg)
                match = re.search(r't(\d+)m([\d.]+)s', filename)
                if match:
                    minutes = int(match.group(1))
                    seconds = float(match.group(2))
                    time_in_seconds = minutes * 60 + seconds
                    self.image_map[time_in_seconds] = img_path
                    image_count += 1
                    logger.debug(f"  找到: {filename} -> {time_in_seconds}秒")
        
        logger.info(f"✓ 載入了 {image_count} 張圖片")
        
        # 顯示時間範圍
        if self.image_map:
            min_time = min(self.image_map.keys())
            max_time = max(self.image_map.keys())
            logger.info(f"  時間範圍: {min_time}秒 ~ {max_time}秒")
        
        return image_count > 0
    
    def find_closest_image(self, target_seconds, tolerance=30):
        """找到最接近目標時間的圖片"""
        if not self.image_map:
            return None
        
        closest_time = min(self.image_map.keys(), 
                          key=lambda x: abs(x - target_seconds))
        
        if abs(closest_time - target_seconds) <= tolerance:
            return self.image_map[closest_time]
        
        return None
    
    def fix_docx(self, docx_path, image_folder, output_path=None):
        """修復 DOCX 文件，插入實際圖片"""
        
        # 載入圖片
        if not self.load_images_from_folder(image_folder):
            logger.error("無法載入圖片，程序終止")
            return False
        
        # 載入 DOCX
        logger.info(f"\n📄 處理 DOCX: {docx_path}")
        doc = Document(docx_path)
        new_doc = Document()
        
        # 統計
        timestamps_found = 0
        images_inserted = 0
        
        # 處理每個段落
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            
            # 檢查是否包含圖片時間戳
            if '🖼️' in text and '[' in text:
                # 提取所有時間戳
                timestamps = re.findall(r'\[([^\]]+)\]', text)
                
                if timestamps:
                    timestamps_found += len(timestamps)
                    
                    # 先添加說明文字
                    new_para = new_doc.add_paragraph()
                    self._copy_paragraph_format(para, new_para)
                    new_para.add_run(text)
                    
                    # 為每個時間戳插入圖片
                    for ts in timestamps:
                        time_seconds = self.parse_time(ts)
                        if time_seconds is not None:
                            logger.info(f"\n⏰ 處理時間戳 [{ts}] = {time_seconds}秒")
                            
                            img_path = self.find_closest_image(time_seconds)
                            if img_path:
                                try:
                                    # 插入圖片
                                    new_doc.add_paragraph()  # 空行
                                    new_doc.add_picture(img_path, width=Inches(5.5))
                                    new_doc.add_paragraph()  # 空行
                                    
                                    images_inserted += 1
                                    logger.info(f"  ✓ 插入圖片: {os.path.basename(img_path)}")
                                except Exception as e:
                                    logger.error(f"  ❌ 插入失敗: {e}")
                            else:
                                logger.warning(f"  ⚠️  找不到匹配的圖片")
                        else:
                            logger.warning(f"  ⚠️  無法解析時間: {ts}")
                else:
                    # 有圖示但沒有時間戳，直接複製
                    new_para = new_doc.add_paragraph()
                    self._copy_paragraph_format(para, new_para)
                    self._copy_paragraph_runs(para, new_para)
            else:
                # 普通段落，直接複製
                new_para = new_doc.add_paragraph()
                self._copy_paragraph_format(para, new_para)
                self._copy_paragraph_runs(para, new_para)
        
        # 儲存結果
        if output_path is None:
            output_path = Path(docx_path).with_suffix('.with_images.docx')
        
        new_doc.save(str(output_path))
        
        # 顯示結果
        logger.info(f"\n📊 處理結果:")
        logger.info(f"  找到時間戳: {timestamps_found} 個")
        logger.info(f"  插入圖片: {images_inserted} 張")
        logger.info(f"  輸出檔案: {output_path}")
        
        return images_inserted > 0
    
    def _copy_paragraph_format(self, source, target):
        """複製段落格式"""
        try:
            target.alignment = source.alignment
            target.paragraph_format.space_before = source.paragraph_format.space_before
            target.paragraph_format.space_after = source.paragraph_format.space_after
            target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
        except:
            pass
    
    def _copy_paragraph_runs(self, source, target):
        """複製段落內容和格式"""
        for run in source.runs:
            new_run = target.add_run(run.text)
            try:
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
            except:
                pass


def main():
    if len(sys.argv) < 3:
        print("用法: python fix_docx_with_images.py <docx檔案> <圖片資料夾> [輸出檔名]")
        print("\n範例:")
        print('  python fix_docx_with_images.py output.docx "1. GLP-1 RA"')
        print('  python fix_docx_with_images.py output.docx "/Users/xxx/Desktop/ADA2025/1. GLP-1 RA" fixed.docx')
        sys.exit(1)
    
    docx_file = sys.argv[1]
    image_folder = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    # 執行修復
    fixer = DocxImageFixer()
    success = fixer.fix_docx(docx_file, image_folder, output_file)
    
    if success:
        print("\n✅ 修復完成！")
    else:
        print("\n❌ 修復失敗，請檢查錯誤訊息")


if __name__ == "__main__":
    main()