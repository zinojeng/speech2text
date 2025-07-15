#!/usr/bin/env python3
"""
處理已轉錄文字檔案的腳本
Process Transcription Script for ADA 2025 Conference

此程式用於處理已轉錄的文字檔案：
1. 讀取轉錄文字檔案
2. 使用 Gemini-2.5-pro 進行智能摘要
3. 生成 Markdown 格式摘要
4. 轉換為保留格式的 Word 文件

使用方法:
    python process_transcription.py <transcription_file.txt>
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import Optional
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from dotenv import load_dotenv

# 載入環境變數
load_dotenv()

# 設定日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('process_transcription.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# API 配置
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# 系統提示詞
SYSTEM_PROMPT = """This is American Diabetes Association 2025年會，聚焦在糖尿病、代謝、肥胖等主題。

請將提供的轉錄內容改寫成**學術論文風格的會議筆記**，要求如下：

**寫作風格要求：**
1. **段落式寫作**：以完整的段落敘述為主，但在需要清楚呈現多個要點時，可適度使用項目符號列表。
2. **學術寫作風格**：採用正式的學術寫作風格，句子之間要有良好的邏輯連接和過渡。
3. **重點標記方式**：
   - 在段落中使用 **粗體** 標記關鍵概念、藥物名稱、重要數據
   - 使用 __底線__ 標記最重要的研究發現或結論
   - 當有多個並列的研究結果或治療選項時，可使用簡潔的項目符號列表
4. **內容整合**：將相關的資訊整合在同一段落中，但當資訊較為複雜時，可使用列表增加可讀性。

**內容組織要求：**
1. **章節結構**：使用 # ## ### 來組織內容，讓資訊有清晰的層次
2. **靈活的段落安排**：根據內容需要自然組織段落，不必強制每個章節都有特定的段落結構
3. **引用研究**：在段落中自然地引入研究發現，包括期刊名稱、年份、樣本數等關鍵資訊
4. **數據呈現**：將統計數據和研究結果整合在敘述中，使用括號補充具體數值

**範例風格：**
不要寫成：「• 甲狀腺功能低下與MASLD風險增加相關」
而應寫成：「近期的大型統合分析顯示，**原發性甲狀腺功能低下**是MASLD發展的重要風險因子。這項發表於2024年《Gut》期刊的研究涵蓋了7600萬人的數據，發現甲狀腺功能低下患者發生MASLD的__賠率比高達1.43__，且這種關聯性在更嚴重的MASH或晚期纖維化患者中更為顯著。」

請確保輸出是專業、流暢、易讀的學術風格文章，使用繁體中文。"""


class TranscriptionProcessor:
    """處理已轉錄文字的處理器"""
    
    def __init__(self):
        """初始化處理器"""
        self.setup_api()
    
    def setup_api(self):
        """設定 Google Gemini API"""
        try:
            genai.configure(api_key=GOOGLE_API_KEY)
            logger.info("Google Gemini API 設定完成")
        except Exception as e:
            logger.error(f"API 設定失敗: {e}")
            raise
    
    def read_transcription_file(self, file_path: str) -> str:
        """
        讀取轉錄文字檔案
        
        Args:
            file_path: 文字檔案路徑
            
        Returns:
            檔案內容
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"檔案不存在: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            if not content:
                raise ValueError("檔案內容為空")
            
            logger.info(f"成功讀取檔案: {file_path}")
            logger.info(f"檔案大小: {len(content)} 字元")
            return content
            
        except Exception as e:
            logger.error(f"讀取檔案失敗: {e}")
            raise
    
    def summarize_with_gemini(self, transcription: str) -> str:
        """
        使用 Gemini-2.5-pro 進行摘要處理
        
        Args:
            transcription: 轉錄文字
            
        Returns:
            Markdown 格式的摘要
        """
        try:
            logger.info("開始使用 Gemini-2.5-pro 進行摘要處理")
            
            # 建立模型
            model = genai.GenerativeModel('gemini-2.5-pro')
            
            # 構建提示詞
            user_prompt = f"""請根據以下轉錄內容生成詳細的會議筆記：

轉錄內容：
{transcription}

請確保輸出是完整的 Markdown 格式文件，包含適當的格式化。"""
            
            # 生成摘要
            response = model.generate_content([
                {"role": "user", "parts": [{"text": SYSTEM_PROMPT}]},
                {"role": "user", "parts": [{"text": user_prompt}]}
            ])
            
            summary = response.text
            logger.info(f"摘要完成，長度: {len(summary)} 字元")
            return summary
            
        except Exception as e:
            logger.error(f"摘要處理失敗: {e}")
            raise
    
    def save_markdown(self, content: str, output_path: str) -> str:
        """
        保存 Markdown 檔案
        
        Args:
            content: Markdown 內容
            output_path: 輸出路徑
            
        Returns:
            保存的檔案路徑
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown 檔案已保存: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"保存 Markdown 失敗: {e}")
            raise
    
    def markdown_to_docx(self, markdown_text: str, output_path: str) -> bool:
        """
        將 Markdown 文字轉換為保留格式的 Word 文件
        
        Args:
            markdown_text: Markdown 格式文字
            output_path: 輸出檔案路徑
            
        Returns:
            轉換是否成功
        """
        try:
            logger.info(f"開始轉換為 Word 文件: {output_path}")
            
            doc = Document()
            
            # 添加標題
            title = doc.add_heading('ADA 2025 會議筆記', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加日期
            date_para = doc.add_paragraph()
            date_para.add_run(f"處理日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_paragraph()  # 空行
            
            # 分行處理 Markdown 內容
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # 空行
                    continue
                
                # 處理標題
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    
                    if level == 1:
                        doc.add_heading(title_text, level=1)
                    elif level == 2:
                        doc.add_heading(title_text, level=2)
                    elif level == 3:
                        doc.add_heading(title_text, level=3)
                    else:
                        doc.add_heading(title_text, level=4)
                    
                    continue
                
                # 處理列表項目
                if line.startswith(('- ', '* ', '+ ')):
                    list_text = line[2:].strip()
                    paragraph = doc.add_paragraph(list_text, style='List Bullet')
                elif re.match(r'^\d+\.\s', line):
                    list_text = re.sub(r'^\d+\.\s', '', line).strip()
                    paragraph = doc.add_paragraph(list_text, style='List Number')
                else:
                    # 處理普通段落
                    paragraph = doc.add_paragraph()
                    
                    # 處理行內格式（粗體、斜體、底線）
                    # 使用更複雜的正則表達式來處理嵌套格式
                    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|_[^_]+_|\*[^*]+\*)', line)
                    
                    for part in parts:
                        if not part:
                            continue
                            
                        if part.startswith('**') and part.endswith('**'):
                            # 粗體
                            run = paragraph.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('__') and part.endswith('__'):
                            # 底線
                            run = paragraph.add_run(part[2:-2])
                            run.underline = True
                        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                            # 斜體
                            run = paragraph.add_run(part[1:-1])
                            run.italic = True
                        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                            # 底線（單底線）
                            run = paragraph.add_run(part[1:-1])
                            run.underline = True
                        else:
                            # 普通文字
                            paragraph.add_run(part)
            
            # 儲存文件
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logger.info(f"Word 文件已儲存: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"轉換為 Word 文件失敗: {e}")
            return False
    
    def process_file(self, input_file: str) -> dict:
        """
        處理單一轉錄檔案
        
        Args:
            input_file: 輸入檔案路徑
            
        Returns:
            處理結果
        """
        result = {
            'success': False,
            'input_file': input_file,
            'markdown_file': None,
            'docx_file': None,
            'error': None
        }
        
        try:
            # 讀取轉錄檔案
            transcription = self.read_transcription_file(input_file)
            
            # 使用 Gemini 生成摘要
            summary = self.summarize_with_gemini(transcription)
            
            # 準備輸出路徑
            input_path = Path(input_file)
            output_dir = input_path.parent
            base_name = input_path.stem
            
            # 保存 Markdown
            markdown_path = output_dir / f"{base_name}_summary.md"
            self.save_markdown(summary, str(markdown_path))
            result['markdown_file'] = str(markdown_path)
            
            # 轉換為 Word
            docx_path = output_dir / f"{base_name}_summary.docx"
            if self.markdown_to_docx(summary, str(docx_path)):
                result['docx_file'] = str(docx_path)
                result['success'] = True
            else:
                result['error'] = "Word 轉換失敗"
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"處理檔案失敗: {e}")
        
        return result


def main():
    """主函數"""
    if len(sys.argv) < 2:
        print("使用方法: python process_transcription.py <transcription_file.txt>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    print(f"\n=== 轉錄文字處理程式 ===")
    print(f"輸入檔案: {input_file}")
    print(f"使用模型: Gemini-2.5-pro")
    print("處理中...\n")
    
    try:
        processor = TranscriptionProcessor()
        result = processor.process_file(input_file)
        
        if result['success']:
            print("\n✅ 處理成功！")
            print(f"Markdown 檔案: {result['markdown_file']}")
            print(f"Word 檔案: {result['docx_file']}")
        else:
            print(f"\n❌ 處理失敗: {result['error']}")
            sys.exit(1)
            
    except Exception as e:
        print(f"\n❌ 程式執行失敗: {e}")
        logger.error(f"程式執行失敗: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()