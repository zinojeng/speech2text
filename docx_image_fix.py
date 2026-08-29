#!/usr/bin/env python3
"""
DOCX Image Fix Script
Fixes the issue where images are shown as text paths instead of actual images in DOCX files.

This script processes DOCX files that contain image references in the format:
[IMAGE: folder/filename.jpg]

And replaces them with actual embedded images.
"""

import os
import sys
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def find_and_insert_images(docx_path: str, base_image_dir: str = None) -> bool:
    """
    Process a DOCX file and replace image text references with actual images.
    
    Args:
        docx_path: Path to the DOCX file
        base_image_dir: Base directory where images are located (defaults to docx directory)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Load the document
        doc = Document(docx_path)
        docx_dir = Path(docx_path).parent
        
        if base_image_dir is None:
            base_image_dir = docx_dir
        else:
            base_image_dir = Path(base_image_dir)
        
        # Track changes
        images_found = 0
        images_inserted = 0
        
        # Create a new document to rebuild
        new_doc = Document()
        
        # Copy styles from original document
        for style in doc.styles:
            try:
                if hasattr(new_doc.styles, 'add_style'):
                    if style.name not in [s.name for s in new_doc.styles]:
                        new_doc.styles.add_style(style.name, style.type)
            except:
                pass
        
        # Process each paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check if this paragraph contains an image reference
            if '[IMAGE:' in text:
                # Find all image references
                pattern = r'\[IMAGE:\s*([^\]]+)\]'
                matches = re.findall(pattern, text)
                
                if matches:
                    # This paragraph contains image references
                    for img_ref in matches:
                        images_found += 1
                        
                        # Clean up the image reference
                        img_ref = img_ref.strip()
                        
                        # Try different paths
                        possible_paths = [
                            base_image_dir / img_ref,  # Full path from base
                            docx_dir / img_ref,        # Relative to docx
                            Path(img_ref),             # Absolute path
                        ]
                        
                        # If the reference doesn't include extension, try common ones
                        if not any(img_ref.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.JPG', '.PNG']):
                            for base_path in possible_paths[:]:
                                for ext in ['.jpg', '.jpeg', '.png', '.JPG', '.PNG']:
                                    possible_paths.append(Path(str(base_path) + ext))
                        
                        # Try to find the image
                        image_found = False
                        for img_path in possible_paths:
                            if img_path.exists():
                                try:
                                    # Add the image
                                    new_doc.add_picture(str(img_path), width=Inches(5.5))
                                    images_inserted += 1
                                    logger.info(f"Inserted image: {img_path}")
                                    image_found = True
                                    break
                                except Exception as e:
                                    logger.warning(f"Failed to insert image {img_path}: {e}")
                        
                        if not image_found:
                            logger.warning(f"Image not found: {img_ref}")
                            # Add text indicating missing image
                            new_para = new_doc.add_paragraph()
                            new_para.add_run(f"[Missing image: {img_ref}]").italic = True
                    
                    # If there was text before or after the image reference, add it
                    clean_text = re.sub(pattern, '', text).strip()
                    if clean_text:
                        new_para = new_doc.add_paragraph()
                        copy_paragraph_format(para, new_para)
                        new_para.add_run(clean_text)
                else:
                    # No image reference, copy paragraph as is
                    new_para = new_doc.add_paragraph()
                    copy_paragraph_format(para, new_para)
                    copy_paragraph_runs(para, new_para)
            else:
                # No image reference, copy paragraph as is
                new_para = new_doc.add_paragraph()
                copy_paragraph_format(para, new_para)
                copy_paragraph_runs(para, new_para)
        
        # Save the new document
        output_path = Path(docx_path).with_suffix('.fixed.docx')
        new_doc.save(str(output_path))
        
        logger.info(f"Processing complete!")
        logger.info(f"Images found: {images_found}")
        logger.info(f"Images inserted: {images_inserted}")
        logger.info(f"Output saved to: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return False


def copy_paragraph_format(source_para, target_para):
    """Copy formatting from source paragraph to target paragraph."""
    try:
        target_para.alignment = source_para.alignment
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    except:
        pass


def copy_paragraph_runs(source_para, target_para):
    """Copy all runs from source paragraph to target paragraph."""
    for run in source_para.runs:
        new_run = target_para.add_run(run.text)
        try:
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
        except:
            pass


def main():
    """Main function for command line usage."""
    if len(sys.argv) < 2:
        print("Usage: python docx_image_fix.py <docx_file> [base_image_directory]")
        print("\nExample:")
        print("  python docx_image_fix.py document.docx")
        print("  python docx_image_fix.py document.docx /path/to/images")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    base_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(docx_file):
        print(f"Error: File not found: {docx_file}")
        sys.exit(1)
    
    print(f"Processing: {docx_file}")
    if base_dir:
        print(f"Using image directory: {base_dir}")
    
    if find_and_insert_images(docx_file, base_dir):
        print("Success! Check the .fixed.docx file.")
    else:
        print("Failed to process the document.")
        sys.exit(1)


if __name__ == "__main__":
    main()