#!/usr/bin/env python3
"""
智慧批次音訊處理程式
Smart Batch Audio Processor with Skip Functionality

此程式用於自動化處理音訊檔案，包含：
1. 遞歸搜索資料夾和子資料夾中的音訊檔案
2. 智能跳過已處理的檔案（檢查同名轉錄檔案）
3. 使用 GPT-4o 進行語音轉錄
4. 支援多種輸出格式 (text, srt, markdown)
5. 可選的 Word 文件輸出

使用方法:
    python batch_audio_smart.py <資料夾路徑> [選項]
    
選項:
    --model MODEL        轉錄模型 (預設: gpt-transcribe)
    --format FORMAT      輸出格式: text, srt, markdown (預設: text)
    --force             強制重新處理所有檔案（忽略已存在的轉錄）
    --docx              同時產生 Word 文件
    --combined          合併所有轉錄到單一檔案
"""

import os
import sys
import logging
import argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import time
import json

# 延遲載入大型模組
def lazy_imports():
    global OpenAI, split_large_audio, check_file_size, AudioSegment, load_dotenv
    from dotenv import load_dotenv
    from openai import OpenAI
    from utils import split_large_audio, check_file_size
    from pydub import AudioSegment
    return load_dotenv

# 初始化時載入
OpenAI = None
split_large_audio = None
check_file_size = None
AudioSegment = None
load_dotenv = None

# 設定日誌
def setup_logging(log_file: str = None):
    """設定日誌系統"""
    if log_file is None:
        log_file = f'batch_smart_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger(__name__)

# 支援的音訊格式
SUPPORTED_AUDIO_FORMATS = [
    '.mp3', '.wav', '.m4a', '.aac', '.flac', '.ogg',
    '.wma', '.mp4', '.mov', '.avi', '.mkv', '.webm'
]

# 支援的轉錄輸出格式
TRANSCRIPT_FORMATS = ['.txt', '.srt', '.md']

class SmartBatchProcessor:
    """智慧批次音訊處理器"""
    
    def __init__(self, model: str = "gpt-transcribe", 
                 output_format: str = "text",
                 force_reprocess: bool = False,
                 generate_docx: bool = False,
                 combined_output: bool = False):
        """
        初始化處理器
        
        Args:
            model: 使用的轉錄模型
            output_format: 輸出格式 (text, srt, markdown)
            force_reprocess: 是否強制重新處理已存在的檔案
            generate_docx: 是否生成 Word 文件
            combined_output: 是否合併所有輸出到單一檔案
        """
        # 延遲載入模組
        global load_dotenv
        load_dotenv_func = lazy_imports()
        load_dotenv_func()
        
        self.logger = setup_logging()
        self.openai_client = None
        self.transcribe_model = model
        self.output_format = output_format
        self.force_reprocess = force_reprocess
        self.generate_docx = generate_docx
        self.combined_output = combined_output
        self.setup_apis()
        
        # 統計資訊
        self.stats = {
            'total_files': 0,
            'processed': 0,
            'skipped': 0,
            'failed': 0,
            'processing_time': 0
        }
    
    def setup_apis(self):
        """設定 API 客戶端"""
        openai_key = os.getenv('OPENAI_API_KEY')
        if not openai_key:
            raise ValueError("請在 .env 檔案中設定 OPENAI_API_KEY")
        
        self.openai_client = OpenAI(api_key=openai_key)
        self.logger.info(f"OpenAI API 設定完成，使用模型: {self.transcribe_model}")
    
    def get_output_extension(self) -> str:
        """取得輸出檔案副檔名"""
        format_map = {
            'text': '.txt',
            'srt': '.srt',
            'markdown': '.md',
            'md': '.md'
        }
        return format_map.get(self.output_format.lower(), '.txt')
    
    def check_existing_transcript(self, audio_path: str) -> Optional[str]:
        """
        檢查是否已存在轉錄檔案
        
        Args:
            audio_path: 音訊檔案路徑
            
        Returns:
            已存在的轉錄檔案路徑，或 None
        """
        audio_path_obj = Path(audio_path)
        base_path = audio_path_obj.with_suffix('')
        
        # 檢查各種可能的轉錄檔案格式
        for ext in TRANSCRIPT_FORMATS:
            transcript_path = base_path.with_suffix(ext)
            if transcript_path.exists() and transcript_path.stat().st_size > 0:
                return str(transcript_path)
        
        # 檢查 Word 檔案
        docx_path = base_path.with_suffix('.docx')
        if docx_path.exists() and docx_path.stat().st_size > 0:
            return str(docx_path)
        
        return None
    
    def scan_audio_files(self, folder_path: str) -> List[str]:
        """
        遞歸掃描資料夾中的所有音訊檔案
        
        Args:
            folder_path: 資料夾路徑
            
        Returns:
            音訊檔案路徑列表
        """
        folder_path = Path(folder_path)
        if not folder_path.exists():
            self.logger.error(f"資料夾不存在: {folder_path}")
            return []
        
        audio_files = []
        
        # 遞歸搜索所有音訊檔案
        for ext in SUPPORTED_AUDIO_FORMATS:
            pattern = f"**/*{ext}"
            files = list(folder_path.glob(pattern))
            audio_files.extend([str(f) for f in files])
        
        # 去重並排序
        audio_files = sorted(list(set(audio_files)))
        
        self.logger.info(f"在 {folder_path} 中找到 {len(audio_files)} 個音訊檔案")
        
        return audio_files
    
    def filter_files_to_process(self, audio_files: List[str]) -> Tuple[List[str], List[str]]:
        """
        過濾需要處理的檔案
        
        Args:
            audio_files: 所有音訊檔案列表
            
        Returns:
            (需要處理的檔案列表, 跳過的檔案列表)
        """
        to_process = []
        to_skip = []
        
        for audio_file in audio_files:
            if self.force_reprocess:
                to_process.append(audio_file)
            else:
                existing = self.check_existing_transcript(audio_file)
                if existing:
                    to_skip.append((audio_file, existing))
                    self.logger.info(f"跳過已處理: {Path(audio_file).name} -> {Path(existing).name}")
                else:
                    to_process.append(audio_file)
        
        return to_process, to_skip
    
    def transcribe_audio(self, audio_path: str) -> Optional[str]:
        """
        使用 GPT-4o 轉錄音訊檔案
        
        Args:
            audio_path: 音訊檔案路徑
            
        Returns:
            轉錄文字或 None
        """
        try:
            self.logger.info(f"開始轉錄: {Path(audio_path).name}")
            
            # 檢查檔案大小
            if check_file_size(audio_path):
                self.logger.info(f"檔案較大，進行分割處理")
                segments = split_large_audio(audio_path)
                
                if not segments:
                    self.logger.error("音訊分割失敗")
                    return None
                
                # 轉錄每個片段
                full_transcript = ""
                for i, segment_path in enumerate(segments):
                    self.logger.info(f"轉錄片段 {i+1}/{len(segments)}")
                    
                    try:
                        with open(segment_path, "rb") as audio_file:
                            response = self.openai_client.audio.transcriptions.create(
                                model=self.transcribe_model.replace('-transcribe', ''),
                                file=audio_file,
                                response_format=self.output_format if self.output_format in ['text', 'srt'] else 'text'
                            )
                        
                        if self.output_format == 'srt':
                            full_transcript += response + "\n\n"
                        else:
                            full_transcript += response + " "
                        
                        # 清理臨時檔案
                        if os.path.exists(segment_path):
                            os.remove(segment_path)
                    
                    except Exception as e:
                        self.logger.error(f"片段 {i+1} 轉錄失敗: {e}")
                        if os.path.exists(segment_path):
                            os.remove(segment_path)
                        continue
                    
                    # 避免 API 限制
                    time.sleep(2)
                
                return full_transcript.strip() if full_transcript else None
            
            else:
                # 直接轉錄
                with open(audio_path, "rb") as audio_file:
                    response = self.openai_client.audio.transcriptions.create(
                        model=self.transcribe_model.replace('-transcribe', ''),
                        file=audio_file,
                        response_format=self.output_format if self.output_format in ['text', 'srt'] else 'text'
                    )
                
                return response
        
        except Exception as e:
            self.logger.error(f"轉錄失敗 {audio_path}: {e}")
            return None
    
    def save_transcript(self, transcript: str, audio_path: str) -> bool:
        """
        儲存轉錄結果
        
        Args:
            transcript: 轉錄文字
            audio_path: 原始音訊檔案路徑
            
        Returns:
            是否成功儲存
        """
        try:
            audio_path_obj = Path(audio_path)
            output_path = audio_path_obj.with_suffix(self.get_output_extension())
            
            # 如果是 markdown 格式，加入標題
            if self.output_format == 'markdown':
                transcript = f"# {audio_path_obj.stem}\n\n{transcript}"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(transcript)
            
            self.logger.info(f"轉錄已儲存: {output_path.name}")
            
            # 如果需要生成 Word 文件
            if self.generate_docx:
                self.create_docx(transcript, str(audio_path_obj.with_suffix('.docx')))
            
            return True
        
        except Exception as e:
            self.logger.error(f"儲存失敗: {e}")
            return False
    
    def create_docx(self, content: str, output_path: str):
        """
        建立 Word 文件
        
        Args:
            content: 文件內容
            output_path: 輸出路徑
        """
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # 加入標題
            title = Path(output_path).stem
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 加入內容
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.save(output_path)
            self.logger.info(f"Word 文件已建立: {Path(output_path).name}")
        
        except Exception as e:
            self.logger.error(f"建立 Word 文件失敗: {e}")
    
    def save_combined_output(self, transcripts: Dict[str, str], output_folder: str):
        """
        儲存合併的輸出檔案
        
        Args:
            transcripts: {檔案路徑: 轉錄內容} 字典
            output_folder: 輸出資料夾
        """
        if not transcripts:
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = Path(output_folder) / f"combined_transcript_{timestamp}{self.get_output_extension()}"
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                if self.output_format == 'srt':
                    # SRT 格式需要重新編號
                    srt_index = 1
                    for audio_file, transcript in transcripts.items():
                        f.write(f"# {Path(audio_file).name}\n\n")
                        
                        # 重新編號 SRT 條目
                        lines = transcript.split('\n')
                        for line in lines:
                            if line.strip().isdigit():
                                f.write(f"{srt_index}\n")
                                srt_index += 1
                            else:
                                f.write(f"{line}\n")
                        f.write("\n\n")
                
                else:
                    # Text 或 Markdown 格式
                    for audio_file, transcript in transcripts.items():
                        f.write(f"{'='*60}\n")
                        f.write(f"檔案: {Path(audio_file).name}\n")
                        f.write(f"{'='*60}\n\n")
                        f.write(transcript)
                        f.write("\n\n")
            
            self.logger.info(f"合併檔案已儲存: {output_path}")
            
            # 如果需要生成 Word 文件
            if self.generate_docx:
                self.create_combined_docx(transcripts, str(output_path.with_suffix('.docx')))
        
        except Exception as e:
            self.logger.error(f"儲存合併檔案失敗: {e}")
    
    def create_combined_docx(self, transcripts: Dict[str, str], output_path: str):
        """建立合併的 Word 文件"""
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # 加入總標題
            title = doc.add_heading("批次轉錄結果", level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 加入日期
            doc.add_paragraph(f"生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")
            
            # 加入各個轉錄內容
            for audio_file, transcript in transcripts.items():
                # 檔案標題
                doc.add_heading(Path(audio_file).name, level=2)
                
                # 轉錄內容
                paragraphs = transcript.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # 分頁
                doc.add_page_break()
            
            doc.save(output_path)
            self.logger.info(f"合併 Word 文件已建立: {Path(output_path).name}")
        
        except Exception as e:
            self.logger.error(f"建立合併 Word 文件失敗: {e}")
    
    def process_folder(self, folder_path: str) -> Dict:
        """
        處理整個資料夾
        
        Args:
            folder_path: 資料夾路徑
            
        Returns:
            處理統計結果
        """
        start_time = time.time()
        self.logger.info(f"開始處理資料夾: {folder_path}")
        
        # 掃描音訊檔案
        audio_files = self.scan_audio_files(folder_path)
        self.stats['total_files'] = len(audio_files)
        
        if not audio_files:
            self.logger.warning("未找到任何音訊檔案")
            return self.stats
        
        # 過濾需要處理的檔案
        to_process, to_skip = self.filter_files_to_process(audio_files)
        self.stats['skipped'] = len(to_skip)
        
        # 顯示統計
        self.logger.info(f"總共 {len(audio_files)} 個檔案:")
        self.logger.info(f"  - 需要處理: {len(to_process)}")
        self.logger.info(f"  - 跳過已存在: {len(to_skip)}")
        
        if not to_process:
            self.logger.info("沒有需要處理的檔案")
            return self.stats
        
        # 處理檔案
        transcripts = {} if self.combined_output else None
        
        for i, audio_file in enumerate(to_process, 1):
            self.logger.info(f"\n處理進度: {i}/{len(to_process)}")
            self.logger.info(f"檔案: {Path(audio_file).name}")
            
            # 轉錄
            transcript = self.transcribe_audio(audio_file)
            
            if transcript:
                if self.combined_output:
                    transcripts[audio_file] = transcript
                else:
                    # 個別儲存
                    if self.save_transcript(transcript, audio_file):
                        self.stats['processed'] += 1
                    else:
                        self.stats['failed'] += 1
                
                # 如果不是合併模式，直接計數
                if not self.combined_output:
                    continue
                    
                self.stats['processed'] += 1
            else:
                self.stats['failed'] += 1
                self.logger.error(f"轉錄失敗: {Path(audio_file).name}")
            
            # 進度顯示
            if i % 5 == 0:
                elapsed = time.time() - start_time
                avg_time = elapsed / i
                remaining = avg_time * (len(to_process) - i)
                self.logger.info(f"預計剩餘時間: {remaining/60:.1f} 分鐘")
        
        # 儲存合併檔案
        if self.combined_output and transcripts:
            self.save_combined_output(transcripts, folder_path)
        
        # 統計
        self.stats['processing_time'] = time.time() - start_time
        
        # 儲存處理報告
        self.save_report(folder_path)
        
        return self.stats
    
    def save_report(self, folder_path: str):
        """儲存處理報告"""
        report_path = Path(folder_path) / f"processing_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        report = {
            'timestamp': datetime.now().isoformat(),
            'folder': str(folder_path),
            'model': self.transcribe_model,
            'output_format': self.output_format,
            'statistics': self.stats,
            'processing_time_minutes': self.stats['processing_time'] / 60
        }
        
        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, ensure_ascii=False, indent=2)
            self.logger.info(f"處理報告已儲存: {report_path.name}")
        except Exception as e:
            self.logger.error(f"儲存報告失敗: {e}")
    
    def print_summary(self):
        """列印處理摘要"""
        print("\n" + "="*60)
        print("處理摘要")
        print("="*60)
        print(f"總檔案數: {self.stats['total_files']}")
        print(f"已處理: {self.stats['processed']}")
        print(f"已跳過: {self.stats['skipped']}")
        print(f"失敗: {self.stats['failed']}")
        print(f"處理時間: {self.stats['processing_time']/60:.1f} 分鐘")
        
        if self.stats['processed'] > 0:
            avg_time = self.stats['processing_time'] / self.stats['processed']
            print(f"平均處理時間: {avg_time:.1f} 秒/檔案")
        
        print("="*60)


def main():
    """主程式"""
    parser = argparse.ArgumentParser(
        description='智慧批次音訊轉錄處理程式',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument('folder', help='要處理的資料夾路徑')
    parser.add_argument('--model', default='gpt-transcribe',
                        choices=['gpt-transcribe', 'gemini-3.5-transcribe'],
                        help='轉錄模型 (預設: gpt-transcribe)')
    parser.add_argument('--format', default='text',
                        choices=['text', 'srt', 'markdown'],
                        help='輸出格式 (預設: text)')
    parser.add_argument('--force', action='store_true',
                        help='強制重新處理所有檔案')
    parser.add_argument('--docx', action='store_true',
                        help='同時產生 Word 文件')
    parser.add_argument('--combined', action='store_true',
                        help='合併所有轉錄到單一檔案')
    
    args = parser.parse_args()
    
    # 建立處理器
    processor = SmartBatchProcessor(
        model=args.model,
        output_format=args.format,
        force_reprocess=args.force,
        generate_docx=args.docx,
        combined_output=args.combined
    )
    
    # 處理資料夾
    try:
        processor.process_folder(args.folder)
        processor.print_summary()
    except KeyboardInterrupt:
        print("\n處理已中斷")
        processor.print_summary()
        sys.exit(1)
    except Exception as e:
        print(f"\n錯誤: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()