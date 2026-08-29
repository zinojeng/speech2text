#!/usr/bin/env python3
"""
Enhanced script to fix DOCX files with timestamp-based image references
Specifically handles the format: > 🖼️ **投影片圖表說明**（[00:04:28]）：
"""

import os
import sys
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
from typing import Dict, List, Optional, Tuple

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ImageTimestampFixer:
    def __init__(self, docx_path: str, slides_with_images: List[Tuple[str, str]]):
        """
        Args:
            docx_path: Path to the DOCX file to fix
            slides_with_images: List of (slide_file, image_folder) tuples
        """
        self.docx_path = Path(docx_path)
        self.slides_with_images = slides_with_images
        self.all_slide_images = {}  # {time_sec: [(slide_index, img_path), ...]}
        
    def parse_time_format(self, time_str: str) -> Optional[float]:
        """Parse time formats including HH:MM:SS"""
        # Remove brackets if present
        time_str = time_str.strip('[]')
        
        # HH:MM:SS format
        match = re.match(r'(\d{1,2}):(\d{2}):(\d{2})', time_str)
        if match:
            hours = int(match.group(1))
            minutes = int(match.group(2))
            seconds = int(match.group(3))
            return hours * 3600 + minutes * 60 + seconds
        
        # M:SS format
        match = re.match(r'(\d+):(\d+)', time_str)
        if match:
            minutes = int(match.group(1))
            seconds = int(match.group(2))
            return minutes * 60 + seconds
        
        return None
    
    def parse_slide_time(self, filename: str) -> Optional[float]:
        """Parse time from slide image filename"""
        # Format: slide_009_t1m4.7s.jpg
        match = re.search(r't(\d+)m([\d.]+)s', filename)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        return None
    
    def load_all_images(self):
        """Load all images from specified folders"""
        for slide_index, (slide_file, img_folder) in enumerate(self.slides_with_images):
            img_path = Path(img_folder)
            if not img_path.exists():
                logger.warning(f"Image folder not found: {img_folder}")
                continue
                
            # Load images from this folder
            for img_file in img_path.glob("*.jpg"):
                time_sec = self.parse_slide_time(img_file.name)
                if time_sec is not None:
                    if time_sec not in self.all_slide_images:
                        self.all_slide_images[time_sec] = []
                    self.all_slide_images[time_sec].append((slide_index, str(img_file)))
            
            for img_file in img_path.glob("*.png"):
                time_sec = self.parse_slide_time(img_file.name)
                if time_sec is not None:
                    if time_sec not in self.all_slide_images:
                        self.all_slide_images[time_sec] = []
                    self.all_slide_images[time_sec].append((slide_index, str(img_file)))
        
        logger.info(f"Loaded {len(self.all_slide_images)} unique timestamps with images")
    
    def find_closest_image(self, target_time: float) -> Optional[str]:
        """Find the closest image to the target time"""
        if not self.all_slide_images:
            return None
            
        # Find closest time within 30 seconds
        closest_time = min(self.all_slide_images.keys(), 
                          key=lambda x: abs(x - target_time))
        
        if abs(closest_time - target_time) <= 30:
            # Get the first image (prioritize earlier slides)
            images = sorted(self.all_slide_images[closest_time], key=lambda x: x[0])
            return images[0][1]
        
        return None
    
    def fix_document(self, output_path: str = None):
        """Fix the document by replacing timestamp references with images"""
        if output_path is None:
            output_path = self.docx_path.with_suffix('.fixed.docx')
        
        # Load all images first
        self.load_all_images()
        
        # Load the document
        doc = Document(str(self.docx_path))
        new_doc = Document()
        
        # Process each paragraph
        for para in doc.paragraphs:
            text = para.text
            
            # Check for timestamp patterns
            if '🖼️' in text and '（[' in text:
                # Extract all timestamps from this line
                timestamps = re.findall(r'\[(\d{2}:\d{2}:\d{2})\]', text)
                
                if timestamps:
                    # Add the text first (optional - you might want to skip this)
                    # new_para = new_doc.add_paragraph(text)
                    
                    # Insert images for each timestamp
                    images_inserted = False
                    for ts in timestamps:
                        time_sec = self.parse_time_format(ts)
                        if time_sec is not None:
                            img_path = self.find_closest_image(time_sec)
                            if img_path and os.path.exists(img_path):
                                try:
                                    new_doc.add_picture(img_path, width=Inches(5.5))
                                    new_doc.add_paragraph()  # Add space after image
                                    logger.info(f"Inserted image for timestamp {ts}: {os.path.basename(img_path)}")
                                    images_inserted = True
                                except Exception as e:
                                    logger.error(f"Failed to insert image {img_path}: {e}")
                    
                    # If no images were inserted, add the original text
                    if not images_inserted:
                        new_para = new_doc.add_paragraph(text)
                        # Copy formatting
                        self.copy_paragraph_format(para, new_para)
                else:
                    # No timestamps found, copy as is
                    new_para = new_doc.add_paragraph()
                    self.copy_paragraph_format(para, new_para)
                    self.copy_paragraph_runs(para, new_para)
            else:
                # No image markers, copy as is
                new_para = new_doc.add_paragraph()
                self.copy_paragraph_format(para, new_para)
                self.copy_paragraph_runs(para, new_para)
        
        # Save the fixed document
        new_doc.save(str(output_path))
        logger.info(f"Fixed document saved to: {output_path}")
        
    def copy_paragraph_format(self, source_para, target_para):
        """Copy paragraph formatting"""
        try:
            target_para.alignment = source_para.alignment
            pf = target_para.paragraph_format
            spf = source_para.paragraph_format
            pf.space_before = spf.space_before
            pf.space_after = spf.space_after
            pf.line_spacing = spf.line_spacing
            pf.left_indent = spf.left_indent
            pf.right_indent = spf.right_indent
            pf.first_line_indent = spf.first_line_indent
        except:
            pass
    
    def copy_paragraph_runs(self, source_para, target_para):
        """Copy all runs with formatting"""
        for run in source_para.runs:
            new_run = target_para.add_run(run.text)
            try:
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
            except:
                pass


def main():
    """Main function for command line usage"""
    if len(sys.argv) < 2:
        print("Usage: python fix_image_timestamps.py <docx_file> <slide1.md:images1/> [slide2.md:images2/] ...")
        print("\nExample:")
        print("  python fix_image_timestamps.py document.docx slide1.md:images1/ slide2.md:images2/")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    slides_with_images = []
    
    # Parse slide:image pairs
    for arg in sys.argv[2:]:
        if ':' in arg:
            slide, images = arg.split(':', 1)
            slides_with_images.append((slide, images))
    
    if not slides_with_images:
        print("Error: No slide:image pairs provided")
        sys.exit(1)
    
    print(f"Processing: {docx_file}")
    print(f"With {len(slides_with_images)} slide/image folders")
    
    # Create fixer and process
    fixer = ImageTimestampFixer(docx_file, slides_with_images)
    fixer.fix_document()
    
    print("Done! Check the .fixed.docx file")


if __name__ == "__main__":
    main()