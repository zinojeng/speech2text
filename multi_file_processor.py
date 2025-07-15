import streamlit as st
import os
import tempfile
import shutil
from pathlib import Path
import pandas as pd
from datetime import datetime
import zipfile
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

import os
from dotenv import load_dotenv

load_dotenv()

from audio2text.gpt4o_stt import transcribe_audio_gpt4o
from gemini_utils import call_gemini_api
from utils import check_file_constraints, split_large_audio

# Optional imports
try:
    from docx import Document
    from markdown2 import markdown
    from bs4 import BeautifulSoup
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx, markdown2, or beautifulsoup4 not installed. Word export disabled.")

st.set_page_config(
    page_title="Multi-File Audio Processor",
    page_icon="🎵",
    layout="wide"
)

st.title("🎵 Multi-File Audio Processor")
st.markdown("### Batch process audio files with GPT-4o transcription and Gemini 2.5 Pro summarization")
st.markdown("---")

# Check API keys
if not os.getenv("OPENAI_API_KEY"):
    st.error("⚠️ OPENAI_API_KEY not found in environment. Please set it in .env file.")
    st.stop()

if not os.getenv("GOOGLE_API_KEY"):
    st.warning("⚠️ GOOGLE_API_KEY not found. Summary generation will be disabled.")
    has_gemini = False
else:
    has_gemini = True

if 'processed_files' not in st.session_state:
    st.session_state.processed_files = []
if 'processing_status' not in st.session_state:
    st.session_state.processing_status = {}
if 'results' not in st.session_state:
    st.session_state.results = {}
if 'batch_id' not in st.session_state:
    st.session_state.batch_id = None

def find_audio_files(directory):
    """Recursively find all audio files in directory and subdirectories"""
    audio_extensions = {'.mp3', '.wav', '.m4a', '.flac', '.aac', '.ogg', '.wma', '.opus', '.webm'}
    audio_files = []
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            if Path(file).suffix.lower() in audio_extensions:
                audio_files.append(os.path.join(root, file))
    
    return sorted(audio_files)

def process_single_file(file_path, model="gpt-4o-mini-transcribe", summary_prompt=None, max_retries=3, has_gemini=True):
    """Process a single audio file through transcription and summarization with retry logic"""
    result = {
        'file_path': file_path,
        'file_name': os.path.basename(file_path),
        'status': 'processing',
        'transcription': None,
        'summary': None,
        'error': None,
        'processing_time': 0,
        'retries': 0
    }
    
    start_time = time.time()
    
    for retry in range(max_retries):
        try:
            # Check file constraints
            constraint_ok, constraint_msg = check_file_constraints(file_path)
            
            # Get file size for splitting decision
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
            size_ok = file_size <= 25  # 25MB limit
            duration_ok = True  # We'll let split_large_audio handle duration
            
            if not size_ok or not duration_ok:
                # Split large files
                with tempfile.TemporaryDirectory() as temp_dir:
                    segments = split_large_audio(file_path, temp_dir)
                    transcriptions = []
                    
                    for i, segment in enumerate(segments):
                        st.write(f"Processing segment {i+1}/{len(segments)} of {os.path.basename(file_path)}")
                        segment_transcription = transcribe_audio_gpt4o(
                            segment, 
                            api_key=os.getenv("OPENAI_API_KEY"),
                            model=model
                        )
                        if segment_transcription:
                            transcriptions.append(segment_transcription)
                    
                    result['transcription'] = "\n\n".join(transcriptions)
            else:
                # Process normally
                result['transcription'] = transcribe_audio_gpt4o(
                    file_path,
                    api_key=os.getenv("OPENAI_API_KEY"),
                    model=model
                )
        
            if result['transcription']:
                # Generate summary if prompt provided
                if summary_prompt and has_gemini:
                    summary_text = f"{summary_prompt}\n\nTranscription:\n{result['transcription']}"
                    result['summary'] = call_gemini_api(summary_text, model="gemini-2.0-flash-exp")
                
                result['status'] = 'completed'
                break  # Success, exit retry loop
            else:
                result['status'] = 'failed'
                result['error'] = 'Transcription failed'
            
        except Exception as e:
            result['error'] = str(e)
            result['retries'] = retry + 1
            
            if retry < max_retries - 1:
                # Wait before retry (exponential backoff)
                wait_time = 2 ** retry
                time.sleep(wait_time)
            else:
                result['status'] = 'failed'
    
    result['processing_time'] = time.time() - start_time
    return result

def markdown_to_docx(markdown_text, title="Document"):
    """Convert markdown text to Word document"""
    if not HAS_DOCX:
        return None
        
    doc = Document()
    doc.add_heading(title, 0)
    
    # Convert markdown to HTML
    html = markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')
    
    # Process HTML elements
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'pre']):
        if element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name in ['h1', 'h2', 'h3']:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
        elif element.name == 'pre':
            doc.add_paragraph(element.get_text(), style='Normal')
    
    return doc

def save_batch_state(batch_id, audio_files, results, status):
    """Save batch processing state to file"""
    state_dir = os.path.join(tempfile.gettempdir(), "batch_processor_states")
    os.makedirs(state_dir, exist_ok=True)
    
    state_file = os.path.join(state_dir, f"batch_{batch_id}.json")
    state_data = {
        'batch_id': batch_id,
        'timestamp': datetime.now().isoformat(),
        'audio_files': audio_files,
        'results': results,
        'status': status
    }
    
    with open(state_file, 'w', encoding='utf-8') as f:
        json.dump(state_data, f, ensure_ascii=False, indent=2)
    
    return state_file

def load_batch_state(batch_id):
    """Load batch processing state from file"""
    state_dir = os.path.join(tempfile.gettempdir(), "batch_processor_states")
    state_file = os.path.join(state_dir, f"batch_{batch_id}.json")
    
    if os.path.exists(state_file):
        with open(state_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def list_saved_batches():
    """List all saved batch states"""
    state_dir = os.path.join(tempfile.gettempdir(), "batch_processor_states")
    if not os.path.exists(state_dir):
        return []
    
    batches = []
    for file in os.listdir(state_dir):
        if file.startswith("batch_") and file.endswith(".json"):
            batch_id = file[6:-5]  # Remove "batch_" and ".json"
            state = load_batch_state(batch_id)
            if state:
                batches.append({
                    'id': batch_id,
                    'timestamp': state['timestamp'],
                    'file_count': len(state['audio_files']),
                    'completed': sum(1 for r in state['results'].values() if r['status'] == 'completed')
                })
    
    return sorted(batches, key=lambda x: x['timestamp'], reverse=True)

# Main UI
col1, col2 = st.columns([1, 2])

with col1:
    st.header("📁 Input Settings")
    
    # Recovery option
    saved_batches = list_saved_batches()
    if saved_batches:
        st.subheader("📋 Resume Previous Batch")
        batch_options = ["New batch"] + [f"Batch {b['id']} ({b['timestamp'][:16]}, {b['completed']}/{b['file_count']} done)" for b in saved_batches]
        selected_batch = st.selectbox("Select batch:", batch_options)
        
        if selected_batch != "New batch":
            batch_id = saved_batches[batch_options.index(selected_batch) - 1]['id']
            if st.button("📥 Load Batch"):
                state = load_batch_state(batch_id)
                if state:
                    st.session_state.batch_id = batch_id
                    st.session_state.processing_status = state['status']
                    st.session_state.results = state['results']
                    st.success(f"Loaded batch {batch_id}")
                    st.rerun()
    
    # File upload options
    upload_method = st.radio(
        "Select input method:",
        ["Upload multiple files", "Select folder path"]
    )
    
    audio_files = []
    
    if upload_method == "Upload multiple files":
        uploaded_files = st.file_uploader(
            "Choose audio files",
            type=['mp3', 'wav', 'm4a', 'flac', 'aac', 'ogg', 'wma', 'opus', 'webm'],
            accept_multiple_files=True
        )
        
        if uploaded_files:
            # Save uploaded files temporarily
            temp_dir = tempfile.mkdtemp()
            for uploaded_file in uploaded_files:
                file_path = os.path.join(temp_dir, uploaded_file.name)
                with open(file_path, 'wb') as f:
                    f.write(uploaded_file.getbuffer())
                audio_files.append(file_path)
    
    else:
        folder_path = st.text_input("Enter folder path:", placeholder="/path/to/audio/folder")
        if folder_path and os.path.exists(folder_path):
            audio_files = find_audio_files(folder_path)
            st.success(f"Found {len(audio_files)} audio files")
        elif folder_path:
            st.error("Folder path does not exist")
    
    # Processing settings
    st.subheader("⚙️ Processing Settings")
    
    model = st.selectbox(
        "Select transcription model:",
        ["gpt-4o-mini-transcribe", "gpt-4o-transcribe"],
        index=0
    )
    
    use_summary = st.checkbox("Generate summaries with Gemini 2.5 Pro", value=True, disabled=not has_gemini)
    
    summary_prompt = ""
    if use_summary:
        summary_prompt = st.text_area(
            "Summary prompt:",
            value="Please provide a concise summary of the following transcription, highlighting key points and main topics discussed:",
            height=100
        )
    
    # Batch processing controls
    st.markdown("---")
    st.subheader("🚀 Batch Processing")
    
    if audio_files:
        # File summary info box
        st.info(f"""
        📁 **Files found:** {len(audio_files)}  
        🎵 **Supported formats:** MP3, WAV, M4A, FLAC, AAC, OGG, WMA, OPUS, WebM
        """)
        
        # Show file list
        with st.expander("📋 View file list", expanded=False):
            for i, file in enumerate(audio_files, 1):
                st.text(f"{i}. {os.path.basename(file)}")
        
        max_workers = st.slider("Parallel processing threads:", 1, 5, 2)
        
        if st.button("🎯 Start Batch Processing", type="primary"):
            # Generate batch ID
            st.session_state.batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.session_state.processing_status = {file: 'pending' for file in audio_files}
            st.session_state.results = {}
            
            # Save initial state
            save_batch_state(
                st.session_state.batch_id,
                audio_files,
                st.session_state.results,
                st.session_state.processing_status
            )
            
            # Process files
            progress_bar = st.progress(0)
            status_text = st.empty()
            processing_container = st.container()
            
            with processing_container:
                st.write("🔄 Processing Status:")
                status_cols = st.columns(4)
                with status_cols[0]:
                    pending_metric = st.metric("Pending", len(audio_files))
                with status_cols[1]:
                    processing_metric = st.metric("Processing", 0)
                with status_cols[2]:
                    completed_metric = st.metric("Completed", 0)
                with status_cols[3]:
                    failed_metric = st.metric("Failed", 0)
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {
                    executor.submit(process_single_file, file, model, summary_prompt, has_gemini=has_gemini): file 
                    for file in audio_files
                }
                
                completed = 0
                for future in as_completed(future_to_file):
                    file = future_to_file[future]
                    try:
                        result = future.result()
                        st.session_state.results[file] = result
                        st.session_state.processing_status[file] = result['status']
                    except Exception as e:
                        st.session_state.results[file] = {
                            'file_path': file,
                            'file_name': os.path.basename(file),
                            'status': 'failed',
                            'error': str(e)
                        }
                        st.session_state.processing_status[file] = 'failed'
                    
                    completed += 1
                    progress_bar.progress(completed / len(audio_files))
                    status_text.text(f"Processed {completed}/{len(audio_files)} files")
                    
                    # Update metrics
                    pending_count = sum(1 for s in st.session_state.processing_status.values() if s == 'pending')
                    processing_count = sum(1 for s in st.session_state.processing_status.values() if s == 'processing')
                    completed_count = sum(1 for s in st.session_state.processing_status.values() if s == 'completed')
                    failed_count = sum(1 for s in st.session_state.processing_status.values() if s == 'failed')
                    
                    pending_metric.metric("Pending", pending_count)
                    processing_metric.metric("Processing", processing_count)
                    completed_metric.metric("Completed", completed_count)
                    failed_metric.metric("Failed", failed_count)
                    
                    # Save state after each file
                    save_batch_state(
                        st.session_state.batch_id,
                        audio_files,
                        st.session_state.results,
                        st.session_state.processing_status
                    )
            
            st.success("✅ Batch processing completed!")
    else:
        st.info("Please select audio files to process")

with col2:
    st.header("📊 Processing Results")
    
    if st.session_state.results:
        # Summary statistics
        completed = sum(1 for r in st.session_state.results.values() if r['status'] == 'completed')
        failed = sum(1 for r in st.session_state.results.values() if r['status'] == 'failed')
        
        col2_1, col2_2, col2_3 = st.columns(3)
        with col2_1:
            st.metric("Total Files", len(st.session_state.results))
        with col2_2:
            st.metric("Completed", completed)
        with col2_3:
            st.metric("Failed", failed)
        
        # Results table with better formatting
        st.subheader("📋 Processing Details")
        
        results_data = []
        for file, result in st.session_state.results.items():
            status_emoji = '✅' if result['status'] == 'completed' else '❌'
            results_data.append({
                'File': result['file_name'],
                'Status': f"{status_emoji} {result['status']}",
                'Transcription': '✅' if result.get('transcription') else '❌',
                'Summary': '✅' if result.get('summary') else ('➖' if not has_gemini else '❌'),
                'Time': f"{result.get('processing_time', 0):.1f}s",
                'Error': result.get('error', '')[:50] + '...' if result.get('error') and len(result.get('error', '')) > 50 else result.get('error', '')
            })
        
        df = pd.DataFrame(results_data)
        
        # Style the dataframe
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "File": st.column_config.TextColumn("📄 File Name", width="medium"),
                "Status": st.column_config.TextColumn("📊 Status", width="small"),
                "Transcription": st.column_config.TextColumn("🎤 Transcribed", width="small"),
                "Summary": st.column_config.TextColumn("📝 Summary", width="small"),
                "Time": st.column_config.TextColumn("⏱️ Time", width="small"),
                "Error": st.column_config.TextColumn("⚠️ Error", width="medium"),
            }
        )
        
        # Export options
        st.subheader("💾 Export Results")
        
        export_format = st.selectbox(
            "Export format:",
            ["Individual files", "Combined document"]
        )
        
        if st.button("📥 Generate Export"):
            with st.spinner("Generating export..."):
                export_dir = tempfile.mkdtemp()
                
                if export_format == "Individual files":
                    # Create individual files for each result
                    for file, result in st.session_state.results.items():
                        if result['status'] == 'completed':
                            base_name = Path(result['file_name']).stem
                            
                            # Save transcription
                            if result.get('transcription'):
                                trans_path = os.path.join(export_dir, f"{base_name}_transcription.txt")
                                with open(trans_path, 'w', encoding='utf-8') as f:
                                    f.write(result['transcription'])
                            
                            # Save summary
                            if result.get('summary'):
                                summary_path = os.path.join(export_dir, f"{base_name}_summary.md")
                                with open(summary_path, 'w', encoding='utf-8') as f:
                                    f.write(result['summary'])
                                
                                # Convert to docx if available
                                if HAS_DOCX:
                                    doc = markdown_to_docx(result['summary'], f"{base_name} Summary")
                                    if doc:
                                        doc.save(os.path.join(export_dir, f"{base_name}_summary.docx"))
                
                else:  # Combined document
                    # Create combined markdown
                    combined_md = "# Batch Processing Results\n\n"
                    combined_md += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                    
                    for file, result in st.session_state.results.items():
                        if result['status'] == 'completed':
                            combined_md += f"## {result['file_name']}\n\n"
                            
                            if result.get('transcription'):
                                combined_md += "### Transcription\n\n"
                                combined_md += result['transcription'] + "\n\n"
                            
                            if result.get('summary'):
                                combined_md += "### Summary\n\n"
                                combined_md += result['summary'] + "\n\n"
                            
                            combined_md += "---\n\n"
                    
                    # Save combined markdown
                    md_path = os.path.join(export_dir, "combined_results.md")
                    with open(md_path, 'w', encoding='utf-8') as f:
                        f.write(combined_md)
                    
                    # Convert to docx if available
                    if HAS_DOCX:
                        doc = markdown_to_docx(combined_md, "Batch Processing Results")
                        if doc:
                            doc.save(os.path.join(export_dir, "combined_results.docx"))
                
                # Create zip file
                zip_path = os.path.join(tempfile.gettempdir(), "batch_results.zip")
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for root, dirs, files in os.walk(export_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, export_dir)
                            zipf.write(file_path, arcname)
                
                # Download button
                with open(zip_path, 'rb') as f:
                    st.download_button(
                        label="📥 Download Results (ZIP)",
                        data=f.read(),
                        file_name=f"batch_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip"
                    )
                
                # Cleanup
                shutil.rmtree(export_dir)
                os.remove(zip_path)
        
        # Individual file details
        st.markdown("---")
        st.subheader("📄 File Details")
        
        completed_files = [r['file_name'] for r in st.session_state.results.values() if r['status'] == 'completed']
        
        if completed_files:
            selected_file = st.selectbox(
                "Select file to view details:",
                completed_files,
                format_func=lambda x: f"📄 {x}"
            )
        
            if selected_file:
                for file, result in st.session_state.results.items():
                    if result['file_name'] == selected_file:
                        # Show processing info
                        col1_info, col2_info = st.columns(2)
                        with col1_info:
                            st.metric("Processing Time", f"{result.get('processing_time', 0):.1f}s")
                        with col2_info:
                            st.metric("Status", "✅ Completed")
                        
                        # Tabs for content
                        tabs = st.tabs(["🎤 Transcription", "📝 Summary"])
                        
                        with tabs[0]:
                            if result.get('transcription'):
                                st.text_area("", result['transcription'], height=400, label_visibility="collapsed")
                            else:
                                st.info("No transcription available")
                        
                        with tabs[1]:
                            if result.get('summary'):
                                st.markdown(result['summary'])
                            else:
                                if not has_gemini:
                                    st.warning("Summary generation is disabled (no Google API key)")
                                else:
                                    st.info("No summary available")
                        
                        break
        else:
            st.info("No completed files to display. Process some audio files first!")
    else:
        st.info("No processing results yet. Start batch processing to see results here.")

# Sidebar with instructions
with st.sidebar:
    st.header("📖 Instructions")
    st.markdown("""
    ### How to use:
    
    1. **Select input method:**
       - Upload multiple audio files directly
       - Or specify a folder path to process all audio files recursively
    
    2. **Configure settings:**
       - Choose transcription model (GPT-4o or GPT-4o-mini)
       - Enable/disable summary generation
       - Customize summary prompt if needed
    
    3. **Start batch processing:**
       - Set number of parallel threads
       - Click "Start Batch Processing"
    
    4. **Export results:**
       - Individual files: Each audio gets separate transcription and summary files
       - Combined document: All results in one file
       - Downloads as ZIP archive with .txt, .md, and .docx formats
    
    ### Features:
    - ✅ Automatic file splitting for large audio
    - ✅ Recursive folder scanning
    - ✅ Parallel processing
    - ✅ Progress tracking
    - ✅ Multiple export formats
    - ✅ Word document generation
    """)
    
    st.divider()
    
    st.markdown("""
    ### Supported formats:
    - MP3, WAV, M4A, FLAC
    - AAC, OGG, WMA, OPUS, WebM
    
    ### File limits:
    - Max 25MB per file (auto-split if larger)
    - Files split into 5-minute segments
    """)