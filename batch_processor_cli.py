#!/usr/bin/env python3
"""
Command-line interface for batch audio processing
支援多檔案批次處理，自動化轉錄和摘要生成
"""

import os
import sys
import argparse
import json
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from tqdm import tqdm
from dotenv import load_dotenv

load_dotenv()

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from audio2text.gpt4o_stt import transcribe_audio_gpt4o
from gemini_utils import call_gemini_api
from utils import check_file_constraints, split_large_audio
from docx import Document
from markdown2 import markdown
from bs4 import BeautifulSoup
import tempfile
import shutil

class BatchProcessor:
    def __init__(self, model="gpt-4o-mini-transcribe", output_format="markdown", max_workers=2):
        self.model = model
        self.output_format = output_format
        self.max_workers = max_workers
        self.results = {}
        
        # Check API keys
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY not found in environment")
        if not os.getenv("GOOGLE_API_KEY"):
            raise ValueError("GOOGLE_API_KEY not found in environment")
    
    def find_audio_files(self, path):
        """Find all audio files in path (file or directory)"""
        audio_extensions = {'.mp3', '.wav', '.m4a', '.flac', '.aac', '.ogg', '.wma', '.opus', '.webm'}
        audio_files = []
        
        if os.path.isfile(path):
            if Path(path).suffix.lower() in audio_extensions:
                audio_files.append(path)
        elif os.path.isdir(path):
            for root, dirs, files in os.walk(path):
                for file in files:
                    if Path(file).suffix.lower() in audio_extensions:
                        audio_files.append(os.path.join(root, file))
        
        return sorted(audio_files)
    
    def process_file(self, file_path, summary_prompt=None):
        """Process a single audio file"""
        result = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'status': 'processing',
            'transcription': None,
            'summary': None,
            'error': None,
            'processing_time': 0
        }
        
        start_time = time.time()
        
        try:
            # Check file constraints
            size_ok, duration_ok, file_size, duration = check_file_constraints(file_path)
            
            if not size_ok or not duration_ok:
                # Split large files
                print(f"  Splitting large file: {os.path.basename(file_path)}")
                with tempfile.TemporaryDirectory() as temp_dir:
                    segments = split_large_audio(file_path, temp_dir)
                    transcriptions = []
                    
                    for i, segment in enumerate(segments):
                        print(f"    Processing segment {i+1}/{len(segments)}")
                        segment_transcription = transcribe_audio_gpt4o(
                            segment,
                            api_key=os.getenv("OPENAI_API_KEY"),
                            model=self.model,
                            output_format=self.output_format if self.output_format in ['text', 'srt'] else 'text'
                        )
                        if segment_transcription:
                            transcriptions.append(segment_transcription)
                    
                    result['transcription'] = "\n\n".join(transcriptions)
            else:
                # Process normally
                result['transcription'] = transcribe_audio_gpt4o(
                    file_path,
                    api_key=os.getenv("OPENAI_API_KEY"),
                    model=self.model,
                    output_format=self.output_format if self.output_format in ['text', 'srt'] else 'text'
                )
            
            if result['transcription']:
                # Generate summary if prompt provided
                if summary_prompt:
                    print(f"  Generating summary...")
                    summary_text = f"{summary_prompt}\n\nTranscription:\n{result['transcription']}"
                    result['summary'] = call_gemini_api(summary_text, model="gemini-2.0-flash-exp")
                
                result['status'] = 'completed'
            else:
                result['status'] = 'failed'
                result['error'] = 'Transcription failed'
                
        except Exception as e:
            result['status'] = 'failed'
            result['error'] = str(e)
            print(f"  Error: {str(e)}")
        
        result['processing_time'] = time.time() - start_time
        return result
    
    def process_batch(self, audio_files, summary_prompt=None, output_dir=None):
        """Process multiple audio files in parallel"""
        if not output_dir:
            output_dir = f"batch_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"\n🎵 Processing {len(audio_files)} audio files...")
        print(f"📁 Output directory: {output_dir}")
        print(f"🔧 Model: {self.model}")
        print(f"📄 Format: {self.output_format}")
        print(f"⚙️  Workers: {self.max_workers}\n")
        
        # Process files with progress bar
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_file = {
                executor.submit(self.process_file, file, summary_prompt): file
                for file in audio_files
            }
            
            with tqdm(total=len(audio_files), desc="Processing") as pbar:
                for future in as_completed(future_to_file):
                    file = future_to_file[future]
                    try:
                        result = future.result()
                        self.results[file] = result
                        
                        # Save individual results
                        if result['status'] == 'completed':
                            self.save_result(result, output_dir)
                            pbar.set_postfix({'status': '✅', 'file': os.path.basename(file)})
                        else:
                            pbar.set_postfix({'status': '❌', 'file': os.path.basename(file)})
                    except Exception as e:
                        self.results[file] = {
                            'file_path': file,
                            'file_name': os.path.basename(file),
                            'status': 'failed',
                            'error': str(e)
                        }
                        pbar.set_postfix({'status': '❌', 'file': os.path.basename(file)})
                    
                    pbar.update(1)
        
        # Generate summary report
        self.generate_report(output_dir)
        
        # Print results summary
        completed = sum(1 for r in self.results.values() if r['status'] == 'completed')
        failed = sum(1 for r in self.results.values() if r['status'] == 'failed')
        
        print(f"\n✅ Completed: {completed}")
        print(f"❌ Failed: {failed}")
        print(f"📊 Total processing time: {sum(r.get('processing_time', 0) for r in self.results.values()):.1f}s")
        print(f"\n📁 Results saved to: {output_dir}")
    
    def save_result(self, result, output_dir):
        """Save individual result files"""
        base_name = Path(result['file_name']).stem
        
        # Save transcription
        if result.get('transcription'):
            ext = 'srt' if self.output_format == 'srt' else 'txt'
            trans_path = os.path.join(output_dir, f"{base_name}_transcription.{ext}")
            with open(trans_path, 'w', encoding='utf-8') as f:
                f.write(result['transcription'])
        
        # Save summary
        if result.get('summary'):
            summary_path = os.path.join(output_dir, f"{base_name}_summary.md")
            with open(summary_path, 'w', encoding='utf-8') as f:
                f.write(result['summary'])
            
            # Also save as docx
            doc = self.markdown_to_docx(result['summary'], f"{base_name} Summary")
            doc.save(os.path.join(output_dir, f"{base_name}_summary.docx"))
    
    def markdown_to_docx(self, markdown_text, title):
        """Convert markdown to Word document"""
        doc = Document()
        doc.add_heading(title, 0)
        
        html = markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'pre']):
            if element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'pre':
                doc.add_paragraph(element.get_text(), style='Normal')
        
        return doc
    
    def generate_report(self, output_dir):
        """Generate processing report"""
        report = {
            'timestamp': datetime.now().isoformat(),
            'model': self.model,
            'format': self.output_format,
            'total_files': len(self.results),
            'completed': sum(1 for r in self.results.values() if r['status'] == 'completed'),
            'failed': sum(1 for r in self.results.values() if r['status'] == 'failed'),
            'total_time': sum(r.get('processing_time', 0) for r in self.results.values()),
            'files': []
        }
        
        for file, result in self.results.items():
            report['files'].append({
                'name': result['file_name'],
                'status': result['status'],
                'time': result.get('processing_time', 0),
                'error': result.get('error', None)
            })
        
        # Save JSON report
        report_path = os.path.join(output_dir, 'processing_report.json')
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        # Save markdown report
        md_report = f"""# Batch Processing Report

Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Summary
- **Model**: {self.model}
- **Format**: {self.output_format}
- **Total Files**: {report['total_files']}
- **Completed**: {report['completed']} ✅
- **Failed**: {report['failed']} ❌
- **Total Time**: {report['total_time']:.1f}s

## File Details
| File | Status | Time (s) | Error |
|------|--------|----------|-------|
"""
        
        for file_info in report['files']:
            status_icon = '✅' if file_info['status'] == 'completed' else '❌'
            error = file_info['error'] or '-'
            md_report += f"| {file_info['name']} | {status_icon} | {file_info['time']:.1f} | {error} |\n"
        
        report_md_path = os.path.join(output_dir, 'processing_report.md')
        with open(report_md_path, 'w', encoding='utf-8') as f:
            f.write(md_report)

def main():
    parser = argparse.ArgumentParser(
        description="Batch process audio files with GPT-4o transcription and Gemini summarization",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process single file
  python batch_processor_cli.py audio.mp3
  
  # Process entire folder
  python batch_processor_cli.py /path/to/audio/folder
  
  # With custom model and format
  python batch_processor_cli.py folder/ --model gpt-4o-transcribe --format srt
  
  # With summary generation
  python batch_processor_cli.py folder/ --summary "Provide key points from this meeting"
  
  # Custom output directory and parallel workers
  python batch_processor_cli.py folder/ --output results/ --workers 4
"""
    )
    
    parser.add_argument('input', help='Audio file or directory path')
    parser.add_argument('--model', default='gpt-4o-mini-transcribe',
                        choices=['gpt-4o-transcribe', 'gpt-4o-mini-transcribe'],
                        help='Transcription model (default: gpt-4o-mini-transcribe)')
    parser.add_argument('--format', default='text',
                        choices=['text', 'srt', 'markdown'],
                        help='Output format (default: text)')
    parser.add_argument('--summary', help='Generate summary with this prompt')
    parser.add_argument('--output', help='Output directory (default: batch_output_TIMESTAMP)')
    parser.add_argument('--workers', type=int, default=2,
                        help='Number of parallel workers (default: 2)')
    
    args = parser.parse_args()
    
    # Initialize processor
    try:
        processor = BatchProcessor(
            model=args.model,
            output_format=args.format,
            max_workers=args.workers
        )
    except ValueError as e:
        print(f"Error: {e}")
        print("Please set required API keys in .env file")
        sys.exit(1)
    
    # Find audio files
    audio_files = processor.find_audio_files(args.input)
    
    if not audio_files:
        print(f"No audio files found in: {args.input}")
        sys.exit(1)
    
    # Process batch
    processor.process_batch(
        audio_files,
        summary_prompt=args.summary,
        output_dir=args.output
    )

if __name__ == "__main__":
    main()