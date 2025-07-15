#!/usr/bin/env python3
"""
æ‰¹æ¬¡éŸ³è¨Šè™•ç†è‡ªå‹•åŒ–ç¨‹å¼
Batch Audio Processor for ADA 2025 Conference

æ­¤ç¨‹å¼ç”¨æ–¼è‡ªå‹•åŒ–è™•ç†éŸ³è¨Šæª”æ¡ˆï¼ŒåŒ…å«ï¼š
1. æœç´¢æŒ‡å®šè³‡æ–™å¤¾ä¸­çš„æ‰€æœ‰éŸ³è¨Šæª”æ¡ˆå’Œæ ¼å¼åŒ–æ–‡å­—æ–‡ä»¶
2. ä½¿ç”¨ GPT-4o é€²è¡ŒèªéŸ³è½‰éŒ„
3. ä½¿ç”¨ Gemini 2.5 Pro é€²è¡Œæ‘˜è¦è™•ç†
4. å°‡çµæœè½‰æ›ç‚º Word æ–‡ä»¶

ä½¿ç”¨æ–¹æ³•:
    python batch_audio_processor.py
    æˆ–
    python batch_audio_processor.py <è³‡æ–™å¤¾è·¯å¾‘>

ç’°å¢ƒéœ€æ±‚:
- .env æª”æ¡ˆä¸­éœ€åŒ…å« OPENAI_API_KEY å’Œ GOOGLE_API_KEY
- éŸ³è¨Šæª”æ¡ˆéœ€æœ‰åŒåçš„æ–‡å­—æª”æ¡ˆä½œç‚º agenda å…§å®¹
"""

import os
import sys
import glob
import logging
from pathlib import Path
from typing import List, Dict, Optional
from dotenv import load_dotenv
from openai import OpenAI
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import time

# å°å…¥ç¾æœ‰çš„å·¥å…·å‡½æ•¸
from utils import split_large_audio, check_file_size
from audio2text.gpt4o_stt import transcribe_audio_gpt4o
from pydub import AudioSegment

# è¼‰å…¥ç’°å¢ƒè®Šæ•¸
load_dotenv()

# è¨­å®šæ—¥èªŒ
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('batch_processor.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# æ”¯æ´çš„éŸ³è¨Šæ ¼å¼
SUPPORTED_AUDIO_FORMATS = [
    '*.mp3', '*.wav', '*.m4a', '*.aac', '*.flac', '*.ogg',
    '*.wma', '*.mp4', '*.mov', '*.avi', '*.mkv', '*.webm'
]

# æ”¯æ´çš„æ–‡å­—æ–‡ä»¶æ ¼å¼
SUPPORTED_TEXT_FORMATS = [
    '*.txt', '*.md', '*.rtf', '*.doc', '*.docx', '*.pdf',
    '*.html', '*.htm', '*.xml', '*.json', '*.csv'
]

# ç³»çµ±æç¤ºè©
SYSTEM_PROMPT = """This is American Diabetes Association 2025å¹´æœƒï¼Œèšç„¦åœ¨ç³–å°¿ç—…ï¼Œä»£è¬ï¼Œè‚¥èƒ–ç­‰ç­‰ä¸»é¡Œã€‚å„ªåŒ–æ¼”è¬›è€… transcribe content æˆå®Œæ•´çš„ç­†è¨˜.

rewrite or é †æš¢æ¼”è¬›è€…å…§å®¹ï¼Œä¸¦å„˜å¯èƒ½å®Œæ•´å‘ˆç¾ï¼ŒåŠ å…¥ä¸€äº›ç²—é«”æˆ–åº•ç·šä¾†å‘ˆç¾é‡é»ï¼Œæˆ–åŠ å…¥å°çµä¾†ç¸½çµæˆ–å»¶ä¼¸æ‚¨çš„æƒ³æ³•ã€‚ ä¸¦ä¾ä½¿ç”¨è€…æä¾›çš„ agenda æ•´ç†ã€‚as detail , professional and comprehensive as you can in zh-tw

ä¾ä»¥ä¸‹çš„ agenda ä¾†åˆ†ä¸»è¦æ®µè½ã€‚"""


class BatchAudioProcessor:
    """æ‰¹æ¬¡éŸ³è¨Šè™•ç†å™¨"""
    
    def __init__(self, model="gpt-4o-mini-transcribe", output_format="text"):
        """
        åˆå§‹åŒ–è™•ç†å™¨
        
        Args:
            model: ä½¿ç”¨çš„è½‰éŒ„æ¨¡å‹ (gpt-4o-transcribe æˆ– gpt-4o-mini-transcribe)
            output_format: è¼¸å‡ºæ ¼å¼ (text, markdown, srt)
        """
        self.openai_client = None
        self.transcribe_model = model
        self.output_format = output_format
        self.setup_apis()
    
    def setup_apis(self):
        """è¨­å®š API å®¢æˆ¶ç«¯"""
        # è¨­å®š OpenAI API
        openai_key = os.getenv('OPENAI_API_KEY')
        if not openai_key:
            raise ValueError("è«‹åœ¨ .env æª”æ¡ˆä¸­è¨­å®š OPENAI_API_KEY")

        self.openai_client = OpenAI(api_key=openai_key)
        logger.info("OpenAI API è¨­å®šå®Œæˆ")

        # è¨­å®š Google Gemini API
        google_key = 'AIzaSyBUNvJo_D2KZV3UVVgQxvFlZC1aFfXIw9k'  # ä½¿ç”¨æä¾›çš„ API key
        if not google_key:
            google_key = os.getenv('GOOGLE_API_KEY')
            if not google_key:
                raise ValueError("è«‹åœ¨ .env æª”æ¡ˆä¸­è¨­å®š GOOGLE_API_KEY")

        genai.configure(api_key=google_key)
        logger.info("Google Gemini API è¨­å®šå®Œæˆ")
    
    def get_folder_path(self) -> str:
        """
        å–å¾—ä½¿ç”¨è€…è¼¸å…¥çš„è³‡æ–™å¤¾è·¯å¾‘

        Returns:
            è³‡æ–™å¤¾è·¯å¾‘å­—ä¸²
        """
        if len(sys.argv) > 1:
            folder_path = sys.argv[1]
        else:
            print("=== æ‰¹æ¬¡éŸ³è¨Šè™•ç†è‡ªå‹•åŒ–ç¨‹å¼ ===")
            print("è«‹è¼¸å…¥è¦è™•ç†çš„è³‡æ–™å¤¾è·¯å¾‘:")
            folder_path = input("> ").strip()

        # ç§»é™¤è·¯å¾‘å…©ç«¯çš„å¼•è™Ÿ
        folder_path = folder_path.strip('"\'')

        # é©—è­‰è·¯å¾‘æ˜¯å¦å­˜åœ¨
        if not Path(folder_path).exists():
            raise ValueError(f"è³‡æ–™å¤¾ä¸å­˜åœ¨: {folder_path}")

        return folder_path

    def find_files(self, folder_path: str) -> Dict[str, List[str]]:
        """
        åœ¨æŒ‡å®šè³‡æ–™å¤¾ä¸­éæ­¸æœç´¢éŸ³è¨Šæª”æ¡ˆå’Œæ–‡å­—æ–‡ä»¶

        Args:
            folder_path: è¦æœç´¢çš„è³‡æ–™å¤¾è·¯å¾‘

        Returns:
            åŒ…å«éŸ³è¨Šæª”æ¡ˆå’Œæ–‡å­—æ–‡ä»¶çš„å­—å…¸
        """
        folder_path = Path(folder_path)

        if not folder_path.exists():
            logger.error(f"è³‡æ–™å¤¾ä¸å­˜åœ¨: {folder_path}")
            return {'audio': [], 'text': []}

        logger.info(f"æœç´¢è³‡æ–™å¤¾: {folder_path}")

        # æœç´¢éŸ³è¨Šæª”æ¡ˆ
        audio_files = []
        for pattern in SUPPORTED_AUDIO_FORMATS:
            files = glob.glob(str(folder_path / "**" / pattern),
                              recursive=True)
            audio_files.extend(files)

        # æœç´¢æ–‡å­—æ–‡ä»¶
        text_files = []
        for pattern in SUPPORTED_TEXT_FORMATS:
            files = glob.glob(str(folder_path / "**" / pattern),
                              recursive=True)
            text_files.extend(files)

        # å»é‡ä¸¦æ’åº
        audio_files = sorted(list(set(audio_files)))
        text_files = sorted(list(set(text_files)))

        logger.info(f"æ‰¾åˆ° {len(audio_files)} å€‹éŸ³è¨Šæª”æ¡ˆ")
        for file in audio_files:
            logger.info(f"  - {file}")

        logger.info(f"æ‰¾åˆ° {len(text_files)} å€‹æ–‡å­—æ–‡ä»¶")
        for file in text_files:
            logger.info(f"  - {file}")

        return {'audio': audio_files, 'text': text_files}
    
    def transcribe_audio(self, audio_path: str) -> Optional[str]:
        """
        ä½¿ç”¨ GPT-4o è½‰éŒ„éŸ³è¨Šæª”æ¡ˆï¼Œæ”¯æ´å¤§æª”æ¡ˆåˆ†å‰²

        Args:
            audio_path: éŸ³è¨Šæª”æ¡ˆè·¯å¾‘

        Returns:
            è½‰éŒ„æ–‡å­—ï¼Œå¤±æ•—æ™‚è¿”å› None
        """
        try:
            logger.info(f"é–‹å§‹è½‰éŒ„éŸ³è¨Š: {audio_path}")
            
            # æª¢æŸ¥æª”æ¡ˆå¤§å°ï¼Œå¦‚æœå¤ªå¤§å‰‡åˆ†å‰²
            if check_file_size(audio_path):
                logger.info(f"æª”æ¡ˆè¼ƒå¤§ï¼Œé€²è¡Œåˆ†å‰²è™•ç†: {audio_path}")
                segments = split_large_audio(audio_path)
                
                if not segments:
                    logger.error("éŸ³è¨Šåˆ†å‰²å¤±æ•—")
                    return None
                
                # åˆ†åˆ¥è½‰éŒ„æ¯å€‹ç‰‡æ®µ
                full_transcript = ""
                for i, segment_path in enumerate(segments):
                    logger.info(f"è½‰éŒ„ç‰‡æ®µ {i+1}/{len(segments)}: {segment_path}")
                    
                    try:
                        with open(segment_path, "rb") as audio_file:
                            transcript = self.openai_client.audio.transcriptions.create(
                                model=self.transcribe_model,
                                file=audio_file,
                                language="zh",
                                response_format="text"
                            )
                        
                        full_transcript += transcript.text + " "
                        logger.info(f"ç‰‡æ®µ {i+1} è½‰éŒ„å®Œæˆ")
                        
                        # æ¸…ç†è‡¨æ™‚æª”æ¡ˆ
                        if os.path.exists(segment_path):
                            os.remove(segment_path)
                            
                    except Exception as e:
                        logger.error(f"ç‰‡æ®µ {i+1} è½‰éŒ„å¤±æ•—: {e}")
                        # æ¸…ç†è‡¨æ™‚æª”æ¡ˆ
                        if os.path.exists(segment_path):
                            os.remove(segment_path)
                        continue
                    
                    # é¿å… API é™åˆ¶ï¼ŒåŠ å…¥å»¶é²
                    time.sleep(2)
                
                if full_transcript.strip():
                    logger.info(f"æ‰€æœ‰ç‰‡æ®µè½‰éŒ„å®Œæˆï¼Œç¸½é•·åº¦: {len(full_transcript)}")
                    return full_transcript.strip()
                else:
                    logger.error("æ‰€æœ‰ç‰‡æ®µè½‰éŒ„éƒ½å¤±æ•—")
                    return None
            
            else:
                # æª”æ¡ˆå¤§å°æ­£å¸¸ï¼Œç›´æ¥è½‰éŒ„
                with open(audio_path, "rb") as audio_file:
                    transcript = self.openai_client.audio.transcriptions.create(
                        model=self.transcribe_model,
                        file=audio_file,
                        language="zh",
                        response_format="text"
                    )

                logger.info(f"è½‰éŒ„å®Œæˆï¼Œæ–‡å­—é•·åº¦: {len(transcript.text)}")
                return transcript.text

        except Exception as e:
            logger.error(f"è½‰éŒ„å¤±æ•—: {e}")
            return None
    
    def read_agenda_file(self, audio_path: str,
                         text_files: List[str]) -> Optional[str]:
        """
        è®€å–èˆ‡éŸ³è¨Šæª”æ¡ˆåŒåçš„è­°ç¨‹æ–‡å­—æª”æ¡ˆ

        Args:
            audio_path: éŸ³è¨Šæª”æ¡ˆè·¯å¾‘
            text_files: å¯ç”¨çš„æ–‡å­—æª”æ¡ˆåˆ—è¡¨

        Returns:
            è­°ç¨‹å…§å®¹ï¼Œæ‰¾ä¸åˆ°æ™‚è¿”å› None
        """
        audio_path = Path(audio_path)
        audio_stem = audio_path.stem

        # å°‹æ‰¾åŒåçš„æ–‡å­—æª”æ¡ˆ
        for text_file in text_files:
            text_path = Path(text_file)
            if text_path.stem == audio_stem:
                try:
                    with open(text_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                    if content:
                        logger.info(f"æ‰¾åˆ°è­°ç¨‹æª”æ¡ˆ: {text_path}")
                        return content
                except Exception as e:
                    logger.warning(f"è®€å–è­°ç¨‹æª”æ¡ˆå¤±æ•— {text_path}: {e}")

        # å¦‚æœæ²’æœ‰æ‰¾åˆ°åŒåæª”æ¡ˆï¼Œå˜—è©¦å¸¸è¦‹çš„è­°ç¨‹æª”æ¡ˆå
        agenda_names = ['agenda', 'schedule', 'program', 'è­°ç¨‹']
        for text_file in text_files:
            text_path = Path(text_file)
            if any(name in text_path.stem.lower() for name in agenda_names):
                try:
                    with open(text_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                    if content:
                        logger.info(f"æ‰¾åˆ°è­°ç¨‹æª”æ¡ˆ: {text_path}")
                        return content
                except Exception as e:
                    logger.warning(f"è®€å–è­°ç¨‹æª”æ¡ˆå¤±æ•— {text_path}: {e}")

        logger.warning(f"æœªæ‰¾åˆ°è­°ç¨‹æª”æ¡ˆ: {audio_stem}")
        return None
    
    def summarize_with_gemini(self, transcript: str,
                              agenda: Optional[str] = None) -> Optional[str]:
        """
        ä½¿ç”¨ Gemini 2.5 Pro é€²è¡Œæ‘˜è¦è™•ç†

        Args:
            transcript: è½‰éŒ„æ–‡å­—
            agenda: è­°ç¨‹å…§å®¹

        Returns:
            æ‘˜è¦æ–‡å­—ï¼Œå¤±æ•—æ™‚è¿”å› None
        """
        try:
            logger.info("é–‹å§‹ä½¿ç”¨ Gemini 2.5 Pro é€²è¡Œæ‘˜è¦è™•ç†")

            # å»ºç«‹æ¨¡å‹ - ä½¿ç”¨æœ€æ–°çš„ Gemini 2.5 Pro
            model = genai.GenerativeModel('gemini-2.5-pro')

            # æ§‹å»ºæç¤ºè©
            user_prompt = f"""è«‹æ ¹æ“šä»¥ä¸‹è½‰éŒ„å…§å®¹é€²è¡Œæ‘˜è¦è™•ç†ï¼š

è½‰éŒ„å…§å®¹ï¼š
{transcript}
"""

            if agenda:
                user_prompt += f"""

è­°ç¨‹å…§å®¹ï¼š
{agenda}
"""

            # ç”Ÿæˆæ‘˜è¦
            response = model.generate_content([
                {"role": "user", "parts": [{"text": SYSTEM_PROMPT}]},
                {"role": "user", "parts": [{"text": user_prompt}]}
            ])

            summary = response.text
            logger.info(f"æ‘˜è¦å®Œæˆï¼Œé•·åº¦: {len(summary)}")
            return summary

        except Exception as e:
            logger.error(f"æ‘˜è¦è™•ç†å¤±æ•—: {e}")
            return None
    
    def markdown_to_docx(self, markdown_text: str, output_path: str) -> bool:
        """
        å°‡ Markdown æ–‡å­—è½‰æ›ç‚º Word æ–‡ä»¶

        Args:
            markdown_text: Markdown æ ¼å¼æ–‡å­—
            output_path: è¼¸å‡ºæª”æ¡ˆè·¯å¾‘

        Returns:
            è½‰æ›æ˜¯å¦æˆåŠŸ
        """
        try:
            logger.info(f"é–‹å§‹è½‰æ›ç‚º Word æ–‡ä»¶: {output_path}")

            doc = Document()

            # åˆ†è¡Œè™•ç†
            lines = markdown_text.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # è™•ç†æ¨™é¡Œ
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()

                    if level == 1:
                        doc.add_heading(title_text, level=1)
                    elif level == 2:
                        doc.add_heading(title_text, level=2)
                    elif level == 3:
                        doc.add_heading(title_text, level=3)
                    else:
                        doc.add_heading(title_text, level=4)

                    continue

                # è™•ç†æ™®é€šæ®µè½
                paragraph = doc.add_paragraph()

                # è™•ç†ç²—é«”å’Œåº•ç·š
                parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', line)

                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # ç²—é«”
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('__') and part.endswith('__'):
                        # åº•ç·š
                        run = paragraph.add_run(part[2:-2])
                        run.underline = True
                    elif part.startswith('*') and part.endswith('*'):
                        # æ–œé«”
                        run = paragraph.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('_') and part.endswith('_'):
                        # åº•ç·š
                        run = paragraph.add_run(part[1:-1])
                        run.underline = True
                    else:
                        # æ™®é€šæ–‡å­—
                        paragraph.add_run(part)

            # å„²å­˜æ–‡ä»¶
            doc.save(output_path)
            logger.info(f"Word æ–‡ä»¶å·²å„²å­˜: {output_path}")
            return True

        except Exception as e:
            logger.error(f"è½‰æ›ç‚º Word æ–‡ä»¶å¤±æ•—: {e}")
            return False
    
    def process_single_file(self, audio_path: str,
                            text_files: List[str]) -> Dict:
        """
        è™•ç†å–®ä¸€éŸ³è¨Šæª”æ¡ˆ

        Args:
            audio_path: éŸ³è¨Šæª”æ¡ˆè·¯å¾‘
            text_files: å¯ç”¨çš„æ–‡å­—æª”æ¡ˆåˆ—è¡¨

        Returns:
            è™•ç†çµæœå­—å…¸
        """
        result = {
            'file': audio_path,
            'success': False,
            'transcript': None,
            'summary': None,
            'docx_path': None,
            'error': None
        }

        try:
            # 1. è½‰éŒ„éŸ³è¨Š
            transcript = self.transcribe_audio(audio_path)
            if not transcript:
                result['error'] = "è½‰éŒ„å¤±æ•—"
                return result

            result['transcript'] = transcript

            # 2. è®€å–è­°ç¨‹æª”æ¡ˆ
            agenda = self.read_agenda_file(audio_path, text_files)

            # 3. ç”Ÿæˆæ‘˜è¦
            summary = self.summarize_with_gemini(transcript, agenda)
            if not summary:
                result['error'] = "æ‘˜è¦è™•ç†å¤±æ•—"
                return result

            result['summary'] = summary

            # 4. è½‰æ›ç‚º Word æ–‡ä»¶
            audio_path_obj = Path(audio_path)
            docx_path = audio_path_obj.with_suffix('.docx')

            if self.markdown_to_docx(summary, str(docx_path)):
                result['docx_path'] = str(docx_path)
                result['success'] = True
            else:
                result['error'] = "Word æ–‡ä»¶è½‰æ›å¤±æ•—"

        except Exception as e:
            result['error'] = str(e)
            logger.error(f"è™•ç†æª”æ¡ˆæ™‚ç™¼ç”ŸéŒ¯èª¤: {e}")

        return result
    
    def process_folder(self, folder_path: str) -> List[Dict]:
        """
        è™•ç†æ•´å€‹è³‡æ–™å¤¾ä¸­çš„éŸ³è¨Šæª”æ¡ˆ

        Args:
            folder_path: è³‡æ–™å¤¾è·¯å¾‘

        Returns:
            è™•ç†çµæœåˆ—è¡¨
        """
        logger.info(f"é–‹å§‹è™•ç†è³‡æ–™å¤¾: {folder_path}")

        # æ‰¾åˆ°æ‰€æœ‰æª”æ¡ˆ
        files = self.find_files(folder_path)
        audio_files = files['audio']
        text_files = files['text']

        if not audio_files:
            logger.warning("æœªæ‰¾åˆ°ä»»ä½•éŸ³è¨Šæª”æ¡ˆ")
            return []

        results = []

        for i, audio_file in enumerate(audio_files, 1):
            logger.info(f"è™•ç†æª”æ¡ˆ {i}/{len(audio_files)}: {audio_file}")

            result = self.process_single_file(audio_file, text_files)
            results.append(result)

            if result['success']:
                logger.info(f"âœ… æˆåŠŸè™•ç†: {audio_file}")
            else:
                logger.error(f"âŒ è™•ç†å¤±æ•—: {audio_file} - {result['error']}")

            # é¿å… API é™åˆ¶ï¼ŒåŠ å…¥å»¶é²
            time.sleep(1)

        return results
    
    def generate_report(self, results: List[Dict], output_path: str):
        """
        ç”Ÿæˆè™•ç†çµæœå ±å‘Š

        Args:
            results: è™•ç†çµæœåˆ—è¡¨
            output_path: å ±å‘Šè¼¸å‡ºè·¯å¾‘
        """
        try:
            doc = Document()

            # æ¨™é¡Œ
            title = doc.add_heading('æ‰¹æ¬¡éŸ³è¨Šè™•ç†å ±å‘Š', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # çµ±è¨ˆè³‡è¨Š
            total_files = len(results)
            success_files = sum(1 for r in results if r['success'])
            failed_files = total_files - success_files

            stats = doc.add_paragraph()
            stats.add_run(f"ç¸½æª”æ¡ˆæ•¸: {total_files}\n").bold = True
            stats.add_run(f"æˆåŠŸè™•ç†: {success_files}\n").bold = True
            stats.add_run(f"è™•ç†å¤±æ•—: {failed_files}\n").bold = True

            # è©³ç´°çµæœ
            doc.add_heading('è™•ç†çµæœè©³ç´°', level=2)

            for result in results:
                file_name = Path(result['file']).name

                if result['success']:
                    status = "âœ… æˆåŠŸ"
                    doc.add_paragraph(f"{file_name}: {status}")
                    doc.add_paragraph(f"  è¼¸å‡ºæ–‡ä»¶: {result['docx_path']}")
                else:
                    status = "âŒ å¤±æ•—"
                    doc.add_paragraph(f"{file_name}: {status}")
                    doc.add_paragraph(f"  éŒ¯èª¤: {result['error']}")

                doc.add_paragraph()  # ç©ºè¡Œ

            doc.save(output_path)
            logger.info(f"å ±å‘Šå·²å„²å­˜: {output_path}")

        except Exception as e:
            logger.error(f"ç”Ÿæˆå ±å‘Šå¤±æ•—: {e}")


def main():
    """ä¸»å‡½æ•¸"""
    try:
        # æª¢æŸ¥å‘½ä»¤è¡Œåƒæ•¸ä¸­æ˜¯å¦æŒ‡å®šæ¨¡å‹
        model = "gpt-4o-mini-transcribe"  # é è¨­ä½¿ç”¨ mini ç‰ˆæœ¬
        if len(sys.argv) > 2 and sys.argv[2] in ["gpt-4o-transcribe", "gpt-4o-mini-transcribe"]:
            model = sys.argv[2]
        
        print(f"ğŸ¤– ä½¿ç”¨è½‰éŒ„æ¨¡å‹: {model}")
        processor = BatchAudioProcessor(model=model)
        folder_path = processor.get_folder_path()
        results = processor.process_folder(folder_path)

        # ç”Ÿæˆå ±å‘Š
        report_path = Path(folder_path) / "processing_report.docx"
        processor.generate_report(results, str(report_path))

        # é¡¯ç¤ºçµæœçµ±è¨ˆ
        total_files = len(results)
        success_files = sum(1 for r in results if r['success'])

        print("\n=== è™•ç†å®Œæˆ ===")
        print(f"ç¸½æª”æ¡ˆæ•¸: {total_files}")
        print(f"æˆåŠŸè™•ç†: {success_files}")
        print(f"è™•ç†å¤±æ•—: {total_files - success_files}")
        print(f"å ±å‘Šä½ç½®: {report_path}")

    except Exception as e:
        logger.error(f"ç¨‹å¼åŸ·è¡Œå¤±æ•—: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
