#!/usr/bin/env python3
"""
SRT 格式演講稿與投影片智能合併系統
自動根據時間插入所有截圖，確保 DOCX 中顯示實際圖片
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import List, Tuple, Dict, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import glob
from dotenv import load_dotenv
from llm_provider_kit import GeminiTextModel
from llm_provider_kit import GEMINI_REFINE

# 載入環境變數
load_dotenv()

# 設定日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('srt_merge.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# API 配置
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')


class SRTSegment:
    """SRT 段落資料結構"""
    def __init__(self, index: int, start_time: float, end_time: float, text: str):
        self.index = index
        self.start_time = start_time
        self.end_time = end_time
        self.text = text.strip()
    
    def contains_time(self, time: float) -> bool:
        """檢查時間是否在此段落內"""
        return self.start_time <= time <= self.end_time


class SRTMergeProcessor:
    """SRT 格式合併處理器"""
    
    def __init__(self):
        """初始化處理器"""
        self.segments: List[SRTSegment] = []
        self.all_images: Dict[float, str] = {}  # {time_seconds: image_path}
        self.setup_api()
    
    def setup_api(self):
        """設定 Google Gemini API"""
        if GOOGLE_API_KEY:
            logger.info("Google Gemini API 設定完成")
    
    def parse_srt_time(self, time_str: str) -> float:
        """解析 SRT 時間格式 (HH:MM:SS,mmm) 為秒數"""
        time_str = time_str.strip()
        # 處理逗號（SRT 使用逗號分隔毫秒）
        time_str = time_str.replace(',', '.')
        
        parts = time_str.split(':')
        if len(parts) == 3:
            hours = int(parts[0])
            minutes = int(parts[1])
            seconds = float(parts[2])
            return hours * 3600 + minutes * 60 + seconds
        return 0.0
    
    def parse_srt_file(self, srt_path: str) -> List[SRTSegment]:
        """解析 SRT 檔案"""
        segments = []
        
        try:
            with open(srt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 分割成段落
            blocks = content.strip().split('\n\n')
            
            for block in blocks:
                lines = block.strip().split('\n')
                if len(lines) >= 3:
                    # 第一行是序號
                    index = int(lines[0])
                    
                    # 第二行是時間範圍
                    time_line = lines[1]
                    time_match = re.match(r'(\S+)\s*-->\s*(\S+)', time_line)
                    if time_match:
                        start_time = self.parse_srt_time(time_match.group(1))
                        end_time = self.parse_srt_time(time_match.group(2))
                        
                        # 剩餘行是文字
                        text = ' '.join(lines[2:])
                        
                        segment = SRTSegment(index, start_time, end_time, text)
                        segments.append(segment)
            
            logger.info(f"成功解析 SRT 檔案，共 {len(segments)} 個段落")
            return segments
            
        except Exception as e:
            logger.error(f"解析 SRT 檔案失敗: {e}")
            raise
    
    def parse_image_time(self, filename: str) -> Optional[float]:
        """從圖片檔名解析時間（如 slide_007_t2m39.0s.jpg）"""
        match = re.search(r't(\d+)m([\d.]+)s', filename)
        if match:
            minutes = int(match.group(1))
            seconds = float(match.group(2))
            return minutes * 60 + seconds
        return None
    
    def load_all_images(self, image_folders: List[str]) -> Dict[float, str]:
        """載入所有資料夾中的圖片"""
        all_images = {}
        
        for folder in image_folders:
            if not os.path.exists(folder):
                logger.warning(f"圖片資料夾不存在: {folder}")
                continue
            
            folder_name = os.path.basename(folder)
            image_count = 0
            
            # 支援的圖片格式
            patterns = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']
            
            for pattern in patterns:
                for img_path in glob.glob(os.path.join(folder, pattern)):
                    filename = os.path.basename(img_path)
                    time_sec = self.parse_image_time(filename)
                    
                    if time_sec is not None:
                        # 使用完整路徑
                        all_images[time_sec] = img_path
                        image_count += 1
            
            logger.info(f"從 {folder_name} 載入了 {image_count} 張圖片")
        
        logger.info(f"總共載入 {len(all_images)} 張圖片")
        return all_images
    
    def find_best_insertion_point(self, image_time: float, segments: List[SRTSegment]) -> int:
        """找到圖片的最佳插入位置"""
        # 策略：找到包含此時間或剛結束的段落
        for i, segment in enumerate(segments):
            if segment.contains_time(image_time):
                # 圖片時間在段落內，在段落後插入
                return i + 1
            elif i > 0 and segments[i-1].end_time <= image_time < segment.start_time:
                # 圖片時間在兩段落之間
                return i
        
        # 如果沒找到，根據最接近的時間插入
        if image_time < segments[0].start_time:
            logger.info(f"圖片時間 {self.format_time(image_time)} 在 SRT 開始之前，將插入到開頭")
            return 0
        else:
            logger.info(f"圖片時間 {self.format_time(image_time)} 在 SRT 結束之後，將插入到結尾")
            return len(segments)
    
    def merge_srt_with_images(self, srt_segments: List[SRTSegment], 
                             images: Dict[float, str],
                             slide_contents: List[Tuple[str, str]]) -> str:
        """合併 SRT 內容與圖片"""
        # 建立段落與圖片的對應關係
        segment_images = {}  # {segment_index: [image_paths]}
        used_images = set()  # 追蹤已使用的圖片
        
        # 為每張圖片找到最佳插入點
        for image_time, image_path in sorted(images.items()):
            insert_pos = self.find_best_insertion_point(image_time, srt_segments)
            if insert_pos not in segment_images:
                segment_images[insert_pos] = []
            segment_images[insert_pos].append(image_path)
            used_images.add(image_path)
        
        # 報告未使用的圖片
        all_images = set(images.values())
        unused_images = all_images - used_images
        if unused_images:
            logger.warning(f"⚠️  有 {len(unused_images)} 張圖片未被使用")
            for img in list(unused_images)[:5]:  # 只顯示前5張
                logger.warning(f"  - {os.path.basename(img)}")
            if len(unused_images) > 5:
                logger.warning(f"  ... 還有 {len(unused_images) - 5} 張")
        
        logger.info(f"✅ 成功分配 {len(used_images)} / {len(all_images)} 張圖片")
        
        # 如果有 Gemini API，使用它來潤飾內容
        if GOOGLE_API_KEY and slide_contents:
            return self.merge_with_gemini(srt_segments, segment_images, slide_contents)
        else:
            # 否則直接合併
            return self.simple_merge(srt_segments, segment_images)
    
    def simple_merge(self, segments: List[SRTSegment], 
                    segment_images: Dict[int, List[str]]) -> str:
        """簡單合併（不使用 AI）"""
        lines = []
        
        # 添加標題
        lines.append("# 演講內容與投影片整合")
        lines.append("")
        
        # 處理開頭的圖片（在第一個段落之前）
        if 0 in segment_images and len(segment_images[0]) > 0:
            lines.append("## 開場投影片")
            lines.append("")
            for img_path in segment_images[0]:
                lines.append(f"[IMAGE: {img_path}]")
                lines.append("")
            lines.append("---")
            lines.append("")
        
        # 處理主要內容
        lines.append("## 演講內容")
        lines.append("")
        
        for i, segment in enumerate(segments):
            # 添加段落前的圖片（除了第一個段落前的圖片，已經處理過）
            if i > 0 and i in segment_images:
                for img_path in segment_images[i]:
                    lines.append(f"[IMAGE: {img_path}]")
                    lines.append("")
            
            # 添加時間戳和內容
            start_time = self.format_time(segment.start_time)
            lines.append(f"[{start_time}] {segment.text}")
            lines.append("")
        
        # 處理結尾的圖片（在最後一個段落之後）
        if len(segments) in segment_images and len(segment_images[len(segments)]) > 0:
            lines.append("---")
            lines.append("")
            lines.append("## 結尾投影片")
            lines.append("")
            for img_path in segment_images[len(segments)]:
                lines.append(f"[IMAGE: {img_path}]")
                lines.append("")
        
        return '\n'.join(lines)
    
    def merge_with_gemini(self, segments: List[SRTSegment],
                         segment_images: Dict[int, List[str]],
                         slide_contents: List[Tuple[str, str]]) -> str:
        """使用 Gemini API 進行智能合併"""
        try:
            model = GeminiTextModel(GEMINI_REFINE, api_key=GOOGLE_API_KEY)
            
            # 準備 SRT 內容
            srt_text = self.simple_merge(segments, segment_images)
            
            # 準備投影片內容
            slides_text = "\n\n".join([f"=== {name} ===\n{content}" 
                                      for name, content in slide_contents])
            
            # 構建提示詞
            prompt = f"""請根據以下演講內容和投影片分析，創建一份整合的會議筆記。

重要任務：整合演講內容和投影片分析，創建詳細的會議筆記。

核心要求：
1. 演講內容作為主要架構
2. 從投影片 .md 檔案中提取以下資訊並整合：
   - 圖表中的具體數據和統計結果
   - 研究方法和實驗設計細節
   - 參考文獻和研究背景
   - 視覺元素的詳細描述
   - 投影片中提到但演講未詳述的要點

3. 整合方式：
   - 當演講者提到某研究時，從投影片補充具體數據
   - 當演講者描述結果時，加入投影片中的圖表說明
   - 使用「根據投影片顯示...」或「投影片中的數據指出...」來引入補充資訊
   - 將投影片的詳細數據自然融入演講內容中

4. 格式要求：
   - 不要在正文中包含時間戳記 [HH:MM:SS] 或 [MM:SS]
   - 重要數據用 **粗體** 標示
   - 在每個主要段落後，如果投影片有補充資訊，用「📊 投影片補充」標示
   - 保留所有的 [IMAGE: 路徑] 標記在適當位置，不要修改路徑
   - 使用繁體中文

5. 必須確保：
   - 投影片中的所有重要數據都被納入
   - 圖表分析被整合到相關段落
   - 研究方法、統計數據、臨床試驗結果等具體資訊都要包含
   - 形成真正的二合一筆記，而非分開的兩部分

演講內容（含圖片標記）：
{srt_text}

投影片詳細分析：
{slides_text}

請創建一份整合的筆記，確保投影片中的重要資訊都被適當地整合到演講內容中。"""
            
            response = model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            logger.error(f"Gemini API 處理失敗: {e}")
            # 失敗時返回簡單合併結果
            return self.simple_merge(segments, segment_images)
    
    def format_time(self, seconds: float) -> str:
        """格式化時間為 HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
    
    def save_markdown(self, content: str, output_path: str) -> str:
        """保存 Markdown 檔案"""
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown 檔案已保存: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"保存 Markdown 失敗: {e}")
            raise
    
    def markdown_to_docx(self, markdown_text: str, output_path: str) -> bool:
        """將 Markdown 轉換為 DOCX，確保圖片正確插入"""
        try:
            doc = Document()
            
            # 添加標題
            title = doc.add_heading('演講內容與投影片整合', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加日期
            date_para = doc.add_paragraph()
            date_para.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_paragraph()  # 空行
            
            # 處理內容
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    doc.add_paragraph()  # 空行
                    continue
                
                # 處理圖片標記 [IMAGE: path]
                if '[IMAGE:' in line:
                    matches = re.findall(r'\[IMAGE:\s*([^\]]+)\]', line)
                    for img_path in matches:
                        img_path = img_path.strip()
                        if os.path.exists(img_path):
                            try:
                                doc.add_paragraph()  # 空行
                                doc.add_picture(img_path, width=Inches(5.5))
                                doc.add_paragraph()  # 空行
                                logger.info(f"成功插入圖片: {os.path.basename(img_path)}")
                            except Exception as e:
                                logger.error(f"插入圖片失敗 {img_path}: {e}")
                                doc.add_paragraph(f"[圖片無法插入: {os.path.basename(img_path)}]")
                        else:
                            logger.warning(f"圖片不存在: {img_path}")
                            doc.add_paragraph(f"[圖片不存在: {os.path.basename(img_path)}]")
                
                # 處理標題
                elif line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    doc.add_heading(title_text, level=min(level, 4))
                
                # 處理一般段落
                else:
                    paragraph = doc.add_paragraph()
                    # 處理粗體等格式
                    self._add_formatted_text(paragraph, line)
            
            # 儲存文件
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logger.info(f"DOCX 文件已儲存: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"轉換為 DOCX 失敗: {e}")
            return False
    
    def _add_formatted_text(self, paragraph, text):
        """處理文字格式"""
        # 簡單處理，可以擴展支援更多格式
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗體
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # 普通文字
                paragraph.add_run(part)
    
    def process_files(self, srt_file: str, slide_folders: List[str], 
                     slide_files: List[str] = None, output_base: str = None) -> dict:
        """處理 SRT 檔案與投影片"""
        result = {
            'success': False,
            'srt_file': srt_file,
            'slide_folders': slide_folders,
            'markdown_file': None,
            'docx_file': None,
            'error': None
        }
        
        try:
            # 1. 解析 SRT 檔案
            self.segments = self.parse_srt_file(srt_file)
            
            # 2. 載入所有圖片
            self.all_images = self.load_all_images(slide_folders)
            
            # 3. 載入投影片內容（如果有提供）
            slide_contents = []
            if slide_files:
                for slide_file in slide_files:
                    if os.path.exists(slide_file):
                        with open(slide_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                        slide_contents.append((os.path.basename(slide_file), content))
            
            # 4. 合併內容
            merged_content = self.merge_srt_with_images(
                self.segments, self.all_images, slide_contents
            )
            
            # 5. 準備輸出路徑
            if not output_base:
                srt_path = Path(srt_file)
                output_base = srt_path.stem
            
            output_dir = Path(srt_file).parent
            
            # 6. 保存 Markdown
            markdown_path = output_dir / f"{output_base}_srt_merged.md"
            self.save_markdown(merged_content, str(markdown_path))
            result['markdown_file'] = str(markdown_path)
            
            # 7. 轉換為 DOCX
            docx_path = output_dir / f"{output_base}_srt_merged.docx"
            if self.markdown_to_docx(merged_content, str(docx_path)):
                result['docx_file'] = str(docx_path)
                result['success'] = True
            else:
                result['error'] = "DOCX 轉換失敗"
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"處理檔案失敗: {e}")
        
        return result


def main():
    """主函數"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='SRT 格式演講稿與投影片智能合併系統'
    )
    
    parser.add_argument('srt_file', help='SRT 字幕檔案路徑')
    parser.add_argument('slide_folders', nargs='+', 
                       help='投影片圖片資料夾路徑（可多個）')
    parser.add_argument('--slides', nargs='*',
                       help='投影片 Markdown 檔案（選用）')
    parser.add_argument('--output', '-o',
                       help='輸出檔案基礎名稱')
    
    args = parser.parse_args()
    
    print(f"\n=== SRT 演講稿與投影片合併系統 ===")
    print(f"SRT 檔案: {args.srt_file}")
    print(f"圖片資料夾數量: {len(args.slide_folders)}")
    print("處理中...\n")
    
    processor = SRTMergeProcessor()
    result = processor.process_files(
        args.srt_file, 
        args.slide_folders,
        args.slides,
        args.output
    )
    
    if result['success']:
        print("\n✅ 處理成功！")
        print(f"Markdown: {result['markdown_file']}")
        print(f"DOCX: {result['docx_file']}")
    else:
        print(f"\n❌ 處理失敗: {result['error']}")


if __name__ == "__main__":
    main()